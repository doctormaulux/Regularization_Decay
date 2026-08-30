"""
Build the updated paper as a Word document (.docx)
Reads results dynamically from CSV files.
"""
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import sys
import numpy as np
import pandas as _pd

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(SCRIPT_DIR)
FIGURES_DIR = os.path.join(SCRIPT_DIR, 'figures')

# Allow importing from project root
sys.path.insert(0, ROOT_DIR)
from analysis.stats_utils import (
    load_csv, get_val, welch_ttest, cohens_d, hedges_g, tost_equivalence,
    load_runs, paired_ttest, bootstrap_ci,
    benjamini_hochberg, bonferroni,
)

# Practical-equivalence margins, fixed BEFORE the analysis (REVIEWER-8). Every "no harm" /
# "matches within noise" statement in the manuscript is tested by TOST against these, not
# inferred from a failure to reject. Changing them post hoc would invalidate the claims.
EQUIV_MARGIN_PPL = 0.20   # perplexity points, language modelling
EQUIV_MARGIN_ACC = 0.25   # percentage points, classification


def equiv_verdict(bname, method_a, method_b, metric='test_ppl', margin=None, n=None):
    """TOST verdict for a pair of methods on a benchmark, as a phrase for the prose.

    Returns (verdict_text, equivalent_bool, p_tost). Keeps every 'no harm' sentence tied
    to a computation instead of to an author's reading of a non-significant p-value.
    """
    df = all_data[bname]
    margin = margin if margin is not None else (
        EQUIV_MARGIN_PPL if metric == 'test_ppl' else EQUIV_MARGIN_ACC)
    n = n if n is not None else N_SEEDS.get(bname, 5)
    am, asd = get_val(df, method_a, metric, 'mean'), get_val(df, method_a, metric, 'std')
    bm, bsd = get_val(df, method_b, metric, 'mean'), get_val(df, method_b, metric, 'std')
    if am != am or bm != bm:
        return 'not in this roster', False, float('nan')
    if metric == 'test_acc':          # CSVs store fractions; the margin is in percentage points
        am, asd, bm, bsd = am * 100, asd * 100, bm * 100, bsd * 100
    eq, p_tost, lo, hi = tost_equivalence(am, asd, bm, bsd, margin=margin, n=n)
    if eq:
        return (f'practically equivalent (TOST p = {p_tost:.3f}, whole CI inside '
                f'±{margin:g})', True, p_tost)
    _, p_w = welch_ttest(am, asd, bm, bsd, n=n)
    if p_w < 0.05:
        return (f'a real difference of {am - bm:+.3f} beyond the ±{margin:g} margin '
                f'(p = {p_w:.3g})', False, p_tost)
    return (f'inconclusive — neither different nor demonstrably equivalent at n = {n}',
            False, p_tost)

# ── Load results from CSVs ──
# Competitor family for the statistical scoreboard: WD-tuned (decoupled weight
# decay) IS a competitor; Tau(alpha=0) is an ablation of \u03c4(w) and is therefore
# presented (TABLE_METHODS) but never counted as a win/loss.
METHODS_ORDER = ['Baseline', 'L1', 'L2', 'ElasticNet', 'SCAD', 'MCP', 'LSP', 'WD-tuned', '\u03c4(w)']
TABLE_METHODS = ['Baseline', 'L1', 'L2', 'ElasticNet', 'SCAD', 'MCP', 'LSP',
                 'WD-tuned', 'Tau(alpha=0)', '\u03c4(w)']

# Seeds per benchmark (ExperimentTracker runs), FALLBACK only: overridden below from the
# per-seed JSONs (n = 5 on the single-scale benchmarks, 3 on tiny/medium/WT-103/SmolLM2,
# 10 at 66M); used by Welch t-tests / Cohen d.
N_SEEDS = {
    'gpt2_tiny_wikitext': 3, 'gpt2_medium_wikitext': 3, 'gpt2_large_wikitext': 3,
    'gpt2_wt103': 3,
    'smollm2_wikitext': 3,
}

# Single-scale benchmarks of Table 1 / Figure 1. Scope restricted on 2026-08-29: the
# pretrained SST-2 encoders and the 0.5B / 1B pretrained WikiText-2 LMs are no longer part
# of the paper; only these eight remain (order = column order of Table 1).
BENCHMARKS = [
    ('sin_regression', 'test_mse', 'min', 1e3, 'MSE\u00d710\u207b\u00b3'),
    ('complex_regression', 'test_mse', 'min', 10, 'MSE\u00d710\u207b\u00b9'),
    ('mnist', 'test_acc', 'max', 100, '%'),
    ('cifar', 'test_acc', 'max', 100, '%'),
    ('vit_cifar', 'test_acc', 'max', 100, '%'),
    ('bert_sst2', 'test_acc', 'max', 100, '%'),
    ('gpt2_wikitext', 'test_ppl', 'min', 1, 'PPL'),
    ('smollm2_wikitext', 'test_ppl', 'min', 1, 'PPL'),
]

# From-scratch GPT-2 scale sweep (core-6 roster, n=3 seeds, 12-epoch budget).
# Kept out of BENCHMARKS so Table 1 / Figure 1 keep the single-scale layout;
# included in the statistical family and presented in Section 5.6 (Table 2).
SCALE_BENCHMARKS = [
    ('gpt2_tiny_wikitext', 'test_ppl', 'min', 1, 'PPL'),
    ('gpt2_medium_wikitext', 'test_ppl', 'min', 1, 'PPL'),
    ('gpt2_large_wikitext', 'test_ppl', 'min', 1, 'PPL'),
    ('gpt2_wt103', 'test_ppl', 'min', 1, 'PPL'),
]

def compute_stats(df, metric, mode, n=5):
    """Compute tau(w) vs all others: rank, wins, sig_wins, losses, sig_losses."""
    tau_mean = get_val(df, '\u03c4(w)', metric, 'mean')
    tau_std = get_val(df, '\u03c4(w)', metric, 'std')
    methods = [m for m in METHODS_ORDER if m in df['method'].values]
    # Compute rank
    means = [(m, get_val(df, m, metric, 'mean')) for m in methods]
    if mode == 'max':
        ranked = sorted(means, key=lambda x: -x[1])
    else:
        ranked = sorted(means, key=lambda x: x[1])
    rank = next(i+1 for i, (m, _) in enumerate(ranked) if m == '\u03c4(w)')
    total = len(methods)
    # Wins/losses
    wins, sig_wins, losses, sig_losses = 0, 0, 0, 0
    details = {}
    for m in methods:
        if m == '\u03c4(w)':
            continue
        m_mean = get_val(df, m, metric, 'mean')
        m_std = get_val(df, m, metric, 'std')
        _, p = welch_ttest(tau_mean, tau_std, m_mean, m_std, n=n)
        d = cohens_d(tau_mean, tau_std, m_mean, m_std, n=n)
        if mode == 'max':
            better = tau_mean > m_mean
        else:
            better = tau_mean < m_mean
        details[m] = {'p': p, 'd': d, 'better': better, 'tau_mean': tau_mean, 'm_mean': m_mean}
        if better:
            wins += 1
            if p < 0.05:
                sig_wins += 1
        else:
            losses += 1
            if p < 0.05:
                sig_losses += 1
    return {
        'rank': rank, 'total': total,
        'wins': wins, 'sig_wins': sig_wins,
        'losses': losses, 'sig_losses': sig_losses,
        'details': details,
    }

# Per-benchmark n is read from the per-seed JSON where one exists (the dictionary above is
# only the fallback): the 66M sweep member was re-run at n = 10 after the early-stopping
# fix, and every Welch / Cohen computation must use the n that was actually run.
for _bn, _metric, _mode, _sc, _u in BENCHMARKS + SCALE_BENCHMARKS:
    _runs = load_runs(_bn, 'τ(w)', _metric)
    if _runs:
        N_SEEDS[_bn] = len(_runs)

# Load all data
all_data = {}
all_stats = {}
for bname, metric, mode, scale, _ in BENCHMARKS + SCALE_BENCHMARKS:
    df = load_csv(bname)
    all_data[bname] = df
    all_stats[bname] = compute_stats(df, metric, mode, n=N_SEEDS.get(bname, 5))

# 66M scope x adaptivity decomposition, needed as early as the Abstract (see the table in
# Section 6.1). Defined here so the Abstract and the Discussion read the same numbers.
_d66 = all_data['gpt2_large_wikitext']
_dv = lambda m: get_val(_d66, m, 'test_ppl', 'mean')

# ── Sweep accessors (2026-08-30 numeric pass) ───────────────────────────────
# One NaN-safe accessor per scale so that the Abstract, Section 5.6, Section 6.1 and the
# Conclusion read the same regenerated numbers. A method missing from a provisional CSV
# (the 66M L2 / ElasticNet rows, SCAD at 7M, SCAD/MCP/LSP on ViT) renders as '—' and never
# crashes the build.
_SCALE_KEYS = {'tiny': 'gpt2_tiny_wikitext', 'small': 'gpt2_wikitext',
               'medium': 'gpt2_medium_wikitext', 'large': 'gpt2_large_wikitext',
               'wt103': 'gpt2_wt103', 'smollm2': 'smollm2_wikitext'}

def _pv(bname, method, stat='mean', metric='test_ppl'):
    return get_val(all_data[bname], method, metric, stat)

def _gap(bname, method):
    """method − τ(w) in PPL on `bname` (positive = τ(w) better); NaN if absent."""
    return _pv(bname, method) - _pv(bname, 'τ(w)')

def _pct(bname, method='Baseline'):
    """τ(w) relative to `method` in % of the competitor's PPL (negative = τ(w) better)."""
    b = _pv(bname, method)
    return (_pv(bname, 'τ(w)') - b) / b * 100

def _and_join(items):
    """'a', 'a and b', 'a, b and c' — for lists of names in the prose."""
    items = list(items)
    if not items:
        return ''
    return items[0] if len(items) == 1 else ', '.join(items[:-1]) + ' and ' + items[-1]

def _sgn(v, nd=2):
    """Signed fixed-point value that never prints '-0.00'."""
    if v != v:
        return '—'
    return f'{0.0:.{nd}f}' if round(v, nd) == 0 else f'{v:+.{nd}f}'

def _fmt(v, nd=2, unit='', sign=False):
    """'—' for a missing value, otherwise a fixed-point string."""
    if v != v:
        return '—'
    return f"{v:{'+' if sign else ''}.{nd}f}{unit}"

def _alpha_gap(bname):
    """Pre-registered primary contrast, paired by seed: τ(w) − Tau(α = 0) on `bname`.
    Returns (t, p, mean_diff, n); (nan, nan, nan, 0) if either series is missing."""
    _ta = load_runs(bname, 'τ(w)', 'test_ppl')
    _aa = load_runs(bname, 'Tau(alpha=0)', 'test_ppl')
    if not _ta or not _aa:
        return float('nan'), float('nan'), float('nan'), 0
    return paired_ttest(_ta, _aa)

_AG = {s: _alpha_gap(k) for s, k in _SCALE_KEYS.items()}   # scale -> (t, p, Δ, n)

def _pfmt(p):
    """p-value for the prose: exact to four decimals, or a bound when it underflows that."""
    if p != p:
        return 'p = —'
    return 'p < 0.0001' if p < 1e-4 else f'p = {p:.4f}'

_dv7 = lambda m: _pv('gpt2_wikitext', m)
_dv18 = lambda m: _pv('gpt2_medium_wikitext', m)
_dv2 = lambda m: _pv('gpt2_tiny_wikitext', m)
_dvw3 = lambda m: _pv('gpt2_wt103', m)

# Trainable-parameter counts (millions) as recorded in the result files.
_MP = {s: get_val(all_data[k], 'Baseline', 'total_params', 'mean') / 1e6
       for s, k in _SCALE_KEYS.items()}

# Best validation epoch (mean over seeds) of each method, as recorded by the tracker:
# the overfitting-pressure reading of Section 5.6 rests on where the Baseline's best
# epoch falls relative to the epoch budget.
_EPOCH_BUDGET = {'tiny': 12, 'small': 30, 'medium': 12, 'large': 12, 'wt103': 8}
def _best_epoch(scale, method):
    return get_val(all_data[_SCALE_KEYS[scale]], method, 'convergence_epoch', 'mean')

# ── Data-quantity arm (Section 5.6.1): the 66M model on 25% / 50% / 100% of WikiText-2 ──
# Baseline / Tau(alpha=0) / τ(w), n = 5 seeds, hyperparameters fixed at the 100% optimum,
# optimizer steps held constant; the 100% column is the main 66M run. Everything that reads
# these frames (Abstract clause, Section 5.6.1 table and prose) is guarded on _HAVE_DATA_ARM
# so a missing file can never crash the build or invent numbers.
_DATA_ARM_FILES = {frac: os.path.join(ROOT_DIR, 'new results',
                                      f'gpt2_large_wikitext_standardized_results_data{frac}.csv')
                   for frac in (25, 50)}
_HAVE_DATA_ARM = all(os.path.exists(_p) for _p in _DATA_ARM_FILES.values())
_DA_DFS = ({25: _pd.read_csv(_DATA_ARM_FILES[25]), 50: _pd.read_csv(_DATA_ARM_FILES[50]),
            100: all_data['gpt2_large_wikitext']} if _HAVE_DATA_ARM else {})

def _da_val(frac, method, stat='mean'):
    return get_val(_DA_DFS[frac], method, 'test_ppl', stat)

def _da_pct(frac, method='τ(w)'):
    """method − Baseline at `frac`% of the corpus, in % of the Baseline (negative = better)."""
    _b = _da_val(frac, 'Baseline')
    return (_da_val(frac, method) - _b) / _b * 100

def _da_budget(frac):
    """Epoch budget at `frac`% of the corpus (optimizer steps held constant)."""
    return _EPOCH_BUDGET['large'] * 100 // frac

def _da_best_ep(frac, method):
    return get_val(_DA_DFS[frac], method, 'convergence_epoch', 'mean')

# 66M scope × adaptivity factorial (Section 4.2): the four cells and the four contrasts.
_d_base, _d_wd = _dv('Baseline'), _dv('WD-tuned')
_d_a0, _d_tau, _d_ad = _dv('Tau(alpha=0)'), _dv('τ(w)'), _dv('Tau(AdamW-scope)')
_d_total = _d_base - _d_tau                 # Baseline → τ(w), positive = improvement
_eff_scope_const = _d_a0 - _d_wd            # WD-tuned → Tau(α=0): scope at constant profile
_eff_adapt_adamw = _d_ad - _d_wd            # WD-tuned → Tau(AdamW-scope): adaptivity at AdamW scope
_eff_adapt_tau = _d_tau - _d_a0             # Tau(α=0) → τ(w): adaptivity at τ scope
_eff_scope_adapt = _d_tau - _d_ad           # Tau(AdamW-scope) → τ(w): scope at adaptive profile

# ── Instrumented trajectories (Section 3 worked example, Section 6.2) ───────
# results/instrumentation/gpt2_<scale>_<method>_seed<seed>.json, written by the
# --instrument runs. Files from superseded runs share the same names as the current ones
# (the pre-fix 66M L2 / ElasticNet runs, and one pre-fix seed of each τ-family method), so
# a run is kept ONLY if its hyperparameters equal the benchmark JSON's best_hyperparams for
# that method. Methods without a tuned configuration in the JSON (66M L2 / ElasticNet
# until tonight's run lands) are therefore absent, never stale.
import glob as _glob
import json as _json
_INST_DIR = os.path.join(ROOT_DIR, 'results', 'instrumentation')

def _best_hp(bname):
    _p = os.path.join(ROOT_DIR, 'new results', f'{bname}_standardized_results.json')
    if not os.path.exists(_p):
        return {}
    with open(_p, encoding='utf-8') as _fh:
        return _json.load(_fh).get('best_hyperparams') or {}

def _same_hp(a, b):
    if set(a) != set(b):
        return False
    return all(abs(float(a[k]) - float(b[k])) <= 1e-9 * max(1.0, abs(float(b[k])))
               for k in a)

def _inst_load(scale, bname):
    """{method: [trajectory per seed]} restricted to the final tuned configuration."""
    _best = _best_hp(bname)
    _out = {}
    for _path in sorted(_glob.glob(os.path.join(_INST_DIR, f'gpt2_{scale}_*_seed*.json'))):
        with open(_path, encoding='utf-8') as _fh:
            _d = _json.load(_fh)
        if _d.get('benchmark') != f'gpt2_{scale}' or _d.get('phase', 'eval') != 'eval':
            continue
        _m = _d.get('method')
        if _m not in _best or not _same_hp(_d.get('hyperparams') or {}, _best[_m]):
            continue
        if _d.get('trajectory'):
            _out.setdefault(_m, []).append(_d['trajectory'])
    return _out

def _inst_summary(scale, bname):
    """Per-method weight-distribution summary (mean over seeds): first-to-last-epoch change
    of median |w| and max |w| in %, final median and max, near-zero fraction, n seeds."""
    _by = _inst_load(scale, bname)
    _s = {}
    for _m, _trajs in _by.items():
        def _at(key, idx):
            _v = [t[idx].get(key) for t in _trajs if t[idx].get(key) is not None]
            return sum(_v) / len(_v) if _v else float('nan')
        _med0, _med1 = _at('median_abs_w', 0), _at('median_abs_w', -1)
        _max0, _max1 = _at('max_abs_w', 0), _at('max_abs_w', -1)
        _ne = min(len(t) for t in _trajs)
        _val = [_at('val_ppl', i) for i in range(_ne)]
        _trn = [_at('train_ppl', i) for i in range(_ne)]
        _bi = min(range(_ne), key=lambda i: _val[i])
        _s[_m] = {'d_med': 100 * (_med1 - _med0) / _med0, 'd_max': 100 * (_max1 - _max0) / _max0,
                  'med': _med1, 'max': _max1, 'frac0': _at('frac_below_thr', -1),
                  'best_ep': _bi + 1, 'best_val': _val[_bi], 'gap': _val[_bi] - _trn[_bi],
                  'last_val': _val[-1], 'n_ep': _ne, 'n': len(_trajs)}
    return _s

_INST = {s: _inst_summary(s, k) for s, k in
         (('medium', 'gpt2_medium_wikitext'), ('large', 'gpt2_large_wikitext'),
          ('wt103', 'gpt2_wt103'), ('tiny', 'gpt2_tiny_wikitext'))}
_iv = lambda scale, m, key: _INST.get(scale, {}).get(m, {}).get(key, float('nan'))

def _welch2(mean_a, std_a, n_a, mean_b, std_b, n_b):
    """Welch's t-test from summary statistics with unequal n (stats_utils assumes n_a = n_b)."""
    from scipy import stats as _st
    _se2 = std_a ** 2 / n_a + std_b ** 2 / n_b
    if _se2 == 0:
        return 0.0, 1.0
    _t = (mean_a - mean_b) / np.sqrt(_se2)
    _df = _se2 ** 2 / ((std_a ** 2 / n_a) ** 2 / (n_a - 1) + (std_b ** 2 / n_b) ** 2 / (n_b - 1))
    return float(_t), float(2 * (1 - _st.t.cdf(abs(_t), _df)))

# ── STAT-1: family-wide multiple-comparison correction ─────────────────────
# The full comparison family is all τ(w)-vs-competitor pairs across all
# benchmarks. Apply Benjamini–Hochberg (FDR, q=0.05) AND Bonferroni (FWER,
# α=0.05) once to the WHOLE family, then count significant wins and losses
# using the SAME corrected mask (symmetric: a 0.05 nominal level on wins
# requires the same level on losses, otherwise the "never significantly
# worse" claim is rhetorically inflated by a tighter bar on losses).
# Raw (uncorrected) counts are kept under *_raw for transparency.
_pvals_flat = []
_keys_flat = []  # (bname, competitor_method)
for _bn, _st in all_stats.items():
    for _m, _d in _st['details'].items():
        _keys_flat.append((_bn, _m))
        _pvals_flat.append(_d['p'])

_bh_mask = benjamini_hochberg(_pvals_flat, q=0.05)
_bonf_mask = bonferroni(_pvals_flat, alpha=0.05)

for (_bn, _m), _bh, _bf in zip(_keys_flat, _bh_mask, _bonf_mask):
    _d = all_stats[_bn]['details'][_m]
    _d['sig_bh'] = bool(_bh)
    _d['sig_bonf'] = bool(_bf)

for _bn, _st in all_stats.items():
    _det = _st['details']
    _st['sig_wins_raw']    = _st['sig_wins']
    _st['sig_losses_raw']  = _st['sig_losses']
    _st['sig_wins']        = sum(1 for d in _det.values() if d.get('sig_bh') and d['better'])
    _st['sig_losses']      = sum(1 for d in _det.values() if d.get('sig_bh') and not d['better'])
    _st['sig_wins_bonf']   = sum(1 for d in _det.values() if d.get('sig_bonf') and d['better'])
    _st['sig_losses_bonf'] = sum(1 for d in _det.values() if d.get('sig_bonf') and not d['better'])

_sw  = sum(s['sig_wins'] for s in all_stats.values())
_sl  = sum(s['sig_losses'] for s in all_stats.values())
_swB = sum(s['sig_wins_bonf'] for s in all_stats.values())
_slB = sum(s['sig_losses_bonf'] for s in all_stats.values())
_swR = sum(s['sig_wins_raw'] for s in all_stats.values())
_slR = sum(s['sig_losses_raw'] for s in all_stats.values())
print(f"[OK] Loaded results from CSVs (family of {len(_pvals_flat)} t-tests). "
      f"Wins/losses — BH-FDR(q=0.05): {_sw}/{_sl} | "
      f"Bonferroni(α=0.05): {_swB}/{_slB} | "
      f"raw(p<0.05, uncorrected): {_swR}/{_slR}")

doc = Document()

def add_figure(image_path, caption, width_inches=6.0):
    """Insert a figure with caption into the document."""
    full_path = os.path.join(FIGURES_DIR, image_path)
    if not os.path.exists(full_path):
        p = doc.add_paragraph(f"[FIGURE NOT FOUND: {image_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full_path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True

# ── Style configuration ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)

# ── Helper functions ──
def add_bold_paragraph(text, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_text(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def set_cell_shading(cell, color):
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')

def add_table_row(table, cells, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg_color:
            set_cell_shading(cell, bg_color)

# ============================================================================
# TITLE PAGE
# ============================================================================

add_bold_paragraph("Title:", size=11, space_before=12)
# Retitled (REVIEWER, "come riposizionerei il contributo"). The previous title promised a
# general mechanism for deep networks and transformers; the evidence supports a specific
# benefit in from-scratch, over-capacity autoregressive LMs. "Physically-Inspired" is also
# dropped: "time constant" is an analogy, not a derivation, and the reviewer is right that
# keeping it invites a claim the paper does not make.
# Retitled again (2026-08-30, regenerated results): the previous title sold magnitude
# adaptivity as the mechanism; the 2x2 factorial and the paired primary contrast show the
# gain lies in the decay's scope and schedule, with adaptivity a second-order refinement.
# TITLE: alternative: "Decoupled Fair Weight Decay Under Overfitting Pressure: A Weight-Only,
#   Schedule-Independent Decay Beats Tuned AdamW Weight Decay in From-Scratch Language
#   Models, and Magnitude Adaptivity Is a Second-Order Refinement"
# Retitled 2026-08-31 for the Neural Networks ~80-character title limit: the subtitle
# "Decay Scope and Schedule, Not Magnitude Adaptivity, Carry the Gain" (the mechanism
# claim) moves out of the title; it survives verbatim in the Highlights, the Abstract's
# factorial sentence and Section 6.1's heading.
add_bold_paragraph(
    "Decoupled Fair Weight Decay for Over-Capacity Autoregressive Language Models",
    size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12
)

add_bold_paragraph("Authors:", size=11)
add_text("Giuseppe Maulucci\u00b9\u002a (ORCID: 0000-0002-2154-319X), Tommaso Marchetti\u00b9, Marco De Spirito\u00b9")

add_bold_paragraph("Affiliations:", size=11)
add_text("\u00b9 Università Cattolica del Sacro Cuore, Rome, Italy")

add_bold_paragraph("Corresponding author:", size=11)
add_text("\u002a Giuseppe Maulucci")
add_text("Email: giuseppe.maulucci@unicatt.it")
add_text("ORCID: https://orcid.org/0000-0002-2154-319X")

# ============================================================================
# REFRAME STATUS — pivot to the restricted from-scratch AR-LM claim
# ----------------------------------------------------------------------------
# 2026-07-06: title / highlights / abstract / intro / §6.1 / contributions reframed from
#   "universal choice for transformer training" to "mechanism for from-scratch
#   autoregressive-LM training; graceful degradation to tuned decoupled weight decay
#   elsewhere"; the false identical-to-Baseline claim removed (SmolLM2 τ(w) differs
#   from Baseline with real seed-to-seed variance).
# 2026-07-12: scale sweep (tiny/medium/large) integrated as SCALE_BENCHMARKS + Section 5.6
#   (Table 2, Figure 2B); WD-tuned added to the competitor family; per-benchmark n
#   (N_SEEDS) used in every Welch / Cohen computation.
# 2026-08-29: scope restricted to the 8 single-scale benchmarks in BENCHMARKS plus the
#   3-point scale sweep and the WikiText-103 confirmation (12 benchmarks). The pretrained
#   SST-2 encoders, the 0.5B / 1B pretrained WikiText-2 LMs and every post-training
#   magnitude-pruning result (former §5.11 / §6.4, Tables 6 and 9) are gone; §6.5–6.8
#   became §6.4–6.7 and Tables 7/8 became 6/7. The 'large' sweep member has 65.6M
#   trainable parameters (tied embeddings) and is labelled 66M; 'small' has 7.4M.
# COUNTS NOTE: every family-wide tally — number of t-tests, BH-FDR and Bonferroni win/loss
#   counts, practical win/tie/loss counts — is computed at build time from the CSVs in
#   `new results/` and injected through f-strings (Highlights, Abstract, contributions,
#   §5.7, §6.1). Do not hard-code these numbers anywhere in the prose or in this comment;
#   the build log prints the current values.
# 2026-08-30: numeric pass on the REGENERATED results. Every CSV was rebuilt with the
#   verified best-epoch pipeline (the previous pipeline silently reported the last epoch,
#   which inflated every regularizer's margin over the unregularized Baseline). The claim is
#   reframed accordingly: the gain belongs to the decoupled, weight-only, schedule-independent
#   τ family (Tau(α=0) and τ(w)) under overfitting pressure; magnitude adaptivity is a
#   second-order, scale-dependent refinement (the paired α-gap flips sign at 7M). Title,
#   highlights, abstract, contributions, §5.5–5.10, §6.1–6.7 and §7 now compute every
#   number from the CSVs / per-seed JSONs at build time.
# 2026-08-31: rosters completed (66M L2 / ElasticNet, 7M SCAD, ViT SCAD/MCP/LSP all landed);
#   WD-tuned re-tuned over the extended [1e-6, 10] range after the 0.1 bound proved binding
#   on the sweep; the data-quantity arm (Section 5.6.1) landed and the overfitting-pressure
#   claim was softened accordingly (pressure marks where the gain appears, it does not scale
#   it). The optional WD-tuned(weights) cell renders in Table 6 only if present in the CSV.
# Remaining: re-export docs/articolo.pdf.
# ============================================================================

# ── Highlights ──
doc.add_page_break()
add_bold_paragraph("Highlights", size=13, space_before=12)
_h_gap7 = _gap('gpt2_wikitext', 'WD-tuned')
_h_gap18 = _gap('gpt2_medium_wikitext', 'WD-tuned')
_h_gap66 = _gap('gpt2_large_wikitext', 'WD-tuned')
highlights = [
    "τ-decay: a decoupled, weight-only, schedule-independent Fair weight decay, "
    "w ← w − ρ·w/(1+|w|/δ), one line after optimizer.step, two identifiable parameters.",
    f"Under overfitting pressure it beats PSO-tuned AdamW weight decay by {_h_gap7:.1f} PPL at 7M "
    f"and {_h_gap66:.1f} PPL at 66M ({abs(_pct('gpt2_wikitext')):.0f}% and "
    f"{abs(_pct('gpt2_large_wikitext')):.0f}% below the unregularized baseline), by "
    f"{_h_gap18:.1f} PPL at 18M and not at all at 2M.",
    "The explanatory variable is overfitting pressure (capacity × epochs relative to data), "
    "not model size: the gain appears where the Baseline's best epoch precedes the budget — "
    "though pressure marks where the gain appears rather than scaling it: at fixed capacity "
    "and hyperparameters the margin grows with the amount of data (data-quantity arm).",
    f"A scope × adaptivity factorial at 66M (n = {N_SEEDS['gpt2_large_wikitext']}) attributes the "
    f"margin to the decay's scope and schedule ({abs(_eff_scope_const):.1f} PPL); magnitude "
    f"adaptivity is worth {_AG['large'][2]:+.2f} PPL at 66M, {_AG['medium'][2]:+.2f} at 18M and "
    f"{_AG['small'][2]:+.2f} at 7M (negative = better) — a second-order, scale-dependent refinement.",
    "Huber, pseudo-Huber and log-cosh decays reach parity with the Fair profile: the saturation "
    "profile is immaterial, the decoupled weight-only decay is what matters.",
    f"Elsewhere the family degrades gracefully to parity; no-harm claims are established by "
    f"equivalence tests, and the {_sl} significant loss{'es' if _sl != 1 else ''} after "
    f"Benjamini–Hochberg correction {'are' if _sl != 1 else 'is'} named.",
]
for h in highlights:
    add_bullet(h)

# ── Abstract ──
add_bold_paragraph("Abstract", size=13, space_before=18)

# Condensed 2026-08-31 to one paragraph for the Neural Networks concise-abstract
# requirement (~150 words for full articles). The three-paragraph version lives in git
# history; every number dropped here still appears in §5.5–5.10 / §6.1. Word count is
# checked at build time below — keep it ≤ 155.
_abstract_txt = (
    f"We study τ-decay, a decoupled weight-decay update w ← w − ρ·w/(1 + |w|/δ) applied "
    f"after the optimizer step to weight tensors only, at a schedule-independent rate; its "
    f"implicit penalty is the Fair function of robust statistics. Against seven "
    f"regularizers (L1, L2, ElasticNet, SCAD, MCP, the log-sum penalty and PSO-tuned AdamW "
    f"weight decay) on 12 benchmarks with best-epoch early stopping and pre-specified "
    f"equivalence margins, the finding is regime-specific: in from-scratch autoregressive "
    f"language modelling under overfitting pressure, τ-decay beats tuned AdamW weight "
    f"decay by {_h_gap7:.1f} PPL at 7M and {_h_gap66:.1f} PPL at 66M "
    f"({abs(_pct('gpt2_wikitext')):.0f}% and {abs(_pct('gpt2_large_wikitext')):.0f}% below "
    f"the unregularized baseline); the gain tracks overfitting pressure, not model size. "
    f"A 66M scope × adaptivity factorial attributes the margin to decay scope and "
    f"schedule; magnitude adaptivity is a second-order, scale-dependent refinement, and "
    f"Huber, pseudo-Huber and log-cosh decays match the Fair profile. Elsewhere the family "
    f"degrades gracefully to parity under family-wide Benjamini–Hochberg correction."
)
add_text(_abstract_txt)
_ab_wc = len(_abstract_txt.split())
assert _ab_wc <= 155, f"Abstract too long for Neural Networks: {_ab_wc} words"
print(f"Abstract word count: {_ab_wc}")

_ab_sm_eq = equiv_verdict('smollm2_wikitext', 'τ(w)', 'WD-tuned')[1]

_ab_losses = [(bn, m, d) for bn, s in all_stats.items()
              for m, d in s['details'].items() if d.get('sig_bh') and not d['better']]
_AB_LABEL = {'sin_regression': 'sin(x)', 'complex_regression': 'Friedman', 'mnist': 'MNIST',
             'cifar': 'CIFAR-10 CNN', 'vit_cifar': 'ViT-CIFAR', 'bert_sst2': 'BERT-tiny',
             'gpt2_wikitext': 'GPT-2 7M', 'smollm2_wikitext': 'SmolLM2',
             'gpt2_tiny_wikitext': 'GPT-2 2M', 'gpt2_medium_wikitext': 'GPT-2 18M',
             'gpt2_large_wikitext': 'GPT-2 66M', 'gpt2_wt103': 'GPT-2 124M/WikiText-103'}
_ab_loss_txt = _and_join(f"{m} on {_AB_LABEL.get(bn, bn)}" for bn, m, _ in _ab_losses)

add_bold_paragraph("Keywords", size=11, space_before=12)
add_text(
    "regularization; weight decay; decoupled weight decay; language models; transformers; "
    "overfitting; robust regularization; Fair penalty; Huber decay; adaptive shrinkage",
    italic=True
)

# ============================================================================
# 1. INTRODUCTION
# ============================================================================

doc.add_page_break()
doc.add_heading("1. Introduction", level=1)

add_text(
    "Regularization plays a fundamental role in deep learning by shaping the generalization "
    "properties and complexity of neural networks. Standard convex penalties such as L1 and L2 "
    "remain widely used, yet they present intrinsic limitations. L1 promotes sparsity but "
    "introduces non-smooth shrinkage that may destabilize optimization, especially in large "
    "models. L2 applies uniform shrinkage regardless of weight magnitude and is therefore "
    "ineffective for inducing sparsity. Elastic Net combines the two but retains the "
    "magnitude-independent structure of the underlying penalties."
)

add_text(
    "Non-convex penalties, including SCAD (Fan & Li, 2001), MCP (Zhang, 2010) and the log-sum "
    "penalty (Candès et al., 2008), aim to reduce the bias introduced by convex shrinkage while "
    "allowing selective sparsification. Their theoretical appeal, however, is undermined by "
    "practical difficulties in deep learning. Their penalties are non-convex and defined "
    "piecewise; the derivative is designed to be continuous across the pieces, aside from the "
    "usual non-differentiability at the origin inherited from the absolute value. What causes "
    "trouble is the non-convexity itself: the flat regions introduce additional stationary "
    "points, which in our runs produced oscillatory behaviour, premature weight collapse or "
    "inconsistent convergence across seeds."
)

add_text(
    "At the same time, the growing prevalence of neural networks deployed on mobile devices, "
    "embedded systems and in large-scale pretrained models demands regularization strategies that "
    "promote sparsity in a stable and predictable manner. The rise of transformer architectures\u2014"
    "from BERT (Devlin et al., 2019) and GPT-2 (Radford et al., 2019) to vision transformers "
    "(Dosovitskiy et al., 2021)\u2014has introduced networks with hundreds of millions of parameters, "
    "amplifying the need for effective regularization that scales across architectures."
)

add_text(
    "This work studies τ-decay, a decoupled weight-decay update, and characterises the regime "
    "in which it pays: autoregressive language models trained from scratch under overfitting "
    "pressure, where a decoupled, weight-only, schedule-independent decay outperforms tuned "
    f"AdamW weight decay by {min(_h_gap7, _h_gap18, _h_gap66):.1f}–"
    f"{max(_h_gap7, _h_gap18, _h_gap66):.1f} perplexity points — about "
    f"{abs(_gap('gpt2_large_wikitext', 'WD-tuned')) / _dv('WD-tuned') * 100:.0f}% of the tuned "
    "competitor's perplexity at the scales where the pressure is highest. "
    "Magnitude adaptivity — the property that gives the "
    "method its name — turns out to be a second-order refinement of that decay rather than "
    "its source, and we say so. On pretrained models and on the other architectures tested, "
    "the family degrades gracefully to parity with tuned decoupled weight decay rather than "
    "acting as a universal default. Each weight evolves according to a time constant "
    "τ(w) = τ₀ + α|w| that depends linearly on its current magnitude: small weights decay "
    "faster than large ones. The resulting update is applied after the standard optimizer "
    "step, is compatible with SGD and Adam alike, and adds two identifiable hyperparameters: "
    "a maximum relative decay rate ρ and a magnitude scale δ at which the shrinkage "
    "saturates. (The three symbols η, τ₀ and α of the defining expression are not "
    "independent — the update is invariant under a common rescaling of all three — so the "
    "method is genuinely two-parameter; Section 3 makes this explicit.)"
)

add_text(
    "We are explicit about the scope of the claim. The implicit penalty behind this update is not "
    "new: it is the Fair function of robust statistics (Fair, 1974), as we derive in Section 3.1, "
    "and the broader idea of deriving a decoupled decay from a robust, saturating loss is shared "
    "with recent work on Huber decay for language-model pre-training (Guo & Fan, 2025). We do not "
    "claim a new family of regularizers. What we contribute is a smooth, decoupled, "
    "magnitude-adaptive member of that family \u2014 the Fair variant, whose shrinkage is analytic "
    "everywhere rather than piecewise \u2014 its post-optimizer implementation, a closed-form implicit "
    "penalty that a gradient-driven adaptive decay cannot offer, and, principally, the empirical "
    "characterisation of the regime in which it pays: over-capacity autoregressive language models "
    "trained from scratch."
)

add_text(
    f"Our contribution is threefold. First, a head-to-head evaluation of τ-decay against seven "
    f"competing regularizers (L1, L2, ElasticNet, SCAD, MCP, log-sum penalty, decoupled weight "
    f"decay) across 12 benchmarks spanning four architectural families — feedforward networks, "
    f"convolutional networks, a vision transformer and language models from 1.7M to 135M "
    f"parameters — with verified best-epoch early stopping, equivalence tests against "
    f"pre-specified margins and family-wide multiple-comparison correction. The evidence "
    f"supports a restricted, regime-specific claim: (i) in from-scratch autoregressive language "
    f"modelling under overfitting pressure, the τ family (the constant-rate ablation Tau(α = 0) "
    f"and τ(w)) beats PSO-tuned AdamW weight decay by {_h_gap7:.1f} PPL at 7M and "
    f"{_h_gap66:.1f} PPL at 66M parameters ({abs(_pct('gpt2_wikitext')):.0f}% and "
    f"{abs(_pct('gpt2_large_wikitext')):.0f}% below the unregularized Baseline), by "
    f"{_h_gap18:.1f} PPL at 18M and not at all at 2M; the variable that orders these cases is "
    f"overfitting pressure — capacity relative to data and epochs, read off the Baseline's best "
    f"validation epoch — not model size, and a 124M confirmation at the standard "
    f"GPT-2-small/WikiText-103 operating point, where the corpus binds, closes the regime from "
    f"the data side with parity against tuned weight decay; (ii) elsewhere — pretrained models, "
    f"the vision transformer, the CNNs and the regression tasks — τ-decay is statistically "
    f"indistinguishable from tuned decoupled weight decay, degrading gracefully rather than "
    f"harming performance (family-wide, across {len(_pvals_flat)} pairwise tests: {_sw} "
    f"significant wins vs {_sl} loss{'es' if _sl != 1 else ''} under BH-FDR q = 0.05, the "
    f"losses being named in Section 5.7; {_swB} vs {_slB} under Bonferroni). Second, a "
    f"scope × adaptivity factorial at 66M that attributes the margin: the decay's scope and "
    f"schedule — weight tensors only, biases excluded, at a rate independent of the learning-rate schedule — "
    f"carry {abs(_eff_scope_const):.1f} PPL of it, while magnitude adaptivity adds "
    f"{abs(_eff_adapt_tau):.2f} PPL at that scope; tested paired by seed across scales, the "
    f"pre-registered adaptivity contrast is {_AG['large'][2]:+.2f} PPL at 66M, "
    f"{_AG['medium'][2]:+.2f} at 18M and {_AG['small'][2]:+.2f} at 7M, i.e. small and not "
    f"consistent in sign. Third, because the contribution is positioned inside the robust-decay "
    f"family rather than beside it, we run the comparison that this positioning demands: a "
    f"head-to-head against the nearest neighbours — Huber decay (the mechanism of Guo & Fan, "
    f"2025), pseudo-Huber/Charbonnier decay and log-cosh decay — under an identical decoupled "
    f"post-optimizer protocol in which the saturation profile is the only thing that varies "
    f"(Section 5.10). They land at parity, which is itself informative: what matters is the "
    f"decoupled weight-only decay, not the profile of its saturation."
)

# ============================================================================
# 2. RELATED WORK
# ============================================================================

doc.add_heading("2. Related Work", level=1)

add_text(
    "Weight decay and norm-based penalties have long been central to neural network regularization. "
    "Early formulations focused on L2 norm penalties, corresponding to classical Tikhonov "
    "regularization (Tikhonov, 1963). L1 regularization was later introduced to promote sparse "
    "solutions, particularly in linear models and compressed sensing (Tibshirani, 1996). Elastic "
    "Net provided a compromise between L1- and L2-style shrinkage (Zou & Hastie, 2005)."
)

add_text(
    "More recently, attention has shifted toward non-convex penalties designed to reduce the bias "
    "of L1 while supporting sparse estimation. SCAD (Fan & Li, 2001) and MCP (Zhang, 2010) "
    "introduced piecewise-defined penalties with nearly unbiased regions around large weights. The "
    "log-sum penalty (Candès et al., 2008) further emphasized selective shrinkage by aggressively "
    "suppressing small weights. Despite their advantages in low-dimensional statistics, these "
    "penalties remain challenging to optimize in deep learning and are rarely adopted in practice."
)

add_text(
    "In the transformer era, regularization has taken additional forms. Dropout (Srivastava et al., "
    "2014) and its variants remain standard, while weight decay in Adam optimizers has been shown "
    "to require decoupled formulations (Loshchilov & Hutter, 2019). Structured pruning (Michel "
    "et al., 2019) and knowledge distillation (Hinton et al., 2015) offer complementary compression "
    "strategies but operate at different levels of abstraction."
)

add_text(
    "A natural question is whether \u03c4(w) reduces to one of three existing lines. First, decoupled "
    "weight decay (Loshchilov & Hutter, 2019): \u03c4(w) shares the decoupling \u2014 the decay is applied "
    "after the optimizer step, so it is not rescaled by Adam's per-coordinate step sizes \u2014 but "
    "generalizes the decayed quantity: AdamW's step is the gradient of a quadratic penalty, whereas "
    "\u03c4(w)'s is the gradient of the Fair penalty (Section 3.1), which is quadratic only near the "
    "origin. Setting \u03b1 = 0 recovers constant decoupled decay exactly; both this ablation "
    "(Tau(\u03b1=0)) and a PSO-tuned decoupled weight-decay competitor (WD-tuned) are included in every "
    "benchmark, so the increment attributable to magnitude adaptivity is measured rather than "
    "asserted. Second, the adaptive weight-decay literature makes the decay coefficient a function "
    "of the optimisation state, at two different granularities. Some of it adapts a single global "
    "coefficient over time: Ghiasi et al. (2023) retune \u03bb at each iteration from the ratio of "
    "gradient norms, while layer-wise scaling rules such as LARS/LAMB (You et al., 2017; You et al., "
    "2020) normalise update magnitudes per layer for large-batch stability. Some of it is already "
    "per parameter: AdaDecay (Nakamura & Hong, 2019) assigns each weight its own decay coefficient "
    "at each iteration, by normalising gradient norms within a layer and mapping them through a "
    "sigmoid. \u03c4(w) therefore does not differ from that line in granularity \u2014 a distinction we "
    "explicitly do not claim \u2014 but in the driving signal and in stationarity: AdaDecay's "
    "per-parameter coefficient is a function of the current gradient and changes from step to step, "
    "whereas \u03c4(w)'s is a fixed, static function of the weight's own magnitude. That is what leaves "
    "\u03c4(w) with a closed-form implicit penalty (Section 3.1); a gradient-driven, time-varying "
    "coefficient corresponds to no fixed objective. Third, non-convex sparse penalties (SCAD, MCP, LSP) pursue the "
    "same goal \u2014 shrink small weights, spare large ones \u2014 but do so non-convexly; \u03c4(w) achieves "
    "the same selectivity with a convex implicit penalty by capping the shrinkage force rather "
    "than flattening the penalty, and dominates them empirically in our from-scratch head-to-head "
    "(Section 5). It is tempting to read it as a continuous relaxation of magnitude pruning "
    "(Han et al., 2015) applied during training rather than post hoc, but the analogy should "
    "not be pushed: because the shrinkage is multiplicative and vanishes as w → 0, τ(w) does "
    "not drive weights to exact zeros in finite time, unlike an L1 proximal operator or hard "
    "thresholding. It concentrates the weight distribution near zero, which is a different "
    "and weaker statement. Throughout this paper 'sparsity' accordingly denotes near-zero "
    "density at the |w| < 10⁻³ threshold, not structural or computational sparsity."
)

add_text(
    "A fourth line is closer than any of these, and is the one against which this contribution "
    "should be read: decoupled decays derived from a robust loss. Replacing the quadratic penalty "
    "implicit in AdamW's decay with a function that is quadratic near the origin and sub-quadratic "
    "in the tails yields a shrinkage force that saturates, so that large weights cease to be taxed "
    "in proportion to their size. AdamHuberDecay (Guo & Fan, 2025) instantiates this with the Huber "
    "loss (Huber, 1964) \u2014 parameters decay quadratically while |w| < \u03b4 and linearly beyond \u2014 and "
    "evaluates it in exactly our regime, from-scratch GPT-2 and GPT-3 pre-training, reporting "
    "healthier late-training weight norms and target perplexities in fewer updates. \u03c4-decay belongs "
    "to this family rather than standing outside it: as Section 3.1 shows, it is the decoupled step "
    "on the Fair penalty, which has the same two asymptotic regimes and the same saturating "
    "shrinkage. We state plainly what does and does not separate them. The penalty is not new \u2014 the "
    "Fair function dates to 1974 \u2014 and the quadratic-to-linear design principle is shared with "
    "Huber decay. What differs is how the transition between the two regimes is made. Huber's "
    "shrinkage is piecewise: it equals the \u21132 gradient exactly below the knee and the \u21131 gradient "
    "exactly above it, with a discontinuity in the second derivative at |w| = \u03b4 and a threshold at "
    "which a weight's treatment changes qualitatively. The Fair shrinkage \u03b7\u00b7w/(\u03c4\u2080 + \u03b1|w|) is "
    "analytic everywhere and approaches both regimes only asymptotically, so no weight ever sits at "
    "a kink. The same distinction, with different transition profiles, separates \u03c4-decay from the "
    "other standard smooth robust losses \u2014 pseudo-Huber/Charbonnier (Charbonnier et al., 1994), "
    "whose shrinkage saturates algebraically as w/\u221a(1+(w/\u03b4)\u00b2), and log-cosh, whose \u03b4\u00b7tanh(w/\u03b4) "
    "saturates exponentially. Whether that difference in transition profile matters empirically is "
    "a question about training dynamics, not about the penalties' asymptotics, and it is not "
    "settled by inspecting the formulas. We therefore treat these three as the nearest competitors "
    "and compare against them head-to-head under an identical protocol (Sections 4.2 and 5.10), "
    "rather than resting the case on the distance from SCAD, MCP and LSP, which are functionally "
    "much further away. Related generalisations \u2014 decoupled decays built from generic \u2113p penalties, "
    "whose shrinkage sign(w)|w|^(p\u22121) is monotone in |w| but, unlike the robust family, does not "
    "saturate \u2014 vary the same design axis without producing the bounded-force behaviour at issue here."
)

add_text(
    "The present work therefore does not claim a new mathematical family of regularization. It "
    "contributes a smooth, decoupled member of the robust weight-regularization family — the Fair "
    "variant — together with its post-optimizer implementation, its closed-form implicit penalty, "
    "and the empirical characterisation of where the decoupled, weight-only, schedule-independent "
    "decay it belongs to helps: autoregressive language models trained from scratch under "
    "overfitting pressure, with graceful degradation to tuned decoupled weight decay elsewhere. "
    "Within that regime, we also measure — rather than assume — how much of the gain the "
    "magnitude-adaptive profile adds over its own constant-rate limit."
)

# ============================================================================
# 3. METHOD
# ============================================================================

doc.add_heading("3. Method", level=1)

add_text(
    "The proposed method defines for each weight w a decay time constant:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u03c4(w) = \u03c4\u2080 + \u03b1|w|")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_text(
    "where \u03c4\u2080 > 0 ensures a minimum decay rate and \u03b1 > 0 governs the degree of adaptivity. "
    "The weight update takes the form:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("w \u2190 w \u2212 \u03b7 \u00b7 w / (\u03c4\u2080 + \u03b1|w|)")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_text(
    "where \u03b7 is the decay strength parameter. This update is applied after the standard optimizer "
    "step (gradient-based update), making it compatible with any first-order optimizer including "
    "SGD and Adam."
)

add_bold_paragraph("Two parameters, not three", size=12, space_before=8)

add_text(
    "Written this way the method appears to have three hyperparameters, but it does not: the "
    "triple is over-parameterised and only two combinations are identifiable. Dividing through "
    "by \u03c4\u2080,"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("w \u2190 w \u2212 \u03c1 \u00b7 w / (1 + |w|/\u03b4),   \u03c1 = \u03b7/\u03c4\u2080,   \u03b4 = \u03c4\u2080/\u03b1")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_text(
    "so the rescaling (\u03b7, \u03c4\u2080, \u03b1) \u2192 (c\u03b7, c\u03c4\u2080, c\u03b1) leaves the update exactly invariant for every "
    "c > 0. The consequences are practical, not cosmetic: a three-dimensional search contains a "
    "perfectly flat direction, numerically distinct configurations denote the same method, and "
    "the individual values of \u03b7, \u03c4\u2080 and \u03b1 cannot be interpreted or reported meaningfully. We "
    "therefore adopt the two-parameter form throughout, and tune (\u03c1, \u03b4) rather than the triple. "
    "The same reduction applies to the \u03b1 = 0 ablation, more sharply: there the update is simply "
    "\u03c1\u00b7w, so it has one parameter, not two. Search ranges for (\u03c1, \u03b4) are set to the image of the "
    "previously used three-dimensional box, so the reachable set of methods is unchanged while "
    "the redundant dimension is removed. The two-dimensional (ρ, δ) search receives the same "
    "40-evaluation budget as every other two-parameter competitor, and the one-dimensional "
    "α = 0 ablation the same 12 as every other one-parameter method (Appendix A.1), so τ(w) "
    "holds no tuning advantage over any competitor of equal dimensionality."
)

add_text(
    "The two parameters have direct readings. \u03c1 is the maximum relative decay rate: the fraction "
    "of its own magnitude a weight loses per step in the limit w \u2192 0, where shrinkage is "
    "strongest. \u03b4 is the magnitude scale, or knee: the weight magnitude at which the shrinkage "
    "has fallen to half its small-weight value, separating the quadratic and linear regimes of "
    "the implicit penalty. This is also the parameterisation in which \u03c4-decay is directly "
    "comparable to the other robust decoupled decays, whose knee is conventionally called \u03b4 "
    "as well (Section 2)."
)

add_bold_paragraph("Stability of the decay step", size=12, space_before=8)

add_text(
    "The decay applied on its own maps w to w\u00b7(1 \u2212 \u03c1/(1 + |w|/\u03b4)). The multiplier lies in (0, 1] "
    "\u2014 a monotone contraction towards zero, never a sign change \u2014 precisely when"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("0 < \u03c1 \u2264 1")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_text(
    "Contraction in magnitude alone would permit the weaker \u03c1 < 2, but for 1 < \u03c1 < 2 the decay "
    "overshoots through zero and the weight alternates in sign at every step, which is an "
    "oscillation rather than the relaxation the method is meant to describe. We impose \u03c1 \u2264 1 as "
    "a hard bound of the search space, so no tuned configuration reported in this paper can "
    "violate it; the bound is tight, in the sense that \u03c1 = 1 is stable and any \u03c1 > 1 flips the "
    "sign of sufficiently small weights. Two caveats. This condition governs the decay step "
    "only: it says nothing about the stability of the optimizer as a whole, which depends on "
    "the learning rate, the loss curvature and the adaptive preconditioner, and is not "
    "addressed here. And the relaxation reading of \u03c4(w) is an analogy \u2014 the update has the "
    "form of a discrete relaxation with a magnitude-dependent rate \u2014 not a derivation from "
    "physical principles."
)

add_text(
    "This update corresponds to a discrete relaxation process in which the effective decay rate "
    "\u03b7/\u03c4(w) decreases with weight magnitude. Small weights experience stronger relative "
    "shrinkage than large ones, producing a balanced form of sparsification. The key properties "
    "of this mechanism are:"
)

add_bullet("The update function is continuous and continuously differentiable everywhere (including w = 0, where the two one-sided derivatives coincide at 1/τ₀), "
           "avoiding the discontinuities characteristic of SCAD, MCP and log-sum penalties.",
           bold_prefix="Smoothness: ")
add_bullet("Unlike L1 (constant shrinkage) or L2 (linear shrinkage), the decay rate adapts "
           "nonlinearly to the current weight magnitude.",
           bold_prefix="Magnitude adaptivity: ")
add_bullet("The method introduces two identifiable hyperparameters (\u03c1, \u03b4) and requires "
           "a single line of code after each optimizer step.",
           bold_prefix="Simplicity: ")
add_bullet("Because the decay is applied post-update, it does not interfere with gradient "
           "computations or adaptive learning rate mechanisms.",
           bold_prefix="Optimizer independence: ")

doc.add_heading("3.1 Implicit penalty: τ(w) as a decoupled step on the Fair function", level=2)

add_text(
    "The update admits a closed-form variational interpretation. It is exactly the decoupled "
    "gradient step w ← w − ∇Ω(w) of the penalty"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ω(w) = (η/α²) · [ α|w| − τ₀ · ln(1 + α|w|/τ₀) ]")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_text(
    "since dΩ/dw = η·w/(τ₀ + α|w|). Up to scaling, Ω is the Fair function of robust statistics "
    "(Fair, 1974; Holland & Welsch, 1977) — there a robust loss on residuals, here transplanted "
    "into a penalty on parameters. We stress that this identification is a statement about what "
    "τ-decay is, not a coincidence to be noted in passing: the penalty is fifty years old, and the "
    "contribution lies in the decoupled application, the resulting training dynamics, and the "
    "regime in which they pay off, not in the function. Five properties follow directly:"
)

add_bullet("Ω is convex (Ω″(w) = η·τ₀/(τ₀ + α|w|)² > 0) and continuously differentiable, with "
           "∇Ω(0) = 0. Large-weight preservation is obtained without non-convexity: the shrinkage "
           "force is capped at η/α rather than removed by a flat region, so — unlike SCAD, MCP and "
           "LSP — the penalty contributes no flat region of its own and no additional stationary point of Ω. We claim nothing stronger: L is non-convex, so the stationary points of L + Ω are not determined by Ω's convexity, and in any case the decoupled algorithm does not literally descend L + Ω (see the caveat below).",
           bold_prefix="Convexity and smoothness: ")
add_bullet("for |w| ≪ τ₀/α the penalty is quadratic, Ω ≈ (η/2τ₀)·w², i.e. L2; for |w| ≫ τ₀/α it "
           "is linear, Ω ≈ (η/α)·|w|, i.e. L1. τ(w) is thus L2 where L2 is safe (the bulk of small "
           "weights) and L1 where L2 over-taxes (the large, signal-carrying weights) — the "
           "mirror-image deployment of SCAD/MCP, which are L1-like near zero and flat at infinity.",
           bold_prefix="Two regimes, one knee (τ₀/α): ")
add_bullet("as α → 0, Ω → (η/2τ₀)·w², recovering constant decoupled weight decay at rate η/τ₀ "
           "exactly. The Tau(α=0) ablation used throughout the experiments is therefore a "
           "decoupled-weight-decay control built into the method: any gain of τ(w) over it is "
           "attributable to magnitude adaptivity alone.",
           bold_prefix="Constant-decay limit: ")
# Worked example read from the tuned 18M configuration (canonical (ρ, δ) in the benchmark
# JSON) and from the instrumented final weight distribution of the same runs; nothing here
# is transcribed.
_ex_hp = _best_hp('gpt2_medium_wikitext').get('τ(w)', {})
_ex_rho = float(_ex_hp.get('rho', float('nan')))
_ex_delta = float(_ex_hp.get('delta', float('nan')))
_ex_med, _ex_max = _iv('medium', 'τ(w)', 'med'), _iv('medium', 'τ(w)', 'max')
_ex_rate = lambda w: _ex_rho / (1.0 + w / _ex_delta)
add_bullet(f"the per-step multiplicative decay rate is ρ/(1 + |w|/δ), monotone decreasing in "
           f"|w|, with its knee at |w| = δ. For the PSO-tuned GPT-2 18M instance "
           f"(ρ = {_ex_rho:.2e}, δ = {_ex_delta:.2f}) the knee sits far above the trained bulk "
           f"(final median |w| = {_ex_med:.3f}) and below only the largest weights "
           f"(max |w| = {_ex_max:.2f}): the median weight decays at "
           f"{_ex_rate(_ex_med) / _ex_rho * 100:.0f}% of the small-weight rate ρ and the largest "
           f"at {_ex_rate(_ex_max) / _ex_rho * 100:.0f}%, a "
           f"{_ex_rate(_ex_med) / _ex_rate(_ex_max):.1f}× differential. The tuned profile is thus "
           f"close to a constant decay over almost the whole distribution and graded only in its "
           f"extreme tail — a fact Section 6.2 returns to.",
           bold_prefix="Effective decay rate: ")
add_bullet("the same quadratic-near-zero, linear-in-the-tails design underlies the Huber decay of "
           "Guo & Fan (2025) and, more generally, every decoupled decay built from a robust loss. "
           "Written in the common (λ, δ) form — shrinkage ≈ λ·w near the origin, knee at |w| = δ — "
           "the family members differ only in the transition: λ·w·min(1, δ/|w|) for Huber "
           "(piecewise, C¹ but not C² at the knee), λ·w/√(1+(w/δ)²) for pseudo-Huber, λ·δ·tanh(w/δ) "
           "for log-cosh, and λ·w/(1+|w|/δ) for τ(w), which is the Fair form and recovers the "
           "update above with λ = η/τ₀ and δ = τ₀/α. All four saturate at λ·δ; τ(w) approaches "
           "that ceiling most slowly, so it holds an intermediate, magnitude-graded regime over "
           "the widest range of |w| — the property the α = 0 ablation removes and the head-to-head "
           "of Section 5.10 tests directly.",
           bold_prefix="Position within the robust-decay family: ")

add_text(
    "One caveat is worth stating explicitly: because the decay is applied after the Adam update, "
    "training does not literally follow proximal gradient descent on loss + Ω. This is the same "
    "deliberate decoupling as in AdamW (Loshchilov & Hutter, 2019), adopted for the same reason — "
    "the penalty should not be rescaled by adaptive per-coordinate step sizes."
)

# ============================================================================
# 4. EXPERIMENTAL SETUP
# ============================================================================

doc.add_heading("4. Experimental Setup", level=1)

doc.add_heading("4.1 Benchmarks", level=2)

add_text(
    "The evaluation spans 12 benchmarks organized into four architectural families: "
    "8 single-scale benchmarks, a three-point from-scratch scale sweep, and a "
    "realistic-scale from-scratch confirmation on WikiText-103. "
    "Two classical datasets, MNIST (LeCun et al., 1998) and CIFAR-10 (Krizhevsky, 2009), "
    "anchor the CNN family; the SST-2 sentiment task from GLUE (Socher et al., 2013; "
    "Wang et al., 2018) and the WikiText-2 language-modelling corpus (Merity et al., 2017) "
    "anchor the language-modelling family. Regression benchmarks use synthetic data with the "
    "Friedman (1991) ten-dimensional nonlinear function as the multivariate target."
)

add_text("Regression (feedforward networks). "
         "(i) A one-dimensional sin(x) regression task (1,153 parameters) and "
         "(ii) a ten-dimensional Friedman-#1-style nonlinear function (1,441 parameters). "
         "These serve as controllable baselines for observing regularization behaviour.", bold=False)

add_text("CNN classification. "
         "(iii) MNIST digit classification (118,282 parameters) and "
         "(iv) CIFAR-10 image classification (2,193,226 parameters).", bold=False)

add_text("Vision transformers. "
         "(v) Custom ViT on CIFAR-10 with patch-based attention (Dosovitskiy et al., 2021).", bold=False)

add_text("Custom-trained language models (training from scratch on the target task). "
         "(vi) BERT-tiny (1,709,954 parameters) on SST-2 binary sentiment classification, and "
         "(vii) GPT-2 Small (7,364,608 parameters) on WikiText-2 causal language modelling. "
         "Both follow the standard BERT (Devlin et al., 2019) and GPT-2 (Radford et al., 2019) "
         "architectures at reduced scale.", bold=False)

add_text("Pretrained autoregressive language model on WikiText-2. "
         "(viii) SmolLM2-135M (134,515,008 parameters; Allal et al., 2025).", bold=False)

add_text("From-scratch autoregressive scale sweep on WikiText-2. "
         "(ix) GPT-2 tiny (~2M parameters), (x) GPT-2 medium (~18M) and "
         "(xi) GPT-2 large (66M), trained from scratch with the same recipe, tokenizer and "
         "12-epoch budget, complementing the reference GPT-2 Small configuration (vii). "
         "Together the four scales trace how the effect of each regularizer changes as model "
         "capacity crosses the size of the training corpus (Section 5.6).", bold=False)

add_text("Realistic-scale from-scratch confirmation on WikiText-103. "
         "(xii) GPT-2 small in its standard configuration (124M parameters: 12 layers, "
         "d = 768, 12 heads, GPT-2 BPE vocabulary of 50,257; Radford et al., 2019) trained "
         "from scratch on WikiText-103 (Merity et al., 2017; ~118M training tokens with the "
         "GPT-2 tokenizer, ~100× WikiText-2). This is the standard architecture / tokenizer / "
         "corpus operating point of the GPT-2 literature, and it deliberately inverts the "
         "sweep's capacity-to-data ratio: here the corpus, not model capacity, is the binding "
         "constraint, so it tests the no-harm side of the claim at realistic scale "
         "(Section 5.6).", bold=False)

doc.add_heading("4.2 Methods compared", level=2)

# ERRATUM (REVIEWER, pointwise 3): the previous wording implied WD-tuned and Tau(alpha=0)
# ran on every benchmark. Whether they do is now read from the rosters actually present in
# the CSVs (Table 1 shows a dash wherever a method is absent), and the per-benchmark roster
# is listed in the reproducibility appendix.
_have_wd = sorted(b for b, _m, _, _, _ in BENCHMARKS + SCALE_BENCHMARKS
                  if 'WD-tuned' in set(all_data[b]['method']))
_n_bench_total = len(BENCHMARKS) + len(SCALE_BENCHMARKS)
_wd_presence = (f"all {_n_bench_total} benchmarks" if len(_have_wd) == _n_bench_total
                else f"{len(_have_wd)} of the {_n_bench_total} benchmarks")
add_text(
    f"Ten methods are compared in total, but not every method runs on every benchmark, and "
    f"Table 1 shows a dash wherever one is absent. The six standard penalties (L1, L2, "
    f"ElasticNet, SCAD, MCP, LSP) together with Baseline and τ(w) run on all single-scale "
    f"benchmarks except SmolLM2. The tuned decoupled-weight-decay competitor (WD-tuned) and "
    f"the α = 0 ablation belong to the protocol and are present on "
    f"{_wd_presence}. "
    f"The from-scratch scale sweep, the WikiText-103 confirmation and SmolLM2 use a reduced "
    f"core-six roster (Baseline, L2, ElasticNet, WD-tuned, Tau(α = 0), τ(w)) dictated by "
    f"their training budgets: L1 and the non-convex penalties, dominated on every "
    f"from-scratch transformer benchmark, are dropped there. The ten methods are:")
methods_list = [
    "Baseline: no regularization.",
    "L1: \u03bb\u2211|w\u1d62| added to the loss.",
    "L2: \u03bb\u2211w\u1d62\u00b2 added to the loss.",
    "Elastic Net: \u03bb\u2081\u2211|w\u1d62| + \u03bb\u2082\u2211w\u1d62\u00b2.",
    "SCAD: piecewise non-convex penalty (Fan & Li, 2001).",
    "MCP: minimax concave penalty (Zhang, 2010).",
    "LSP: log-sum penalty (Candès et al., 2008).",
    "WD-tuned: decoupled weight decay (Loshchilov & Hutter, 2019) with the decay "
    "coefficient tuned by PSO \u2014 the fair constant-rate competitor for \u03c4(w).",
    "Tau(\u03b1=0): the \u03c4(w) update with \u03b1 = 0, i.e. constant decoupled decay at rate \u03c1 applied "
    "with τ(w)'s scope and schedule. This is the control of the pre-registered primary "
    "contrast of Section 4.3 (magnitude adaptivity) and, together with WD-tuned, one arm of "
    "the scope × adaptivity factorial. It is an ablation of τ(w) rather than a competing "
    "method from the literature, so it is excluded from the exploratory win/loss tallies and "
    "from the rank denominators (ranks are reported out of the nine competitors) \u2014 excluded "
    "because it is tested separately and more stringently, not because it is unimportant.",
    "\u03c4(w): the proposed weight-dependent decay.",
]
for m in methods_list:
    add_bullet(m)

add_bold_paragraph("Scope × adaptivity factorial", size=12, space_before=8)

add_text(
    "WD-tuned and Tau(α = 0) apply the same constant decay rate but differ in more than one "
    "respect at once: AdamW decays every parameter handed to the optimizer, biases included, "
    "and scales the decay by the current learning rate, so it follows the warmup/decay "
    "schedule; τ-decay's update touches every tensor named weight — the weight matrices, the "
    "embeddings and the normalization gains — leaves the biases alone, and runs at a constant "
    "per-step rate (this is what \u201cweight-only\u201d means throughout the paper). Comparing τ(w) with WD-tuned therefore conflates the shrinkage "
    "profile with the scope and the schedule. To separate them we run the full 2 × 2:"
)

_2x2_tab = doc.add_table(rows=1, cols=3)
_2x2_tab.style = 'Table Grid'
_2x2_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['', 'Constant profile', 'Magnitude-adaptive profile']):
    _2x2_tab.rows[0].cells[_i].text = ""
    _pp = _2x2_tab.rows[0].cells[_i].paragraphs[0]
    _rr = _pp.add_run(_h); _rr.font.size = Pt(9); _rr.font.name = 'Times New Roman'
    _rr.bold = True
    _pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_2x2_tab.rows[0].cells[_i], 'D9D9D9')
add_table_row(_2x2_tab, ['AdamW scope + schedule', 'WD-tuned', 'Tau(AdamW-scope)'])
add_table_row(_2x2_tab, ['τ-decay scope + schedule', 'Tau(α = 0)', 'τ(w)'])

add_text(
    "Tau(AdamW-scope) is the cell that did not previously exist: the Fair shrinkage of τ(w), "
    "applied to every trainable parameter and scaled by the learning-rate schedule, exactly as "
    "AdamW's decoupled decay is. Reading the table down a column isolates scope and schedule; "
    "reading across a row isolates magnitude adaptivity. Section 6.1 reports the decomposition."
)

add_text(
    "A separate roster is used for the robust-decay head-to-head of Section 5.10, which asks the "
    "narrower question of whether the Fair saturation profile matters relative to the other "
    "members of its own family. There, three further competitors join Baseline, WD-tuned, "
    "Tau(\u03b1 = 0) and \u03c4(w):"
)
for m in [
    "Huber-decay: the decoupled Huber shrinkage \u03bb\u00b7w\u00b7min(1, \u03b4/|w|), i.e. the mechanism of "
    "AdamHuberDecay (Guo & Fan, 2025) applied at our post-optimizer call site.",
    "PseudoHuber-decay: \u03bb\u00b7w/\u221a(1 + (w/\u03b4)\u00b2), the Charbonnier smoothing of the above "
    "(Charbonnier et al., 1994).",
    "LogCosh-decay: \u03bb\u00b7\u03b4\u00b7tanh(w/\u03b4).",
]:
    add_bullet(m)
add_text(
    "All four decays \u2014 the three above and \u03c4(w) \u2014 are applied at the same point in the step "
    "(immediately after optimizer.step), to the same parameter scope (every tensor named "
    "weight; biases excluded), and are tuned by PSO over the same two-dimensional (\u03bb, \u03b4) space, with "
    "\u03bb's range set to cover the full effective small-weight rate \u03b7/\u03c4\u2080 that \u03c4(w) itself can reach, "
    "so that no competitor is handicapped by a narrower search. The saturation profile is thus "
    "the only quantity that varies across the comparison."
)

doc.add_heading("4.3 Experimental protocol", level=2)

add_bold_paragraph("Hierarchy of inference", size=12, space_before=6)

add_text(
    "Because the paper's mechanistic claim concerns magnitude adaptivity specifically, and "
    "because only one comparison isolates it, we fix the status of each comparison in advance "
    "rather than pooling them into a single tally:"
)
add_bullet("τ(w) versus Tau(α = 0) — the implementation-matched constant-decay ablation. "
           "This is the only contrast in which scope, schedule, call site and parameter set "
           "are held fixed and adaptivity alone varies. Confirmatory, pre-specified, and "
           "tested paired by seed. Its outcome turns out to be scale-dependent (Section 5.6): "
           "τ(w) wins at 18M and 66M and loses at 7M, each time by a fraction of a perplexity "
           "point. We keep the pre-specification and report it as such; the paper's headline "
           "claim is therefore drawn from the secondary contrast below, and is presented as "
           "such rather than as a test of magnitude adaptivity.",
           bold_prefix="Primary. ")
add_bullet("τ(w) versus WD-tuned (PSO-tuned decoupled AdamW decay), and the scope × "
           "adaptivity factorial that separates the two factors (Section 4.2; results in "
           "Table 6). Confirmatory for the practical question of whether the decoupled-decay "
           "family is worth adopting over standard practice; the factorial attributes the "
           "margin between its two factors.",
           bold_prefix="Secondary. ")
add_bullet("τ(w) versus L1, L2, ElasticNet, SCAD, MCP and LSP, and the head-to-head against "
           "the robust-decay family (Section 5.10). Reported in full, but these are surveys "
           "of a heterogeneous field rather than tests of the mechanism.",
           bold_prefix="Exploratory. ")
add_bold_paragraph("Equivalence margins", size=12, space_before=6)

add_text(
    f"Several claims in this paper are claims of no harm rather than of superiority. A "
    f"non-significant difference does not establish equivalence — at n = 3–5 a test may fail "
    f"to reject purely for want of power — so wherever we assert that τ-decay matches a "
    f"competitor we test it directly, with two one-sided tests (TOST) against a margin fixed "
    f"before the analysis: ±{EQUIV_MARGIN_PPL} PPL for language modelling and "
    f"±{EQUIV_MARGIN_ACC} percentage points for classification. Equivalence is declared only "
    f"when the whole {int((1 - 2 * 0.05) * 100)}% confidence interval of the difference lies "
    f"inside that margin. Where the data do not support equivalence we say the comparison is "
    f"inconclusive, which at these sample sizes is frequently the honest answer."
)

add_bold_paragraph("Estimation and correction", size=12, space_before=6)

add_text(
    "Effect sizes are reported as Hedges' g — Cohen's d carries an upward bias at these "
    "sample sizes — with Cohen's d retained alongside it for comparability with the "
    "literature. Multiple-comparison correction (Benjamini–Hochberg, and Bonferroni as a "
    "conservative check) is applied to the exploratory family, symmetrically to wins and "
    "losses. The primary comparison is a single pre-specified test and is not corrected."
)

add_text(
    "For each method and benchmark, the method-specific hyperparameters are tuned by Particle "
    "Swarm Optimization (PSO; Kennedy & Eberhart, 1995) on the validation set, with a budget "
    "set by the dimensionality of the search space: 12 evaluations for every one-dimensional "
    "search and 40 for every search of two or more dimensions, so that no method — τ-decay "
    "included — receives more tuning effort than a competitor of equal dimensionality "
    "(Appendix A.1 lists every space and budget). The best configuration is then evaluated "
    "across independent seeds: n = 10 at the 66M scale point, from which the primary and "
    "secondary contrasts are drawn; n = 5 for the single-scale benchmarks; n = 3 for the "
    "tiny and medium sweep members, for SmolLM2 and for the 124M WikiText-103 confirmation, "
    "whose training budgets are the largest. We report mean, standard deviation and 95% "
    "confidence intervals of the primary metric (test MSE for regression, test accuracy for "
    "classification, test perplexity for language modelling). All transformer models use the "
    "AdamW optimizer (Loshchilov & Hutter, 2019), itself built on Adam (Kingma & Ba, 2015). "
    "Sparsity is measured as the fraction of weights with |w| < 10\u207b\u00b3."
)

add_bold_paragraph("Early stopping and the reported model", size=12, space_before=6)

add_text(
    "Every run uses patience-based early stopping on the validation metric and reports the "
    "model of the best validation epoch, restored from a checkpoint taken at that epoch. "
    "Because every reported number depends on this step, the restore is verified at run "
    "time by two independent checks: the restored parameters "
    "must reproduce a fingerprint taken of the checkpoint, and the validation metric "
    "re-evaluated on the restored model must reproduce the value that defined the best "
    "epoch. A run failing either check raises an error instead of reporting a number. The "
    "same best-epoch validation value is the objective of the PSO search, so tuning and "
    "evaluation are consistent. The 124M WikiText-103 runs never stopped before their final "
    "epoch, and their per-epoch validation trajectories, released with the code, show that "
    "the reported model is the best-epoch model."
)

add_text(
    "The single exception to full PSO tuning is the 124M WikiText-103 confirmation, where "
    "a full per-method swarm search is computationally infeasible (~7 GPU-hours per "
    "training run). There, hyperparameters are transferred from the PSO optimum of the "
    "66M sweep member: per-step decay strengths are rescaled by the ratio of total "
    "optimizer steps between the two settings (loss-side penalty coefficients are "
    "transferred unchanged), and each transferred configuration is then validated by a "
    "¼× / 1× / 4× confirmation sweep over its dominant strength parameter on a 2-epoch "
    "budget, keeping the best of the three on validation perplexity. The selected "
    "configurations are then run to the full budget (8 epochs, early stopping with "
    "patience 2, n = 3 seeds)."
)

add_text(
    "Statistical comparisons are performed using Welch's t-test (unequal variance) with "
    "the per-benchmark n, and effect sizes are quantified with Cohen's d."
)

# ============================================================================
# 5. RESULTS
# ============================================================================

doc.add_heading("5. Results", level=1)

doc.add_heading("5.1 Overview", level=2)

add_text(
    "Table 1 summarizes the primary metric for all methods across the 8 single-scale "
    "benchmarks; the three from-scratch scale-sweep benchmarks and the 124M WikiText-103 "
    "confirmation are presented "
    "separately in Section 5.6 (Table 2). Figure 1 provides a visual comparison. A dash "
    "indicates that a method is not part of that benchmark's roster."
)

add_figure('fig1_all_benchmarks.png',
    'Figure 1. Performance comparison across the 8 single-scale benchmarks. '
    'Bar height = mean over n = 5 runs (n = 3 for SmolLM2); error bars = \u00b11 std. '
    'Red border = best method per benchmark. Hatched bars = \u03c4(w). '
    f'Note: LSP bar is clipped in the GPT-2 panel (actual value: '
    f'{get_val(all_data["gpt2_wikitext"], "LSP", "test_ppl", "mean"):.1f}).',
    width_inches=6.5)

# ── TABLE 1: Main results ──
add_bold_paragraph(
    "Table 1. Main results across the 8 single-scale benchmarks. For regression, the metric is test MSE "
    "(\u00d710\u207b\u00b3, lower is better). For classification, the metric is test accuracy "
    "(%, higher is better). For GPT-2 and the WikiText-2 language models, the metric is test perplexity (lower is better). "
    "Bold = best. n = 5 runs (n = 3 for SmolLM2).",
    size=9, space_before=12
)

# Create main results table
headers = ["Method", "sin(x)\nMSE\u00d710\u207b\u00b3", "Complex\nMSE\u00d710\u207b\u00b9",
           "MNIST\n%", "CIFAR\n%", "ViT\n%",
           "BERT\n%", "GPT-2\nPPL", "SmolLM2\nPPL"]
assert len(headers) == len(BENCHMARKS) + 1, "Table 1 header count must match BENCHMARKS"

table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "D9E2F3")

# Build data rows dynamically from CSVs
def fmt_val(val, bname):
    """Format a value for the table based on benchmark type."""
    if bname in ('sin_regression', 'complex_regression'):
        return f"{val:.3f}" if val < 10 else f"{val:.2f}"
    return f"{val:.2f}"

data = []
for method in TABLE_METHODS:
    row = [method]
    for bname, metric, mode, scale, _ in BENCHMARKS:
        df = all_data[bname]
        val = get_val(df, method, metric, 'mean') * scale
        row.append('—' if np.isnan(val) else fmt_val(val, bname))
    data.append(row)

# Find best value per column (dynamically; '—' = method absent from that CSV)
best_indices = {}
for col_idx, (bname, metric, mode, scale, _) in enumerate(BENCHMARKS, 1):
    vals = [float(row[col_idx]) if row[col_idx] != '—' else np.nan for row in data]
    if mode == 'max':
        best_indices[col_idx] = int(np.nanargmax(vals))
    else:
        best_indices[col_idx] = int(np.nanargmin(vals))

for row_idx, row_data in enumerate(data):
    row = table.add_row()
    for col_idx, val in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bold best values
        if col_idx in best_indices and best_indices[col_idx] == row_idx:
            run.bold = True
        # Highlight tau row
        if row_data[0] == '\u03c4(w)':
            set_cell_shading(cell, "FFF2CC")

# Add rank row (dynamically computed)
rank_row = table.add_row()
rank_data = ["\u03c4(w) rank"]
for bname, metric, mode, scale, _ in BENCHMARKS:
    s = all_stats[bname]
    rank_data.append(f"{s['rank']}/{s['total']}")
for i, val in enumerate(rank_data):
    cell = rank_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "E2EFDA")

doc.add_paragraph()  # spacing

# ── 5.2 Regression ──
doc.add_heading("5.2 Regression benchmarks", level=2)

# Sin regression - dynamic
_sin = all_stats['sin_regression']
_sin_df = all_data['sin_regression']
_sin_tau = get_val(_sin_df, 'τ(w)', 'test_mse', 'mean') * 1e3
_sin_best_name = [m for m in METHODS_ORDER if m != 'τ(w)' and m in _sin_df['method'].values]
_sin_sorted = sorted(_sin_best_name, key=lambda m: get_val(_sin_df, m, 'test_mse', 'mean'))
_sin_best = _sin_sorted[0]
_sin_best_val = get_val(_sin_df, _sin_best, 'test_mse', 'mean') * 1e3
_sin_p = _sin['details'][_sin_best]['p']
_sin_d = _sin['details'][_sin_best]['d']
_sin_minp = min(d['p'] for d in _sin['details'].values())
_sin_nsig = sum(1 for d in _sin['details'].values() if d['p'] < 0.05)
_sin_same_as_base = abs(get_val(_sin_df, 'τ(w)', 'test_mse', 'mean')
                        - get_val(_sin_df, 'Baseline', 'test_mse', 'mean')) < 1e-9

add_text(
    f"On the sin(x) regression task, τ-decay achieves a test MSE of {_sin_tau:.2f} × 10⁻³, "
    f"placing {_sin['rank']}/{_sin['total']} among the ranked methods (the α = 0 ablation is "
    f"excluded from rankings) in a field that is statistically indistinguishable. {_sin_best} "
    f"achieves the lowest MSE ({_sin_best_val:.2f} × 10⁻³), but the difference from τ-decay is "
    f"not significant (p = {_sin_p:.3f}, Cohen's d = {_sin_d:+.2f}), and "
    + (f"none of the pairwise comparisons between τ-decay and the other methods reaches "
       if _sin_nsig == 0 else
       f"only {_sin_nsig} of the pairwise comparisons between τ-decay and the other methods reach ")
    + f"significance at α = 0.05 (smallest p = {_sin_minp:.2f})."
    + (" The tuned τ(w) coincides with the unregularized Baseline to four significant digits: "
       "the PSO search set its decay rate to the bottom of the range, i.e. switched the "
       "method off, which is the correct outcome on a task with nothing to regularize."
       if _sin_same_as_base else "")
)

# Complex regression - dynamic
_cx = all_stats['complex_regression']
_cx_df = all_data['complex_regression']
_cx_tau = get_val(_cx_df, 'τ(w)', 'test_mse', 'mean')
_cx_vals = sorted(
    get_val(_cx_df, m, 'test_mse', 'mean')
    for m in METHODS_ORDER if m in _cx_df['method'].values
)
_cx_std = get_val(_cx_df, 'Baseline', 'test_mse', 'std')
_cx_maxd = max(abs(d['d']) for d in _cx['details'].values())
add_text(
    f"The complex regression benchmark yields a similar picture. τ-decay ranks "
    f"{_cx['rank']}/{_cx['total']} ({_cx_tau:.3f}) in a tightly clustered field where all counted "
    f"methods achieve MSEs between {_cx_vals[0]:.3f} and {_cx_vals[-1]:.3f}, a between-method "
    f"spread dwarfed by the run-to-run variability (std ≈ {_cx_std:.2f}). All effect sizes are "
    f"negligible (|d| ≤ {_cx_maxd:.2f}), confirming that the methods are functionally equivalent "
    f"on this task."
)

# ── 5.3 CNN ──
doc.add_heading("5.3 CNN classification", level=2)

# ERRATUM (REVIEWER, pointwise 1): this paragraph carried "97.83%, ranks fourth", which
# contradicted Table 1 and belonged to an earlier 9-method roster. Every figure here is
# now read from the CSV, so the two can no longer drift apart.
_mn = all_stats['mnist']
_mn_df = all_data['mnist']
_mn_tau = get_val(_mn_df, 'τ(w)', 'test_acc', 'mean') * 100
_mn_all = [get_val(_mn_df, _m, 'test_acc', 'mean') * 100 for _m in TABLE_METHODS]
_mn_all = [v for v in _mn_all if v == v]
_mn_sig_ahead = [m for m, d in _mn['details'].items() if d['p'] < 0.05 and not d['better']]
_mn_sig_behind = [m for m, d in _mn['details'].items() if d['p'] < 0.05 and d['better']]
_mn_sig_txt = (
    "with no pairwise comparison involving τ-decay reaching α = 0.05"
    if not (_mn_sig_ahead or _mn_sig_behind) else
    "; the only comparison" + ("s" if len(_mn_sig_ahead) + len(_mn_sig_behind) > 1 else "")
    + " involving τ-decay that reach" + ("" if len(_mn_sig_ahead) + len(_mn_sig_behind) > 1 else "es")
    + " α = 0.05 " + ("are" if len(_mn_sig_ahead) + len(_mn_sig_behind) > 1 else "is")
    + (" against " + _and_join(f"{m} ({get_val(_mn_df, m, 'test_acc', 'mean') * 100:.2f}%, "
                               f"p = {_mn['details'][m]['p']:.3f})"
                               for m in _mn_sig_ahead) + ", which "
       + ("is" if len(_mn_sig_ahead) == 1 else "are") + " ahead" if _mn_sig_ahead else "")
    + (" and" if _mn_sig_ahead and _mn_sig_behind else "")
    + (" against " + _and_join(f"{m} ({get_val(_mn_df, m, 'test_acc', 'mean') * 100:.2f}%, "
                               f"p = {_mn['details'][m]['p']:.3f})"
                               for m in _mn_sig_behind) + ", which "
       + ("is" if len(_mn_sig_behind) == 1 else "are") + " behind" if _mn_sig_behind else "")
)
add_text(
    f"On MNIST, all methods achieve accuracies within a narrow band of "
    f"{min(_mn_all):.2f}–{max(_mn_all):.2f}%. τ-decay ({_mn_tau:.2f}%) ranks "
    f"{_mn['rank']}/{_mn['total']}, the whole field being inside "
    f"{max(_mn_all) - min(_mn_all):.2f} percentage points{_mn_sig_txt}."
    + (" That loss survives family-wide correction and is counted as such in Section 5.7."
       if any(_mn['details'][m].get('sig_bh') for m in _mn_sig_ahead) else "")
)

# CIFAR - dynamic
_cif = all_stats['cifar']
_cif_df = all_data['cifar']
_cif_best_name = sorted(
    [m for m in METHODS_ORDER if m != 'τ(w)' and m in _cif_df['method'].values],
    key=lambda m: -get_val(_cif_df, m, 'test_acc', 'mean')
)[0]
_cif_best_val = get_val(_cif_df, _cif_best_name, 'test_acc', 'mean') * 100
_cif_tau_val = get_val(_cif_df, 'τ(w)', 'test_acc', 'mean') * 100
_cif_p = _cif['details'][_cif_best_name]['p']
_cif_d = _cif['details'][_cif_best_name]['d']
_cif_sig_ahead = [m for m, d in _cif['details'].items() if d['p'] < 0.05 and not d['better']]
_cif_sig_behind = [m for m, d in _cif['details'].items() if d['p'] < 0.05 and d['better']]
_cif_val = lambda m: get_val(_cif_df, m, 'test_acc', 'mean') * 100

add_text(
    f"CIFAR-10 provides more separation. {_cif_best_name} leads at {_cif_best_val:.2f}%, while "
    f"τ-decay ranks {_cif['rank']}/{_cif['total']} at {_cif_tau_val:.2f}%, "
    f"{abs(_cif_tau_val - _cif_best_val):.2f} points behind — a gap that "
    f"{'is' if _cif_p < 0.05 else 'is not'} statistically significant (p = {_cif_p:.4f}, "
    f"d = {_cif_d:+.2f})"
    + (f" and {'survives' if _cif['details'][_cif_best_name].get('sig_bh') else 'does not survive'} "
       f"family-wide correction (Section 5.7)" if _cif_p < 0.05 else "")
    + f". τ-decay is indistinguishable from the unregularized Baseline "
    f"({_cif_val('Baseline'):.2f}%, p = {_cif['details']['Baseline']['p']:.2f}) and from tuned "
    f"decoupled weight decay ({_cif_val('WD-tuned'):.2f}%, p = {_cif['details']['WD-tuned']['p']:.2f})"
    + (", and significantly ahead of " + _and_join(f"{m} ({_cif_val(m):.2f}%, p = {_cif['details'][m]['p']:.3f})"
                                                   for m in _cif_sig_behind) if _cif_sig_behind else "")
    + ". A convolutional network trained from initialisation on 50,000 images is not the "
    f"regime this paper is about, and the CNN result is reported as what it is: a "
    f"{'loss to ' + _and_join(_cif_sig_ahead) if _cif_sig_ahead else 'tie'} on a classic benchmark."
)

# ── 5.4 ViT ──
doc.add_heading("5.4 Vision transformer", level=2)

_vit = all_stats['vit_cifar']
_vit_df = all_data['vit_cifar']
_vit_tau = get_val(_vit_df, 'τ(w)', 'test_acc', 'mean') * 100
# Find top 2 competitors
_vit_others = sorted(
    [(m, get_val(_vit_df, m, 'test_acc', 'mean') * 100) for m in METHODS_ORDER if m != 'τ(w)' and m in _vit_df['method'].values],
    key=lambda x: -x[1]
)
_vit_lsp = get_val(_vit_df, 'LSP', 'test_acc', 'mean') * 100
_vit_lsp_std = get_val(_vit_df, 'LSP', 'test_acc', 'std') * 100
_vit_n_comp = _vit['wins'] + _vit['losses']
_vit_val = lambda m: get_val(_vit_df, m, 'test_acc', 'mean') * 100
_vit_sig_down = [m for m, d in _vit['details'].items() if d['p'] < 0.05 and d['better']]
_vit_sig_up = [m for m, d in _vit['details'].items() if d['p'] < 0.05 and not d['better']]
_vit_n_sig = len(_vit_sig_down) + len(_vit_sig_up)
add_text(
    f"On ViT-CIFAR, τ-decay ranks {_vit['rank']}/{_vit['total']} at {_vit_tau:.2f}%, "
    f"behind {_vit_others[0][0]} ({_vit_others[0][1]:.2f}%) and {_vit_others[1][0]} "
    f"({_vit_others[1][1]:.2f}%), in a field whose leading methods are statistically "
    f"inseparable at n = {N_SEEDS.get('vit_cifar', 5)}. The method outperforms "
    f"{_vit['wins']} of {_vit_n_comp} counted competitors"
    + (f"; the comparison{'s' if _vit_n_sig != 1 else ''} involving τ-decay that "
       f"reach{'' if _vit_n_sig != 1 else 'es'} α = 0.05 "
       f"{'are' if _vit_n_sig != 1 else 'is'}"
       + (" against " + _and_join(f"{m} ({_vit_val(m):.2f}%, p = {_vit['details'][m]['p']:.3f})"
                                  for m in _vit_sig_down)
          + f", which trail{'' if len(_vit_sig_down) != 1 else 's'} it" if _vit_sig_down else "")
       + (" and" if _vit_sig_down and _vit_sig_up else "")
       + (" against " + _and_join(f"{m} ({_vit_val(m):.2f}%, p = {_vit['details'][m]['p']:.3f})"
                                  for m in _vit_sig_up)
          + f", which lead{'' if len(_vit_sig_up) != 1 else 's'} it" if _vit_sig_up else "")
       if _vit_n_sig else
       ", with no pairwise comparison involving τ-decay reaching α = 0.05")
    + ". A vision transformer is a transformer, and τ-decay does not help here: the point "
    "matters for Section 5.8, where the regime rather than the architecture is argued to be "
    "the operative variable."
)

# ── 5.5 Language models ──
doc.add_heading("5.5 Language models", level=2)

add_text(
    "The language-model benchmarks separate the regime in which τ-decay pays from the ones in "
    "which it does not."
)

# BERT - dynamic
_bert = all_stats['bert_sst2']
_bert_df = all_data['bert_sst2']
_bert_tau = get_val(_bert_df, 'τ(w)', 'test_acc', 'mean') * 100
_bert_base = get_val(_bert_df, 'Baseline', 'test_acc', 'mean') * 100
_bert_base_p = _bert['details']['Baseline']['p']
_bert_lsp = get_val(_bert_df, 'LSP', 'test_acc', 'mean') * 100
_bert_rank = sorted([(m, get_val(_bert_df, m, 'test_acc', 'mean') * 100) for m in METHODS_ORDER
                     if m != 'τ(w)' and m in _bert_df['method'].values], key=lambda x: -x[1])
_bert_top, _bert_top_val = _bert_rank[0]
_bert_top_p = _bert['details'][_bert_top]['p']
_bert_top_d = _bert['details'][_bert_top]['d']
_bert_sig = [m for m, d in _bert['details'].items() if d['p'] < 0.05]

p = doc.add_paragraph()
run = p.add_run("BERT-tiny on SST-2 (trained from scratch). ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"τ-decay reaches {_bert_tau:.2f}%, ranking {_bert['rank']}/{_bert['total']} in a field whose "
    f"leading methods are statistically inseparable: the deficit to the leader, {_bert_top} "
    f"({_bert_top_val:.2f}%), is not significant (p = {_bert_top_p:.2f}, d = {_bert_top_d:+.2f}), "
    f"nor is the difference from the unregularized Baseline ({_bert_base:.2f}%, "
    f"p = {_bert_base_p:.2f}), and no pairwise comparison among the leading methods reaches "
    f"α = 0.05 at n = {N_SEEDS.get('bert_sst2', 5)}. "
    + (f"The only significant difference{'s' if len(_bert_sig) != 1 else ''} involving τ-decay "
       f"{'are' if len(_bert_sig) != 1 else 'is'} against {_and_join(_bert_sig)}"
       + (f", which collapses to chance on this binary task ({_bert_lsp:.2f}%)." if _bert_sig == ['LSP'] else ".")
       if _bert_sig else "No comparison involving τ-decay reaches α = 0.05.")
)
run.font.name = 'Times New Roman'

# GPT-2 (7.4M) - dynamic
_gpt = all_stats['gpt2_wikitext']
_gpt_df = all_data['gpt2_wikitext']
_gpt_tau = _dv7('τ(w)')
_gpt_a0 = _dv7('Tau(alpha=0)')
# best counted competitor (lowest PPL after tau; the ablation is not counted)
_gpt_others = sorted(
    [(m, _dv7(m)) for m in METHODS_ORDER if m != 'τ(w)' and m in _gpt_df['method'].values],
    key=lambda x: x[1]
)
_gpt_next = _gpt_others[0]
_gpt_min_d = min(abs(d['d']) for d in _gpt['details'].values())
_gpt_max_p = max(d['p'] for d in _gpt['details'].values())
_gpt_t7, _gpt_p7, _gpt_md7, _gpt_n7 = _AG['small']
_gpt_pen_pct = lambda m: (_dv7(m) - _dv7('Baseline')) / _dv7('Baseline') * 100

p = doc.add_paragraph()
run = p.add_run(f"GPT-2 on WikiText-2 ({_MP['small']:.1f}M, trained from scratch, 30-epoch budget). ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"This is the benchmark on which the regime first shows itself. τ-decay reaches a test "
    f"perplexity of {_gpt_tau:.2f}, against {_gpt_next[1]:.2f} for the best counted competitor "
    f"({_gpt_next[0]}) and {_dv7('Baseline'):.2f} for the unregularized Baseline — "
    f"{_gap('gpt2_wikitext', 'WD-tuned'):.1f} PPL below tuned AdamW weight decay and "
    f"{abs(_pct('gpt2_wikitext')):.1f}% below the Baseline, with "
    f"{'p < 0.001' if _gpt_max_p < 1e-3 else f'p ≤ {_gpt_max_p:.3f}'} and |d| ≥ {_gpt_min_d:.0f} "
    f"against every counted competitor. Two features of the field matter for the interpretation. "
    f"First, the constant-rate ablation Tau(α = 0) finishes at {_gpt_a0:.2f}, i.e. "
    f"{abs(_gpt_md7):.2f} PPL {'below' if _gpt_md7 > 0 else 'above'} τ(w) (paired by seed, "
    f"t = {_gpt_t7:+.1f}, {_pfmt(_gpt_p7)}, n = {_gpt_n7}): at this scale the entire margin belongs "
    f"to the decoupled, weight-only, schedule-independent decay, and magnitude adaptivity "
    f"{'costs' if _gpt_md7 > 0 else 'adds'} a fraction of a point. Second, the loss-side L1 and "
    f"log-sum penalties finish {_gpt_pen_pct('L1'):+.0f}% and {_gpt_pen_pct('LSP'):+.0f}% above the "
    f"Baseline; for L1 this is an artefact of the 12-evaluation one-dimensional search on a "
    f"log-scale λ, which settled on a coefficient too large for a 30-epoch run (Section 6.4), not "
    f"a property of the penalty. The Baseline's best validation epoch falls at "
    f"{_best_epoch('small', 'Baseline'):.0f} of {_EPOCH_BUDGET['small']} and WD-tuned's at "
    f"{_best_epoch('small', 'WD-tuned'):.0f}, whereas both τ-family runs are still improving at "
    f"epoch {_EPOCH_BUDGET['small']}: this is the overfitting-delay signature that Section 5.6 "
    f"traces across scales."
)
run.font.name = 'Times New Roman'

add_text("Against the counted competitors, in order of their perplexity:")
for comp_m, comp_v in _gpt_others:
    det = _gpt['details'][comp_m]
    pct = (_gpt_tau - comp_v) / comp_v * 100
    add_bullet(f"vs. {comp_m} ({comp_v:.2f}): {pct:+.1f}% perplexity ({_pfmt(det['p'])}, "
               f"d = {det['d']:+.1f})")

add_figure('fig2_gpt2_detail.png',
    f'Figure 2A. GPT-2 ({_MP["small"]:.1f}M) WikiText-2 language modelling results. '
    f'(A) Performance comparison: τ(w) reaches perplexity {_gpt_tau:.1f}, against '
    f'{_gpt_next[1]:.1f} for the best counted competitor ({_gpt_next[0]}); the constant-rate '
    f'ablation Tau(α=0), at {_gpt_a0:.1f}, is the nominal best of the roster. '
    '(B) 95% confidence intervals; stars mark Welch p-values against τ(w) '
    '(*** p < 0.001, ** p < 0.01, * p < 0.05).',
    width_inches=6.5)

# SmolLM2 - dynamic
_sm = all_stats['smollm2_wikitext']
_sm_df = all_data['smollm2_wikitext']
_sm_tau = get_val(_sm_df, 'τ(w)', 'test_ppl', 'mean')
_sm_base = get_val(_sm_df, 'Baseline', 'test_ppl', 'mean')
_sm_wd = get_val(_sm_df, 'WD-tuned', 'test_ppl', 'mean')
_sm_ab = get_val(_sm_df, 'Tau(alpha=0)', 'test_ppl', 'mean')
_sm_l2 = get_val(_sm_df, 'L2', 'test_ppl', 'mean')
_sm_en = get_val(_sm_df, 'ElasticNet', 'test_ppl', 'mean')
_sm_fam = [_sm_base, _sm_wd, _sm_ab, _sm_tau]
_sm_all = [get_val(_sm_df, m, 'test_ppl', 'mean') for m in TABLE_METHODS if m in _sm_df['method'].values]
_sm_maxstd = max(get_val(_sm_df, m, 'test_ppl', 'std') for m in TABLE_METHODS if m in _sm_df['method'].values)
_sm_sig = [m for m, d in _sm['details'].items() if d['p'] < 0.05]
_sm_eq_wd = equiv_verdict('smollm2_wikitext', 'τ(w)', 'WD-tuned')[0]
_sm_eq_base = equiv_verdict('smollm2_wikitext', 'τ(w)', 'Baseline')[0]
_sm_rho = float(_best_hp('smollm2_wikitext').get('τ(w)', {}).get('rho', float('nan')))
_med_rho = float(_best_hp('gpt2_medium_wikitext').get('τ(w)', {}).get('rho', float('nan')))

p = doc.add_paragraph()
run = p.add_run("SmolLM2-135M on WikiText-2. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"This modern small LLM (run with the reduced core-six roster, Section 4.2) is a null-effect "
    f"benchmark: every method finishes within {max(_sm_all) - min(_sm_all):.2f} PPL of every other. "
    f"The four decay-family configurations — Baseline ({_sm_base:.2f}), WD-tuned ({_sm_wd:.2f}), "
    f"Tau(α=0) ({_sm_ab:.2f}) and τ(w) ({_sm_tau:.2f}) — span {max(_sm_fam) - min(_sm_fam):.2f} PPL, "
    f"and the loss-side penalties are marginally worse (L2 {_sm_l2:.2f}, ElasticNet {_sm_en:.2f}). "
    f"Seed-to-seed variability is minuscule (std ≤ {_sm_maxstd:.2f} PPL), so "
    + (f"the comparison{'s' if len(_sm_sig) != 1 else ''} against {_and_join(_sm_sig)} "
       f"reach{'' if len(_sm_sig) != 1 else 'es'} nominal significance" if _sm_sig else
       "no comparison reaches nominal significance")
    + f"; against tuned decoupled weight decay τ(w) is {_sm_eq_wd}, and against the unregularized "
    f"Baseline it is {_sm_eq_base}. The PSO search drove τ(w)'s decay rate to "
    f"ρ = {_sm_rho:.1e}, {_med_rho / _sm_rho:.0f}× below the from-scratch 18M optimum: on a model "
    f"already saturated by pretraining the method has effectively switched itself off, which is a "
    f"negative result with a positive implication — τ-decay neither helps nor harms on the "
    f"3-epoch fine-tuning budget, degrading gracefully toward the Baseline."
)
run.font.name = 'Times New Roman'

add_figure('fig4_transformer_ci.png',
    'Figure 3. 95% confidence intervals for the four transformer benchmarks of Table 1 '
    '(BERT-tiny on SST-2, GPT-2 on WikiText-2, SmolLM2-135M on WikiText-2 and ViT on CIFAR-10). '
    'Diamond markers = \u03c4(w); dashed gold line = \u03c4(w) mean.',
    width_inches=6.5)

# ── 5.6 Statistical summary ──
doc.add_heading("5.6 From-scratch scale sweep (2M → 66M) and 124M confirmation on WikiText-103", level=2)

_n_sw = {s: N_SEEDS[_SCALE_KEYS[s]] for s in ('tiny', 'small', 'medium', 'large')}
add_text(
    f"The from-scratch scale sweep holds the task (WikiText-2 causal language modelling), "
    f"tokenizer and optimizer fixed and varies model capacity: GPT-2 tiny "
    f"({_MP['tiny']:.1f}M parameters), medium ({_MP['medium']:.0f}M) and large "
    f"({_MP['large']:.0f}M), each on a 12-epoch budget with the core six-method roster, "
    f"dimension-aware PSO tuning and n = {_n_sw['tiny']}, {_n_sw['medium']} and "
    f"{_n_sw['large']} seeds respectively; at 66M the roster also carries the fourth cell of "
    f"the scope × adaptivity factorial, Tau(AdamW-scope). The reference GPT-2 Small "
    f"({_MP['small']:.1f}M, 30-epoch budget, full ten-method roster, n = {_n_sw['small']}) from "
    f"Table 1 is reported alongside. Table 2 gives the test perplexities; Figure 2B plots the sweep."
)

add_figure('fig2b_scale_sweep.png',
           f'Figure 2B. From-scratch GPT-2 scale sweep on WikiText-2 (core roster; '
           f'n = {_n_sw["tiny"]}, {_n_sw["small"]}, {_n_sw["medium"]} and {_n_sw["large"]} seeds '
           f'at 2M, 7M, 18M and 66M). Left: test perplexity per method vs model size; open '
           f'markers mark the GPT-2 Small reference, whose longer 30-epoch budget makes it '
           f'comparable within-column only. Right: τ(w) margin (Δ test PPL) over the '
           f'unregularized Baseline and over tuned decoupled weight decay at each scale. The '
           f'margin does not grow monotonically with size: it appears where the unregularized '
           f'Baseline overfits its budget (7M on 30 epochs and 66M on 12 epochs; not 18M on '
           f'12 epochs, not 2M).')

add_bold_paragraph(
    f"Table 2. From-scratch GPT-2 scale sweep on WikiText-2: test perplexity (mean ± std, "
    f"lower is better). Bold = best per scale. The Small column uses a longer 30-epoch "
    f"budget, so absolute values are not comparable across all columns; within-column "
    f"comparisons are. Tau(α=0) is the ablation of τ(w) and Tau(AdamW-scope) the fourth "
    f"factorial cell (66M only); both are excluded from win/loss statistics. A dash marks a "
    f"method absent from that scale's roster.",
    size=9, space_before=12
)

_SWEEP_COLS = [('gpt2_tiny_wikitext', f'tiny\n{_MP["tiny"]:.1f}M'), ('gpt2_wikitext', f'Small\n{_MP["small"]:.1f}M'),
               ('gpt2_medium_wikitext', f'medium\n{_MP["medium"]:.0f}M'), ('gpt2_large_wikitext', f'large\n{_MP["large"]:.0f}M')]
_SWEEP_ROWS = ['Baseline', 'L2', 'ElasticNet', 'WD-tuned', 'Tau(AdamW-scope)', 'Tau(alpha=0)', 'τ(w)']
# Roster complete (2026-08-31): 66M carries L2 / ElasticNet and 7M carries SCAD; the
#   guarded lookups below still render '—' for any method absent from a CSV.

_sw_table = doc.add_table(rows=1, cols=len(_SWEEP_COLS) + 1)
_sw_table.style = 'Table Grid'
_sw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['Method'] + [c[1] for c in _SWEEP_COLS]):
    _cell = _sw_table.rows[0].cells[_i]
    _cell.text = ""
    _p = _cell.paragraphs[0]
    _run = _p.add_run(_h)
    _run.bold = True
    _run.font.size = Pt(9)
    _run.font.name = 'Times New Roman'
    _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_cell, "D9E2F3")

_sw_best = {}
for _ci, (_bn, _) in enumerate(_SWEEP_COLS, 1):
    _vals = [get_val(all_data[_bn], _m, 'test_ppl', 'mean') for _m in _SWEEP_ROWS]
    _sw_best[_ci] = int(np.nanargmin(_vals)) if any(v == v for v in _vals) else -1

for _ri, _m in enumerate(_SWEEP_ROWS):
    _row = _sw_table.add_row()
    _cell = _row.cells[0]
    _cell.text = ""
    _p = _cell.paragraphs[0]
    _run = _p.add_run(_m)
    _run.font.size = Pt(9)
    _run.font.name = 'Times New Roman'
    for _ci, (_bn, _) in enumerate(_SWEEP_COLS, 1):
        _mean = get_val(all_data[_bn], _m, 'test_ppl', 'mean')
        _std = get_val(all_data[_bn], _m, 'test_ppl', 'std')
        _cell = _row.cells[_ci]
        _cell.text = ""
        _p = _cell.paragraphs[0]
        _run = _p.add_run('—' if np.isnan(_mean) else f"{_mean:.2f} ± {_std:.2f}")
        _run.font.size = Pt(9)
        _run.font.name = 'Times New Roman'
        _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _sw_best.get(_ci) == _ri:
            _run.bold = True

# Per-scale details dicts; every lookup below is guarded so a method missing from a
# provisional CSV cannot crash the build.
_ts = all_stats['gpt2_tiny_wikitext']['details']
_ss = all_stats['gpt2_wikitext']['details']
_ms = all_stats['gpt2_medium_wikitext']['details']
_ls = all_stats['gpt2_large_wikitext']['details']
_pget = lambda details, m: details.get(m, {}).get('p', float('nan'))
_dd = lambda details, m: details.get(m, {}).get('d', float('nan'))
_tiny_rho_tau = float(_best_hp('gpt2_tiny_wikitext').get('τ(w)', {}).get('rho', float('nan')))
_tiny_rho_a0 = float(_best_hp('gpt2_tiny_wikitext').get('Tau(alpha=0)', {}).get('rho', float('nan')))
_tiny_same = abs(_dv2('τ(w)') - _dv2('Baseline')) < 1e-3

add_text(
    f"Read across scales, the sweep does not trace model size; it traces overfitting pressure. "
    f"At {_MP['tiny']:.1f}M parameters no method separates from the Baseline (τ(w) vs Baseline: "
    f"Δ = {_sgn(-_gap('gpt2_tiny_wikitext', 'Baseline'))} PPL, {_pfmt(_pget(_ts, 'Baseline'))})"
    + (f"; the PSO search drove τ(w)'s decay rate to the bottom of its range "
       f"(ρ = {_tiny_rho_tau:.1e}), so τ(w), WD-tuned and the Baseline are the same run to four "
       f"significant digits" if _tiny_same else "")
    + f" — the correct null where there is no overfitting to prevent. At {_MP['medium']:.0f}M "
    f"τ(w) leads tuned decoupled weight decay by {_gap('gpt2_medium_wikitext', 'WD-tuned'):.2f} PPL "
    f"({_pfmt(_pget(_ms, 'WD-tuned'))}) and the Baseline by "
    f"{_gap('gpt2_medium_wikitext', 'Baseline'):.2f} PPL ({abs(_pct('gpt2_medium_wikitext')):.1f}%). "
    f"At {_MP['large']:.0f}M the margin is {_gap('gpt2_large_wikitext', 'WD-tuned'):.2f} PPL over "
    f"WD-tuned and {_gap('gpt2_large_wikitext', 'Baseline'):.2f} PPL "
    f"({abs(_pct('gpt2_large_wikitext')):.1f}%) over the Baseline "
    f"({_pfmt(_pget(_ls, 'Baseline'))}, d = {abs(_dd(_ls, 'Baseline')):.1f}, n = {_n_sw['large']}). "
    f"The {_MP['small']:.1f}M reference, trained for 30 rather than 12 epochs, sits with the 66M "
    f"point rather than between 2M and 18M: {_gap('gpt2_wikitext', 'WD-tuned'):.2f} PPL over "
    f"WD-tuned and {abs(_pct('gpt2_wikitext')):.1f}% over the Baseline. Model size alone cannot "
    f"order these four numbers; the ratio of capacity × epochs to data can. One further "
    f"reading of Table 2 matters for the Discussion: at {_MP['large']:.0f}M the loss-side "
    f"penalties are driven to (near) zero effect by their own tuning — L2 finishes at "
    f"{_dv('L2'):.2f} against {_dv('Baseline'):.2f} for the Baseline, and ElasticNet's λ sits "
    f"at the bottom of its search range — so loss-side penalties buy essentially nothing at "
    f"the scale where the decoupled decays gain {_d_base - _d_wd:.1f}–{_d_base - _d_tau:.1f} "
    f"PPL."
)

add_text(
    f"The training dynamics say the same thing directly. The tracker records the epoch of best "
    f"validation perplexity for every run. At {_MP['medium']:.0f}M and at {_MP['tiny']:.1f}M every "
    f"method's best epoch is the last one (Baseline: {_best_epoch('medium', 'Baseline'):.0f} of "
    f"{_EPOCH_BUDGET['medium']} and {_best_epoch('tiny', 'Baseline'):.0f} of "
    f"{_EPOCH_BUDGET['tiny']}): nothing overfits inside the budget, and the gain is accordingly "
    f"modest or absent. At {_MP['large']:.0f}M the Baseline's best epoch falls at "
    f"{_best_epoch('large', 'Baseline'):.1f} of {_EPOCH_BUDGET['large']} and WD-tuned's at "
    f"{_best_epoch('large', 'WD-tuned'):.1f}, both then deteriorating, while Tau(α = 0), "
    f"Tau(AdamW-scope) and τ(w) are still improving when the budget ends (mean best epoch "
    f"{_best_epoch('large', 'Tau(alpha=0)'):.1f}, {_best_epoch('large', 'Tau(AdamW-scope)'):.1f} and "
    f"{_best_epoch('large', 'τ(w)'):.1f}; Section 6.2). At {_MP['small']:.1f}M on a 30-epoch budget "
    f"the Baseline peaks at epoch {_best_epoch('small', 'Baseline'):.1f} and WD-tuned at "
    f"{_best_epoch('small', 'WD-tuned'):.1f}, while both τ-family runs are still improving at epoch "
    f"{_EPOCH_BUDGET['small']}. The gain of the τ family is therefore an overfitting-delay "
    f"effect: it appears exactly where the unregularized model's validation optimum precedes "
    f"the end of training. Where it appears, not how large it is — Section 5.6.1 shows that "
    f"at fixed hyperparameters the margin grows with the amount of training data even as that "
    f"pressure eases, so the pressure should be read as the switch for the effect rather than "
    f"as its dial."
)

add_text(
    f"The pre-registered primary contrast — τ(w) against its implementation-matched "
    f"constant-decay ablation, paired by seed — is scale-dependent in sign, not merely in size. "
    f"τ(w) − Tau(α = 0) is {_AG['small'][2]:+.2f} PPL at {_MP['small']:.1f}M "
    f"(t = {_AG['small'][0]:+.1f}, {_pfmt(_AG['small'][1])}, n = {_AG['small'][3]}: the constant "
    f"decay wins), {_AG['medium'][2]:+.2f} at {_MP['medium']:.0f}M (t = {_AG['medium'][0]:+.1f}, "
    f"{_pfmt(_AG['medium'][1])}, n = {_AG['medium'][3]}) and {_AG['large'][2]:+.2f} at "
    f"{_MP['large']:.0f}M (t = {_AG['large'][0]:+.1f}, {_pfmt(_AG['large'][1])}, "
    f"n = {_AG['large'][3]}). At {_MP['tiny']:.1f}M the two differ by {_AG['tiny'][2]:+.2f} PPL, "
    f"but that is a difference of search outcome rather than of profile: τ(w)'s search settled at "
    f"no decay, the ablation's at a small non-zero rate (ρ = {_tiny_rho_a0:.1e}) that costs about "
    f"a point, and neither run has a knee inside the weight distribution. Each of the "
    f"over-capacity increments is a fraction of a perplexity point against a family margin of "
    f"{min(_gap('gpt2_wikitext', 'WD-tuned'), _gap('gpt2_large_wikitext', 'WD-tuned')):.1f}–"
    f"{max(_gap('gpt2_wikitext', 'WD-tuned'), _gap('gpt2_large_wikitext', 'WD-tuned')):.1f} PPL over "
    f"tuned weight decay. Magnitude adaptivity is, on this evidence, a second-order and "
    f"scale-dependent refinement of the decoupled decay; the factorial of Section 6.1 shows "
    f"where its apparent value at 66M comes from."
)

add_bold_paragraph("Realistic-scale confirmation: GPT-2 small (124M) on WikiText-103.",
                   size=12, space_before=12)

_w3 = all_stats['gpt2_wt103']['details']
_w3d = all_data['gpt2_wt103']
_w3f = lambda m: (get_val(_w3d, m, 'test_ppl', 'mean'), get_val(_w3d, m, 'test_ppl', 'std'))
# Tau(alpha=0) is excluded from the win/loss family by design; the ablation sentence uses the
# paired contrast computed above.
_w3_t, _w3_p, _w3_md, _w3_n = _AG['wt103']
_w3_l2_bh = _w3.get('L2', {}).get('sig_bh', False)
_w3_l2_bonf = _w3.get('L2', {}).get('sig_bonf', False)
_w3_base_bonf = _w3.get('Baseline', {}).get('sig_bonf', False)
_w3_wd_eq = equiv_verdict('gpt2_wt103', 'τ(w)', 'WD-tuned')[0]
add_text(
    f"The 124M run inverts the sweep's capacity-to-data ratio: 118M training tokens against "
    f"124M parameters on an 8-epoch budget leave every method's validation perplexity still "
    f"falling when training ends (every method's best epoch is "
    f"{_best_epoch('wt103', 'Baseline'):.0f} of {_EPOCH_BUDGET['wt103']}) — there is no "
    f"overfitting to prevent, placing this benchmark outside the τ family's target regime by "
    f"design. The results match that prediction. Test perplexities (mean ± std over "
    f"{N_SEEDS['gpt2_wt103']} seeds) are: Baseline "
    f"{_w3f('Baseline')[0]:.2f} ± {_w3f('Baseline')[1]:.2f}, L2 "
    f"{_w3f('L2')[0]:.2f} ± {_w3f('L2')[1]:.2f}, ElasticNet "
    f"{_w3f('ElasticNet')[0]:.2f} ± {_w3f('ElasticNet')[1]:.2f}, WD-tuned "
    f"{_w3f('WD-tuned')[0]:.2f} ± {_w3f('WD-tuned')[1]:.2f}, Tau(α = 0) "
    f"{_w3f('Tau(alpha=0)')[0]:.2f} ± {_w3f('Tau(alpha=0)')[1]:.2f} and τ(w) "
    f"{_w3f('τ(w)')[0]:.2f} ± {_w3f('τ(w)')[1]:.2f} — absolute values in the standard range of "
    f"the from-scratch GPT-2-small/WikiText-103 literature. τ(w) improves on the unregularized "
    f"baseline by {_gap('gpt2_wt103', 'Baseline'):.2f} PPL ({_pfmt(_pget(_w3, 'Baseline'))}"
    f"{', significant under family-wide Bonferroni' if _w3_base_bonf else ''}), is "
    f"{'practically equivalent to' if 'equivalent' in _w3_wd_eq else 'statistically indistinguishable from'} "
    f"tuned decoupled weight decay (Δ = {_gap('gpt2_wt103', 'WD-tuned'):+.2f} PPL in τ(w)'s "
    f"favour, {_pfmt(_pget(_w3, 'WD-tuned'))}), and trails L2 by "
    f"{-_gap('gpt2_wt103', 'L2'):.2f} PPL — a nominal loss ({_pfmt(_pget(_w3, 'L2'))}) that "
    f"{'survives' if _w3_l2_bh else 'does not survive'} family-wide Benjamini–Hochberg correction"
    f"{' but not Bonferroni' if _w3_l2_bh and not _w3_l2_bonf else ''}. The α = 0 ablation "
    f"trails τ(w) here by {abs(_w3_md):.2f} PPL (paired t = {_w3_t:+.1f}, {_pfmt(_w3_p)}, "
    f"n = {_w3_n}) — the largest adaptivity increment in the study, obtained under "
    f"transfer-tuned rather than PSO-tuned hyperparameters (Section 4.3), so we record it "
    f"without leaning on it. Taken together with the sweep, the two endpoints bracket the "
    f"claim honestly: the advantage over tuned weight decay is a property of the overfitting "
    f"regime, and where that regime ends the τ family degrades to parity — never to harm — at "
    f"the standard architecture/tokenizer/corpus operating point of the GPT-2 literature."
)

add_figure('mechanism_overfitting_delay_wt103.png',
           'Figure 2C. Train vs validation perplexity, GPT-2 small (124M) trained from scratch '
           'on WikiText-103, mean over 3 seeds (dashed = train, solid = validation; right panel: '
           'late-epoch zoom). Unlike the over-capacity sweep members, no method shows an '
           'overfitting upturn within the 8-epoch budget — the regime where the τ-family advantage '
           'flattens to parity with tuned weight decay, by design.')

doc.add_heading("5.6.1 Data-quantity arm at fixed capacity", level=3)

# Data-quantity arm (C1): the 66M model trained on 25% / 50% / 100% of the WikiText-2
#   training set (Baseline, Tau(alpha=0), τ(w); n = 5 seeds; hyperparameters fixed at the
#   values tuned on the full corpus; optimizer steps held constant). Everything below is
#   read from the data25 / data50 CSVs via the guarded accessors defined above; when the
#   files are absent an explicit placeholder is emitted instead.
add_text(
    "The sweep varies capacity at fixed data, which leaves the overfitting-pressure reading "
    "suggestive rather than identified: model size and the capacity-to-data ratio move "
    "together. The complementary arm holds the 66M model fixed and trains it on 25%, 50% and "
    "100% of the WikiText-2 training set (Baseline, the constant-decay ablation and τ(w); "
    "n = 5 seeds; hyperparameters fixed at the values tuned on the full corpus; the number "
    "of optimizer steps held constant, so a smaller corpus is revisited proportionally more "
    "often; the 100% column is the main 66M run). If overfitting pressure acted as a dial, "
    "the τ-family margin over the Baseline should widen as the corpus shrinks."
)
if _HAVE_DATA_ARM:
    _da_tab = doc.add_table(rows=1, cols=4)
    _da_tab.style = 'Table Grid'
    _da_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for _i, _h in enumerate(['Method', '25% of corpus', '50% of corpus', '100% of corpus']):
        _c = _da_tab.rows[0].cells[_i]; _c.text = ""
        _r = _c.paragraphs[0].add_run(_h); _r.bold = True
        _r.font.size = Pt(9); _r.font.name = 'Times New Roman'
        _c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(_c, 'D9D9D9')
    for _m in ['Baseline', 'Tau(alpha=0)', 'τ(w)']:
        _cells = ['Tau(α = 0)' if _m == 'Tau(alpha=0)' else _m]
        for _f in (25, 50, 100):
            _mm = _da_val(_f, _m)
            _sd = _da_val(_f, _m, 'std')
            _cells.append('—' if _mm != _mm else f'{_mm:.2f} ± {_sd:.2f}')
        add_table_row(_da_tab, _cells, bold=(_m == 'τ(w)'))
    add_table_row(_da_tab, ['τ(w) vs Baseline']
                  + [f'{_da_pct(_f):+.1f}%' for _f in (25, 50, 100)],
                  bg_color='F2F2F2')
    _p = doc.add_paragraph()
    _r = _p.add_run('Table 2b. Data-quantity arm: test perplexity of the 66M model trained on '
                    '25%, 50% and 100% of the WikiText-2 training set (mean ± std over n = 5 '
                    'seeds; hyperparameters fixed at the 100% optimum; optimizer steps held '
                    'constant). The 100% column is the main 66M run of Table 2.')
    _r.font.size = Pt(9); _r.font.name = 'Times New Roman'; _r.italic = True
    add_text(
        f"The prediction fails, and the failure is informative. At 25% of the corpus τ(w) "
        f"improves on the Baseline by {abs(_da_pct(25)):.1f}% ({_da_val(25, 'τ(w)'):.2f} vs "
        f"{_da_val(25, 'Baseline'):.2f}); at 50% by {abs(_da_pct(50)):.1f}% "
        f"({_da_val(50, 'τ(w)'):.2f} vs {_da_val(50, 'Baseline'):.2f}); at 100% by "
        f"{abs(_da_pct(100)):.1f}%. At fixed hyperparameters the margin grows with the amount "
        f"of data — even though the reduced corpora put the Baseline under visibly harsher "
        f"pressure: at 25% its best validation epoch falls at "
        f"{_da_best_ep(25, 'Baseline'):.1f} of {_da_budget(25):.0f}, i.e. "
        f"{_da_best_ep(25, 'Baseline') / _da_budget(25) * 100:.0f}% of the budget, against "
        f"{_da_best_ep(50, 'Baseline'):.1f} of {_da_budget(50):.0f} at 50% and "
        f"{_da_best_ep(100, 'Baseline'):.1f} of {_da_budget(100):.0f} at 100%. Reducing the "
        f"data at fixed capacity therefore does not widen the gap, and 'overfitting pressure' "
        f"alone is not a sufficient description of the regime: the decay's advantage needs "
        f"both a baseline that overfits its budget and enough data for the regularized model "
        f"to exploit the additional useful training it affords. (The constant-decay ablation "
        f"is indistinguishable from τ(w) at both reduced fractions: "
        f"{_da_val(25, 'Tau(alpha=0)'):.2f} vs {_da_val(25, 'τ(w)'):.2f} at 25% and "
        f"{_da_val(50, 'Tau(alpha=0)'):.2f} vs {_da_val(50, 'τ(w)'):.2f} at 50%.) One caveat "
        f"keeps the result open: because the hyperparameters were transferred rather than "
        f"re-tuned per fraction, an off-optimum decay strength at higher pressure "
        f"('under-regularized at higher pressure') cannot be separated here from a genuine "
        f"data dependence; a per-fraction re-tuning would separate the two, and is left open."
    )
else:
    _pl = doc.add_paragraph()
    _plr = _pl.add_run(
        '[RESULTS PENDING — NOT FOR SUBMISSION] The data-quantity arm has not completed; no '
        'numbers are reported here. This placeholder is emitted automatically whenever the '
        'data25 / data50 result files are absent, so that the manuscript can never present '
        'this arm as done when it is not.')
    _plr.font.size = Pt(11)
    _plr.font.name = 'Times New Roman'
    _plr.bold = True
    _plr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_heading("5.7 Statistical summary", level=2)

# REVIEWER 7: the "N significant wins vs 1 loss" tally is arithmetically right but reads as
# N independent confirmations, which it is not. Lead with a practical win/tie/loss count
# against the pre-registered margins, and demote the significance tally to what it is.
_pw = _pt = _pl = 0
for _bn, _st in all_stats.items():
    _metric = dict((b[0], b[1]) for b in BENCHMARKS + SCALE_BENCHMARKS)[_bn]
    _margin = EQUIV_MARGIN_PPL if _metric == 'test_ppl' else EQUIV_MARGIN_ACC
    _scale_f = 100.0 if _metric == 'test_acc' else 1.0
    for _m, _d in _st['details'].items():
        _delta = (_d['tau_mean'] - _d['m_mean']) * _scale_f
        if abs(_delta) <= _margin:
            _pt += 1
        elif _d['better']:
            _pw += 1
        else:
            _pl += 1
_ptot = _pw + _pt + _pl

add_text(
    f"We report two different counts, because they answer different questions and the second "
    f"is easily over-read. The practical count asks how often the difference exceeds a margin "
    f"anyone would act on, fixed in advance at ±{EQUIV_MARGIN_PPL} PPL and "
    f"±{EQUIV_MARGIN_ACC} accuracy points (Section 4.3): across all {_ptot} τ(w)-vs-competitor "
    f"comparisons in the study, {_pw} ({_pw/_ptot:.0%}) favour τ-decay by more than the "
    f"margin, {_pl} ({_pl/_ptot:.0%}) favour the competitor by more than the margin, and "
    f"{_pt} ({_pt/_ptot:.0%}) fall inside it and are practical ties. That is a materially "
    f"different picture from the significance tally below, and a more honest one: τ-decay "
    f"is ahead more often than behind, but it is neither uniformly better nor rarely worse."
)

add_text(
    f"The significance count asks a narrower question: how many differences are large relative "
    f"to their seed-to-seed noise. Under family-wide Benjamini–Hochberg correction "
    f"(q = 0.05) applied symmetrically, {_sw} comparisons are significant in τ-decay's favour "
    f"and {_sl} against. That number should not be read as {_sw} independent confirmations. "
    f"It counts comparisons, not findings: one benchmark contributes one comparison per "
    f"competitor, so a single model can add seven or eight to the tally, and several of those "
    f"comparisons share the same τ-decay runs, the same dataset and the same pipeline. Several "
    f"significant wins also sit inside the practical-tie band above, being differences that "
    f"are reliable but too small to matter. The tally is reported for completeness; the "
    f"claims of this paper rest on the primary comparison of Section 6.1 and on the "
    f"per-benchmark effects, not on this count."
)

add_bold_paragraph(
    "Table 3. Statistical summary: \u03c4(w) vs. other methods across the 8 single-scale benchmarks. "
    "The 'Sig. wins' and 'Sig. losses' columns count pairwise t-tests reaching "
    "significance after Benjamini\u2013Hochberg correction (q = 0.05) applied symmetrically "
    "to the family of all \u03c4(w)-vs-competitor comparisons.",
    size=9, space_before=12
)

table2 = doc.add_table(rows=1, cols=6)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["Benchmark", "Rank", "Wins", "Sig. wins", "Losses", "Sig. losses"]
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "D9E2F3")

# Generate statistical summary dynamically
_BENCH_LABEL = {
    'sin_regression': 'sin(x)', 'complex_regression': 'Complex', 'mnist': 'MNIST',
    'cifar': 'CIFAR-10', 'vit_cifar': 'ViT-CIFAR', 'bert_sst2': 'BERT SST-2',
    'gpt2_wikitext': 'GPT-2', 'smollm2_wikitext': 'SmolLM2',
}
_bench_keys = [b[0] for b in BENCHMARKS]
_bench_labels = [_BENCH_LABEL[k] for k in _bench_keys]

stat_data = []
_total_wins, _total_sig_wins, _total_losses, _total_sig_losses = 0, 0, 0, 0
_total_comps = 0
for label, key in zip(_bench_labels, _bench_keys):
    s = all_stats[key]
    n_comp = s['wins'] + s['losses']
    stat_data.append([
        label,
        f"{s['rank']}/{s['total']}",
        f"{s['wins']}/{n_comp}",
        str(s['sig_wins']),
        f"{s['losses']}/{n_comp}",
        str(s['sig_losses']),
    ])
    _total_wins += s['wins']
    _total_sig_wins += s['sig_wins']
    _total_losses += s['losses']
    _total_sig_losses += s['sig_losses']
    _total_comps += n_comp

stat_data.append([
    "TOTAL", "",
    f"{_total_wins}/{_total_comps}", str(_total_sig_wins),
    f"{_total_losses}/{_total_comps}", str(_total_sig_losses),
])

for row_data in stat_data:
    row = table2.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_data[0] == "TOTAL":
            run.bold = True
            set_cell_shading(cell, "E2EFDA")

# Dynamic summary text \u2014 data-driven and BH-FDR aware (STAT-1).
_bench_label_map = dict(zip(_bench_keys, _bench_labels))

def _fmt_loss(bn, m, d):
    return f"{_bench_label_map.get(bn, bn)} vs. {m} (p = {d['p']:.3f})"

_bh_losses = [(bn, m, d)
              for bn, s in all_stats.items()
              for m, d in s['details'].items() if d.get('sig_bh') and not d['better']]
_bh_wins   = [(bn, m, d)
              for bn, s in all_stats.items()
              for m, d in s['details'].items() if d.get('sig_bh') and d['better']]
_bonf_losses = [(bn, m, d)
                for bn, s in all_stats.items()
                for m, d in s['details'].items() if d.get('sig_bonf') and not d['better']]
_bonf_wins   = [(bn, m, d)
                for bn, s in all_stats.items()
                for m, d in s['details'].items() if d.get('sig_bonf') and d['better']]

# Name every corrected loss with its size, direction and Bonferroni status, so the reader
# is never asked to trust a ratio.
_metric_of = dict((b[0], b[1]) for b in BENCHMARKS + SCALE_BENCHMARKS)
def _loss_item(bn, m, d):
    _met = _metric_of[bn]
    _sc = 100.0 if _met == 'test_acc' else 1.0
    _delta = abs(d['m_mean'] - d['tau_mean']) * _sc
    _unit = {'test_acc': ' points', 'test_ppl': ' PPL'}.get(_met, '')
    return (f"{m} on {_AB_LABEL.get(bn, bn)} ({_delta:.2f}{_unit} ahead of τ(w), "
            f"p = {d['p']:.4f}, {'also' if d.get('sig_bonf') else 'not'} significant under "
            f"Bonferroni)")
_loss_per_comp = {}
for _bn, _m, _d in _bh_losses:
    _loss_per_comp[_m] = _loss_per_comp.get(_m, 0) + 1
_max_comp_losses = max(_loss_per_comp.values()) if _loss_per_comp else 0
_lm_keys = {'gpt2_wikitext', 'gpt2_tiny_wikitext', 'gpt2_medium_wikitext',
            'gpt2_large_wikitext', 'gpt2_wt103'}
_n_wins_lm = sum(1 for bn, _m, _d in _bh_wins if bn in _lm_keys)
_loss_clause = _and_join(_loss_item(*x) for x in _bh_losses) if _bh_losses else "none"

add_text(
    f"After family-wide correction the tally is {len(_bh_wins)} significant wins and "
    f"{len(_bh_losses)} significant loss{'es' if len(_bh_losses) != 1 else ''} across all "
    f"{len(_pvals_flat)} τ(w)-vs-competitor pairwise t-tests (the {_total_comps} single-scale "
    f"comparisons of Table 3 plus the scale sweep and the 124M WikiText-103 confirmation) under "
    f"Benjamini–Hochberg control at q = 0.05; under Bonferroni family-wise control at α = 0.05 "
    f"the counts are {len(_bonf_wins)} and {len(_bonf_losses)}. We name the losses rather than "
    f"fold them into a ratio: {_loss_clause}. "
    + (f"{'Both' if len(_bh_losses) == 2 else 'These'} "
       f"{'are' if len(_bh_losses) != 1 else 'is'} on classic benchmarks outside the regime in which "
       f"τ-decay is proposed, and {'both exceed' if len(_bh_losses) != 1 else 'it exceeds'} the "
       f"practical margin of Section 4.3, so {'they are' if len(_bh_losses) != 1 else 'it is'} "
       f"real loss{'es' if len(_bh_losses) != 1 else ''}, not ties. "
       if _bh_losses and all(bn in ('mnist', 'cifar', 'sin_regression', 'complex_regression', 'vit_cifar', 'bert_sst2')
                             for bn, _m, _d in _bh_losses) else "")
    + (f"No competitor beats τ(w) significantly on more than one benchmark. "
       if _max_comp_losses <= 1 else
       f"{max(_loss_per_comp, key=_loss_per_comp.get)} beats τ(w) significantly on "
       f"{_max_comp_losses} benchmarks. ")
    + f"The wins concentrate where the claim lives: {_n_wins_lm} of the {len(_bh_wins)} are on the "
    f"from-scratch autoregressive-LM benchmarks, and the remaining {len(_bh_wins) - _n_wins_lm} "
    f"are wins over penalties that degrade or collapse (LSP, L1) or over loss-side penalties on "
    f"the pretrained LLM, where the difference is inside the practical-tie band."
)

# ERRATUM (REVIEWER, pointwise 4): the caption described the 8 single-scale benchmarks
# but quoted the win/loss counts of the FULL family, which also includes the scale sweep
# and WikiText-103. Both figures are now computed on the figure's own scope, and the
# family-wide counts are named as such.
_fig4_wins = sum(1 for (bn, _m), ok in zip(_keys_flat, _bh_mask)
                 if ok and bn in _bench_keys and all_stats[bn]['details'][_m]['better'])
_fig4_losses = sum(1 for (bn, _m), ok in zip(_keys_flat, _bh_mask)
                   if ok and bn in _bench_keys and not all_stats[bn]['details'][_m]['better'])
add_figure('fig5_wins_losses.png',
    f'Figure 4. Statistical wins and losses of \u03c4(w) vs. its competitors across '
    f'the {len(_bench_keys)} single-scale benchmarks. Green = wins (higher is better); red = losses. '
    f'Dark shading = statistically significant under family-wide Benjamini\u2013Hochberg '
    f'correction (q = 0.05). Within the scope of this figure \u03c4(w) accumulates '
    f'{_fig4_wins} significant wins and {_fig4_losses} significant loss'
    f'{"es" if _fig4_losses != 1 else ""}; the corresponding counts over the whole '
    f'{len(_pvals_flat)}-test family, which also covers the scale sweep and WikiText-103, '
    f'are {_sw} and {_sl} and are reported in Section 5.7. These counts are numbers of '
    f'significant comparisons, not of independent findings: a single benchmark contributes '
    f'one comparison per competitor.',
    width_inches=6.0)

# ── 5.7 Architectural affinity ──
doc.add_heading("5.8 Architectural affinity", level=2)

add_bold_paragraph(
    "Table 4. \u03c4(w) performance by architecture family.",
    size=9, space_before=12
)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ["Architecture", "Benchmarks", "Avg. rank", "First place"]
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "D9E2F3")

# Compute architectural affinity dynamically
_arch_groups = {
    'Feedforward (regression)': ['sin_regression', 'complex_regression'],
    'CNN (classification)': ['mnist', 'cifar'],
    'Vision transformer': ['vit_cifar'],
    'Language models': ['bert_sst2', 'gpt2_wikitext', 'smollm2_wikitext'],
}
# A family with no benchmark left in scope gets no row in Table 4.
_arch_groups = {k: v for k, v in _arch_groups.items() if v}

arch_data = []
for arch_name, keys in _arch_groups.items():
    ranks = [all_stats[k]['rank'] for k in keys]
    firsts = sum(1 for k in keys if all_stats[k]['rank'] == 1)
    avg_rank = np.mean(ranks)
    arch_data.append([arch_name, str(len(keys)), f"{avg_rank:.1f}", f"{firsts}/{len(keys)}"])

# Add "All transformers" row
_trans_keys = _arch_groups.get('Vision transformer', []) + _arch_groups.get('Language models', [])
_trans_ranks = [all_stats[k]['rank'] for k in _trans_keys]
_trans_firsts = sum(1 for k in _trans_keys if all_stats[k]['rank'] == 1)
_trans_avg = np.mean(_trans_ranks)
arch_data.append(["All transformers", str(len(_trans_keys)), f"{_trans_avg:.1f}", f"{_trans_firsts}/{len(_trans_keys)}"])

for row_data in arch_data:
    row = table3.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_data[0] in ["Language models", "All transformers"]:
            run.bold = True
            if i >= 2:
                set_cell_shading(cell, "FFF2CC")

# Dynamic affinity summary
_ff_keys = _arch_groups['Feedforward (regression)'] + _arch_groups['CNN (classification)']
_ff_avg = np.mean([all_stats[k]['rank'] for k in _ff_keys])
_ff_firsts = sum(1 for k in _ff_keys if all_stats[k]['rank'] == 1)
_vit_rank = all_stats['vit_cifar']['rank']
_vit_total = all_stats['vit_cifar']['total']
_trans_first_names = [k for k in _trans_keys if all_stats[k]['rank'] == 1]
add_text(
    f"Table 4 shows the pattern: on feedforward and CNN architectures τ-decay averages rank "
    f"{_ff_avg:.1f} and finishes first on {_ff_firsts} of {len(_ff_keys)} benchmarks, while on "
    f"the transformer benchmarks of Table 1 it averages rank {_trans_avg:.1f} and finishes first "
    f"on {_trans_firsts} of {len(_trans_keys)}"
    + (f" — the first place{'s' if _trans_firsts != 1 else ''} being "
       f"{', '.join(_BENCH_LABEL.get(k, k) for k in _trans_first_names)}" if _trans_firsts else "")
    + ". We resist the obvious reading of that contrast. Four factors vary together across "
    f"these rows — architecture, task type, model size, and whether training starts from "
    f"initialisation or from a pretrained checkpoint — and the design does not separate them. "
    f"The grouping in Table 4 is by architecture because that is how the benchmarks were "
    f"originally organised, not because architecture has been shown to be the operative "
    f"variable."
)

add_text(
    f"Two observations argue against the architectural reading specifically. The vision "
    f"transformer is a transformer, and \u03c4-decay ranks {_vit_rank}/{_vit_total} on it. "
    f"The pretrained SmolLM2 is a transformer, and there the effect is negligible in "
    f"practical terms. What the strong results share is not attention but a regime: "
    f"autoregressive language models with capacity large relative to their training corpus, "
    f"trained from scratch. The supportable statement is an affinity with over-capacity "
    f"from-scratch autoregressive-LM training, and that is how we state it throughout. "
    f"Establishing an architectural effect proper would require a controlled design we did "
    f"not run \u2014 MLP, CNN and transformer on the same task at comparable capacity, the same "
    f"optimizer, the same number of examples and steps, all trained from initialisation."
)

_lm_avg = np.mean([all_stats[k]['rank'] for k in _arch_groups['Language models']])
_reg_avg = np.mean([all_stats[k]['rank'] for k in _arch_groups['Feedforward (regression)']])
add_figure('fig3_architectural_affinity.png',
    'Figure 5. Architectural affinity of \u03c4(w). '
    '(A) Ranking of \u03c4(w) on each benchmark, coloured by architecture family. '
    'Gold stars mark first-place finishes. '
    f'(B) Average rank by architecture: \u03c4(w) averages rank {_lm_avg:.1f} on language models '
    f'vs. {_reg_avg:.1f} on regression.',
    width_inches=6.5)

# ── 5.8 LSP instability ──
doc.add_heading("5.9 LSP instability", level=2)

_lsp_mnist = get_val(all_data['mnist'], 'LSP', 'test_acc', 'mean') * 100
_lsp_vit = get_val(all_data['vit_cifar'], 'LSP', 'test_acc', 'mean') * 100
_lsp_bert = get_val(all_data['bert_sst2'], 'LSP', 'test_acc', 'mean') * 100
_lsp_gpt = get_val(all_data['gpt2_wikitext'], 'LSP', 'test_ppl', 'mean')
_tau_gpt = get_val(all_data['gpt2_wikitext'], 'τ(w)', 'test_ppl', 'mean')
_l1_gpt = get_val(all_data['gpt2_wikitext'], 'L1', 'test_ppl', 'mean')
_gpt_base = get_val(all_data['gpt2_wikitext'], 'Baseline', 'test_ppl', 'mean')
_lsp_cases = 2 + (1 if _lsp_vit == _lsp_vit else 0)
_vit_all_means = [get_val(all_data['vit_cifar'], _m_, 'test_acc', 'mean') * 100
                  for _m_ in TABLE_METHODS if _m_ in all_data['vit_cifar']['method'].values]
add_text(
    f"A secondary but important finding concerns the log-sum penalty (LSP). While LSP performs "
    f"adequately on simple tasks (MNIST: {_lsp_mnist:.2f}%, the best of the field there), it "
    f"degrades on larger models, in one case to the point of outright failure. The "
    f"{'three' if _lsp_cases == 3 else 'two'} cases differ in severity and we distinguish them "
    f"rather than grouping them under a single label:"
)
if _lsp_vit == _lsp_vit:
    add_bullet(f"ViT-CIFAR: {_lsp_vit:.2f}% — "
               f"{'the worst of the field, though ' if _lsp_vit <= min(_vit_all_means) else ''}"
               f"well above the 10% chance level of 10-class CIFAR: a degradation, not a "
               f"collapse")
add_bullet(f"BERT SST-2: {_lsp_bert:.2f}% — this one is a genuine collapse to chance on a "
           f"binary task, and is the only unambiguous failure")
add_bullet(f"GPT-2 ({_MP['small']:.1f}M): perplexity {_lsp_gpt:.1f} against {_gpt_base:.1f} for the "
           f"unregularized baseline and {_tau_gpt:.1f} for τ(w) — a clear deterioration, but the "
           f"model still learns; it is not a divergence. L1 shows the same pattern here "
           f"({_l1_gpt:.1f}), and for both the tuned coefficient, not the penalty, is the "
           f"proximate cause (Section 6.4)")

add_text(
    "The pattern is that the most aggressive non-convex penalty in the set is the least "
    "reliable as models grow, which is consistent with the smooth, bounded shrinkage of the "
    "decoupled decays being easier to optimise. It is a single-penalty observation, however, "
    "and we do not read it as a general result about non-convex regularization."
)

# ============================================================================
# 5.10 Robust-decay head-to-head (REVIEWER-1)
# ----------------------------------------------------------------------------
# Nearest-neighbour comparison: same decoupled post-optimizer mechanism, same scope,
# same PSO budget for every method; only the saturation profile of the shrinkage
# varies. Produced by:
#   python gpt2_wikitext_standardized.py --scale large --robust-roster --instrument
# which writes new results/gpt2_large_wikitext_standardized_results_robust.csv.
#
# NUMBERS GUARD: if that CSV is absent the section emits an explicit "not yet run"
# note and the build prints a warning. It must NEVER invent or carry over numbers.
# ============================================================================
doc.add_heading("5.10 Head-to-head against the robust-decay family", level=2)

_ROBUST_METHODS_ORDER = ['Baseline', 'WD-tuned', 'Tau(alpha=0)', 'Huber-decay',
                         'PseudoHuber-decay', 'LogCosh-decay', '\u03c4(w)']
_ROBUST_LABELS = {
    'Baseline': 'Baseline (unregularized)',
    'WD-tuned': 'WD-tuned (decoupled \u21132)',
    'Tau(alpha=0)': 'Tau(\u03b1 = 0) [ablation]',
    'Huber-decay': 'Huber decay (AdamHD)',
    'PseudoHuber-decay': 'Pseudo-Huber / Charbonnier decay',
    'LogCosh-decay': 'Log-cosh decay',
    '\u03c4(w)': '\u03c4(w) [Fair]',
}
_ROBUST_CSV = os.path.join(ROOT_DIR, 'new results',
                           'gpt2_large_wikitext_standardized_results_robust.csv')
_ROBUST_N = 5                                   # seeds of the three robust decays
_ROBUST_N_SHARED = N_SEEDS.get('gpt2_large_wikitext', 3)   # the four methods shared with §5.6
_ROBUST_SHARED = {'Baseline', 'WD-tuned', 'Tau(alpha=0)', 'τ(w)'}

add_text(
    "Sections 5.1\u20135.9 compare \u03c4-decay with the penalties it is conventionally set against. "
    "That family, however, is not the nearest one: as Section 3.1 establishes, \u03c4-decay is the "
    "decoupled step on the Fair penalty, and its true neighbours are the other decoupled decays "
    "derived from robust, saturating losses \u2014 above all the Huber decay of Guo & Fan (2025), "
    "developed for the same from-scratch language-model pre-training setting. A comparison "
    "against SCAD, MCP and LSP cannot speak to that proximity. This section therefore isolates "
    "the one quantity that separates \u03c4-decay from its neighbours: the shape of the transition "
    "between quadratic and linear shrinkage."
)

add_text(
    "The comparison is run on the 66M-parameter from-scratch GPT-2 configuration of Section 5.6 "
    "\u2014 the operating point at which the \u03c4-decay advantage is largest, and therefore the one "
    "where a difference between saturation profiles has the most room to show itself. Every "
    "decay method is applied at the same call site (immediately after optimizer.step), to the "
    "same parameter scope, and is tuned by PSO over the same two-dimensional (\u03bb, \u03b4) space "
    "described in Section 4.2. Under this protocol Huber decay, pseudo-Huber decay and log-cosh "
    "decay differ from \u03c4(w) in nothing but the transition profile, and Tau(\u03b1 = 0) collapses "
    "the transition entirely."
)

if os.path.exists(_ROBUST_CSV):
    _rb_df = _pd.read_csv(_ROBUST_CSV)

    # If the comparison was assembled from a partial run (--robust-new-only plus
    # analysis/merge_robust_results.py), say so in the text rather than presenting it
    # as one self-contained experiment.
    _rb_prov_path = _ROBUST_CSV.replace('.csv', '.provenance.json')
    if os.path.exists(_rb_prov_path):
        add_text(
            "For economy of compute, the four methods this comparison shares with Section 5.6 "
            "(Baseline, WD-tuned, Tau(α = 0) and τ(w)) are the runs already reported there and "
            "were not retrained; only the three new competitors were trained for this section. "
            "Both invocations use the same script, scale configuration, seeds, epoch budget and "
            "PSO budget, and the pipeline is deterministic — a property verified when the 18M "
            "configuration was repeated and returned identical numbers."
        )

    def _rb_get(method, stat):
        row = _rb_df[_rb_df['method'] == method]
        if row.empty:
            return float('nan')
        return float(row.iloc[0][f'test_ppl_{stat}'])

    def _rb_n(method):
        return _ROBUST_N_SHARED if method in _ROBUST_SHARED else _ROBUST_N

    def _rb_d(ma, sa, na, mb, sb, nb):
        _pooled = np.sqrt(((na - 1) * sa ** 2 + (nb - 1) * sb ** 2) / (na + nb - 2))
        return 0.0 if _pooled == 0 else (ma - mb) / _pooled

    add_text(
        "One asymmetry of the protocol must be stated. The three robust decays were tuned "
        "with an earlier two-dimensional PSO budget of 24 evaluations, τ(w) with the current "
        "40 (Appendix A.1). Their runs never stopped before the final epoch and are unaffected by "
        "it — none of them stopped before its final epoch, so the reported model is the "
        "best-epoch model — and the smaller budget can only have handicapped them, so any "
        "parity reported below is conservative with respect to the competitors."
    )

    _rb_tau_m, _rb_tau_s = _rb_get('τ(w)', 'mean'), _rb_get('τ(w)', 'std')

    _rb_table = doc.add_table(rows=1, cols=6)
    _rb_table.style = 'Table Grid'
    _rb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _rb_hdr = _rb_table.rows[0].cells
    for _i, _h in enumerate(['Method', 'n', 'Test PPL (mean ± SD)', 'Δ vs τ(w)',
                             'p (Welch)', "Cohen's d"]):
        _rb_hdr[_i].text = ""
        _rp = _rb_hdr[_i].paragraphs[0]
        _rr = _rp.add_run(_h)
        _rr.font.size = Pt(9)
        _rr.font.name = 'Times New Roman'
        _rr.bold = True
        _rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(_rb_hdr[_i], 'D9D9D9')

    for _m in _ROBUST_METHODS_ORDER:
        _mm, _ms_ = _rb_get(_m, 'mean'), _rb_get(_m, 'std')
        if _mm != _mm:  # NaN -> method absent from this run
            continue
        if _m == 'τ(w)':
            _cells = [_ROBUST_LABELS[_m], str(_rb_n(_m)), f'{_mm:.2f} ± {_ms_:.2f}', '—', '—', '—']
        else:
            _t, _p = _welch2(_mm, _ms_, _rb_n(_m), _rb_tau_m, _rb_tau_s, _rb_n('τ(w)'))
            _d = _rb_d(_mm, _ms_, _rb_n(_m), _rb_tau_m, _rb_tau_s, _rb_n('τ(w)'))
            _cells = [_ROBUST_LABELS[_m], str(_rb_n(_m)), f'{_mm:.2f} ± {_ms_:.2f}',
                      f'{_mm - _rb_tau_m:+.2f}', '< 0.0001' if _p < 1e-4 else f'{_p:.4f}',
                      f'{abs(_d):.2f}']
        add_table_row(_rb_table, _cells,
                      bold=(_m == 'τ(w)'),
                      bg_color=('F2F2F2' if _m == 'τ(w)' else None))

    _p_rb = doc.add_paragraph()
    _r_rb = _p_rb.add_run(
        f'Table 5. Robust-decay head-to-head on the 66M from-scratch GPT-2 configuration '
        f'(WikiText-2). n = {_ROBUST_N} seeds for the three robust decays; the four shared '
        f'methods are the n = {_ROBUST_N_SHARED} runs of Section 5.6, and each Welch test uses '
        f'the two actual sample sizes. All decay methods share the post-optimizer call site and '
        f'the parameter scope; only the saturation profile of the shrinkage differs. p-values '
        f'are Welch’s t-test against τ(w); a positive Δ favours τ(w) (lower '
        f'perplexity is better).')
    _r_rb.font.size = Pt(9)
    _r_rb.font.name = 'Times New Roman'
    _r_rb.italic = True

    # Prose is generated from the table, so it can never contradict it.
    _rb_rivals = [m for m in ('Huber-decay', 'PseudoHuber-decay', 'LogCosh-decay')
                  if _rb_get(m, 'mean') == _rb_get(m, 'mean')]
    _rb_best = min(_rb_rivals, key=lambda m: _rb_get(m, 'mean')) if _rb_rivals else None
    if _rb_best is not None:
        _bm, _bs = _rb_get(_rb_best, 'mean'), _rb_get(_rb_best, 'std')
        _bt, _bp = _welch2(_bm, _bs, _rb_n(_rb_best), _rb_tau_m, _rb_tau_s, _rb_n('τ(w)'))
        _rb_verdict = (
            "τ(w) is ahead of the best robust-decay competitor "
            if _bm > _rb_tau_m else
            "τ(w) does not lead the robust-decay family: the best competitor is ahead ")
        _rb_sig = ("a difference that reaches significance at α = 0.05"
                   if _bp < 0.05 else
                   "a difference that does not reach significance at α = 0.05")
        # Whether the whole family beats Baseline / WD-tuned is a claim about the data,
        # so derive it rather than assert it.
        _rb_wd = _rb_get('WD-tuned', 'mean')
        _rb_base = _rb_get('Baseline', 'mean')
        _rb_a0 = _rb_get('Tau(alpha=0)', 'mean')
        _rb_family = [_rb_get(m, 'mean') for m in _rb_rivals + ['τ(w)']]
        _rb_all_beat = all(v < _rb_wd for v in _rb_family) and all(v < _rb_base
                                                                  for v in _rb_family)
        _rb_fam_all = _rb_family + ([_rb_a0] if _rb_a0 == _rb_a0 else [])
        _rb_spread = max(_rb_fam_all) - min(_rb_fam_all)
        _rb_eq = [m for m in _rb_rivals
                  if tost_equivalence(_rb_get(m, 'mean'), _rb_get(m, 'std'), _rb_tau_m,
                                      _rb_tau_s, margin=EQUIV_MARGIN_PPL, n=_ROBUST_N)[0]]
        _rb_family_note = (
            f"Every member of the family, τ(w) included, improves on both the unregularized "
            f"baseline and tuned decoupled weight decay by {min(_rb_wd - v for v in _rb_family):.1f} "
            f"PPL or more, and the four adaptive profiles together with the constant-rate "
            f"ablation Tau(α = 0) all fall inside a {_rb_spread:.2f}-PPL band"
            + (f"; {len(_rb_eq)} of the {len(_rb_rivals)} robust decays "
               f"({', '.join(_ROBUST_LABELS[m].split(' (')[0].split(' /')[0] for m in _rb_eq)}) "
               f"{'are' if len(_rb_eq) != 1 else 'is'} demonstrably equivalent to τ(w) within the "
               f"pre-specified ±{EQUIV_MARGIN_PPL} PPL margin" if _rb_eq else "")
            + ". This is the result that carries the paper’s claim: what buys the gain at "
            "this scale is the decoupled, weight-only, schedule-independent decay as such, and "
            "the choice of saturation profile — Fair, Huber, pseudo-Huber, log-cosh or none "
            "— is a refinement within it whose value is at most a few tenths of a point."
            if _rb_all_beat else
            "Not every member of the family improves on both the unregularized baseline and "
            "tuned decoupled weight decay at this operating point, so the saturation profile "
            "is not interchangeable: the per-method figures in the table, not the family "
            "membership, carry the comparison."
        )
        add_text(
            f"The strongest neighbour is {_ROBUST_LABELS[_rb_best]} at "
            f"{_bm:.2f} ± {_bs:.2f} PPL, against {_rb_tau_m:.2f} ± {_rb_tau_s:.2f} for "
            f"τ(w). {_rb_verdict}by {abs(_bm - _rb_tau_m):.2f} PPL (p = {_bp:.4f}), "
            f"{_rb_sig}. {_rb_family_note}"
        )
else:
    print('[WARN] build_paper: robust-decay head-to-head CSV not found at\n'
          f'       {_ROBUST_CSV}\n'
          '       Section 5.10 will be emitted as an explicit "not yet run" placeholder.\n'
          '       Produce it with:  python gpt2_wikitext_standardized.py '
          '--scale large --robust-roster', file=sys.stderr)
    _pl = doc.add_paragraph()
    _plr = _pl.add_run(
        '[PENDING RUN \u2014 NOT FOR SUBMISSION] The head-to-head described above has not yet '
        'been executed; no numbers are reported here. This placeholder is emitted '
        'automatically whenever gpt2_large_wikitext_standardized_results_robust.csv is '
        'absent, so that the manuscript can never present this comparison as done when it '
        'is not.')
    _plr.font.size = Pt(11)
    _plr.font.name = 'Times New Roman'
    _plr.bold = True
    _plr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ============================================================================
# 6. DISCUSSION
# ============================================================================

doc.add_heading("6. Discussion", level=1)

doc.add_heading("6.1 What the evidence supports: a decoupled decay for over-capacity autoregressive-LM training", level=2)

add_text(
    "The central recommendation of this work is concrete and scoped. When an autoregressive "
    "language model trained from scratch overfits its training budget — when the unregularized "
    "model's validation optimum precedes the end of training — a decoupled decay applied to "
    "the weight tensors only (biases excluded), at a rate independent of the learning-rate schedule, outperforms "
    f"PSO-tuned AdamW weight decay by a consistent, statistically decisive margin — "
    f"{_gap('gpt2_large_wikitext', 'WD-tuned'):.1f} PPL at 66M, about "
    f"{_gap('gpt2_large_wikitext', 'WD-tuned') / _dv('WD-tuned') * 100:.0f}% of the tuned "
    "competitor's perplexity — whether its profile is constant "
    "(Tau(α = 0)) or magnitude-adaptive (τ(w)). Where that pressure is absent — shorter budgets, "
    "larger corpora, pretrained models, the other architectures we tested — the same decay "
    "degrades to parity with tuned weight decay. Across the suite of competing methods — L1, L2, "
    "ElasticNet, SCAD, MCP, the log-sum penalty and decoupled weight decay — the τ family is "
    "distinguished by the following properties, which we examine in turn."
)

_ss_base_d = abs(_dd(_ss, 'Baseline'))
_ls_base_d = abs(_dd(_ls, 'Baseline'))
p = doc.add_paragraph()
run = p.add_run("(i) Large, regime-dependent gains on from-scratch autoregressive LMs. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"The from-scratch scale sweep — {_MP['tiny']:.1f}M, {_MP['small']:.1f}M, {_MP['medium']:.0f}M "
    f"and {_MP['large']:.0f}M parameters trained on WikiText-2 with the same recipe and core "
    f"roster — is consistent with the claim but varies model size at fixed data, and therefore "
    f"suggests rather than identifies the mechanism. At {_MP['tiny']:.1f}M every method ties the "
    f"unregularized Baseline (τ(w) {_sgn(_pct('gpt2_tiny_wikitext'), 1)}%): the correct null where "
    f"there is no overfitting to prevent. At {_MP['medium']:.0f}M on 12 epochs, where no method "
    f"overfits inside the budget, τ(w) leads tuned decoupled weight decay by "
    f"{_gap('gpt2_medium_wikitext', 'WD-tuned'):.1f} PPL and the Baseline by "
    f"{_gap('gpt2_medium_wikitext', 'Baseline'):.1f} PPL "
    f"({abs(_pct('gpt2_medium_wikitext')):.1f}%). At the two points where the Baseline does "
    f"overfit — {_MP['small']:.1f}M on 30 epochs and {_MP['large']:.0f}M on 12 — the τ family "
    f"leads tuned weight decay by {_gap('gpt2_wikitext', 'WD-tuned'):.1f} and "
    f"{_gap('gpt2_large_wikitext', 'WD-tuned'):.1f} PPL and the Baseline by "
    f"{abs(_pct('gpt2_wikitext')):.0f}% and {abs(_pct('gpt2_large_wikitext')):.0f}% "
    f"({_pfmt(max(_pget(_ss, 'Baseline'), _pget(_ls, 'Baseline')))}, "
    f"|d| ≥ {min(_ss_base_d, _ls_base_d):.0f}), with the Baseline and WD-tuned deteriorating from "
    f"their best epoch while both τ-family runs are still improving when the budget ends "
    f"(Section 6.2). The ordering follows overfitting pressure, not size: the "
    f"{_MP['small']:.1f}M model gains as much as the {_MP['large']:.0f}M one because it is trained "
    f"for two and a half times as many epochs. The 124M WikiText-103 confirmation (Section 5.6) "
    f"closes the claim from the data side: at the standard GPT-2-small operating point, where a "
    f"118M-token corpus leaves no over-capacity to exploit, τ(w) still clears the unregularized "
    f"Baseline by {_gap('gpt2_wt103', 'Baseline'):.2f} PPL"
    f"{' (Bonferroni-significant)' if _w3_base_bonf else ''} and ties tuned decoupled weight "
    f"decay — the advantage compresses exactly where the reading predicts, and does not invert "
    f"into harm."
)
run.font.name = 'Times New Roman'

# Which benchmarks carry WD-tuned, where τ(w) loses to it after correction, and where the two
# are demonstrably equivalent (classification / LM metrics only: the MSE benchmarks have no
# pre-specified margin).
_wd_all = [(b, m) for b, m, _, _, _ in BENCHMARKS + SCALE_BENCHMARKS
           if 'WD-tuned' in set(all_data[b]['method'])]
_wd_bh_losses = [bn for bn, m, d in _bh_losses if m == 'WD-tuned']
_wd_testable = [(b, m) for b, m in _wd_all if m in ('test_ppl', 'test_acc')]
_wd_eq = [b for b, m in _wd_testable if equiv_verdict(b, 'τ(w)', 'WD-tuned', metric=m)[1]]
# Do the non-convex penalties ever separate from the unregularized Baseline?
_ncx_sep = []
for _b, _m, _mode, _sc, _u in BENCHMARKS:
    for _pen in ('SCAD', 'MCP'):
        _df_ = all_data[_b]
        if _pen not in _df_['method'].values:
            continue
        _, _pp = welch_ttest(get_val(_df_, _pen, _m, 'mean'), get_val(_df_, _pen, _m, 'std'),
                             get_val(_df_, 'Baseline', _m, 'mean'), get_val(_df_, 'Baseline', _m, 'std'),
                             n=N_SEEDS.get(_b, 5))
        if _pp < 0.05:
            _ncx_sep.append(f"{_pen} on {_AB_LABEL.get(_b, _b)}")

p = doc.add_paragraph()
run = p.add_run("(ii) Statistically indistinguishable from tuned weight decay outside the regime, with named exceptions. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"Across the {len(_pvals_flat)} τ(w)-vs-competitor pairwise Welch t-tests in the suite, with "
    f"family-wide Benjamini–Hochberg correction (q = 0.05) applied symmetrically to wins and "
    f"losses, τ-decay registers {_sw} significant wins against {_sl} significant "
    f"loss{'es' if _sl != 1 else ''}; under the more conservative Bonferroni correction "
    f"(α = 0.05), the counts become {_swB} wins and {_slB} losses. The loss"
    f"{'es are' if len(_bh_losses) != 1 else ' is'} "
    f"{_and_join(f'{m} on {_AB_LABEL.get(bn, bn)}' for bn, m, d in _bh_losses) if _bh_losses else 'none'}"
    + (": classic benchmarks on which a different regularizer is genuinely better and "
       "τ-decay, tuned by the same protocol, is not. " if _bh_losses else ". ")
    + f"Against tuned decoupled weight decay specifically — the comparison that matters for "
    f"adoption — τ(w) is "
    f"{'never significantly worse on any benchmark' if not _wd_bh_losses else 'significantly worse on ' + _and_join(_AB_LABEL.get(b, b) for b in _wd_bh_losses)}, "
    f"and it is demonstrably equivalent within the pre-specified margins on {len(_wd_eq)} of the "
    f"{len(_wd_testable)} classification and language-modelling benchmarks that carry it "
    f"({_and_join(_AB_LABEL.get(b, b) for b in _wd_eq)}); on the others the difference is "
    f"either in τ(w)'s favour or inconclusive at the sample size. Adopting the τ family "
    f"therefore carries a small and known downside relative to standard practice — we do not "
    f"claim it never loses — while the loss-side penalties do not share even this: L1 and the "
    f"log-sum penalty win individual classic benchmarks but degrade or collapse on the "
    f"from-scratch language models (LSP on BERT-tiny and on the {_MP['small']:.1f}M GPT-2; L1 on "
    f"the same GPT-2), and SCAD and MCP "
    f"{'never separate from the unregularized Baseline' if not _ncx_sep else 'separate from the unregularized Baseline only for ' + _and_join(_ncx_sep)}."
)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("(iii) Graceful no-harm degradation when no improvement is available. ")
run.bold = True
run.font.name = 'Times New Roman'
_sm_pen_lo = min(_sm_l2, _sm_en) - _sm_base
_sm_pen_hi = max(_sm_l2, _sm_en) - _sm_base
_sm_near = [m.replace('Tau(alpha=0)', 'Tau(α = 0)') for m in TABLE_METHODS
            if m in _sm_df['method'].values and m != 'Baseline'
            and abs(get_val(_sm_df, m, 'test_ppl', 'mean') - _sm_base) <= 0.02]
run = p.add_run(
    f"On a modern small LLM already saturated by pretraining (SmolLM2-135M), τ-decay finishes "
    f"within {abs(_sm_tau - _sm_base):.3f} PPL of the unregularized Baseline — marginally "
    f"{'ahead' if _sm_tau < _sm_base else 'behind'} ({_sm_tau:.3f} vs {_sm_base:.3f}) — a margin "
    f"that is operationally negligible in either direction, because the tuning search switched "
    f"the decay off (Section 5.5). This is the failure mode one wants: where no regularizer can "
    f"extract additional signal, the method converges back to the Baseline instead of to a worse "
    f"operating point. The loss-side penalties L2 and ElasticNet, by contrast, cost "
    f"{_sm_pen_lo:.2f}–{_sm_pen_hi:.2f} PPL on the same budget. The property is not unique to "
    f"τ(w): {_and_join(_sm_near)} all finish within 0.02 PPL of the Baseline, which is the point "
    f"— a decoupled decay whose strength is tuned on validation data can always retreat to zero, "
    f"and a loss-side penalty tuned the same way should be able to as well, but the two we ran "
    f"did not."
)
run.font.name = 'Times New Roman'

add_text(
    "The practical adoption cost is correspondingly small. τ-decay is a one-line update applied "
    "after the standard optimizer step, w ← w − ρ·w/(1 + |w|/δ), compatible with any first-order "
    "optimizer and adding only constant per-step compute. Its two identifiable hyperparameters "
    "(ρ, δ) admit a narrow, transferable search range — discussed in Section 6.4 — and the "
    "constant-rate special case has one."
)

# ============================================================================
# Scope x adaptivity factorial (REVIEWER-5). The single most important correction to the
# previous version of this manuscript: tau-decay's margin over tuned AdamW is NOT mostly
# magnitude adaptivity. Every number below is read from the 66M CSV / per-seed JSON.
# ============================================================================
add_bold_paragraph("Decomposing the margin: the scope × adaptivity factorial", size=12,
                   space_before=10)

_dsd = lambda m: get_val(_d66, m, 'test_ppl', 'std')
def _pc66(a, b):
    """Paired (by seed) contrast b − a on the 66M benchmark: (Δ, p)."""
    _ra, _rb = load_runs('gpt2_large_wikitext', a, 'test_ppl'), load_runs('gpt2_large_wikitext', b, 'test_ppl')
    if not _ra or not _rb:
        return float('nan'), float('nan')
    _t, _p_, _md, _n = paired_ttest(_rb, _ra)
    return _md, _p_
_c_ad_adamw = _pc66('WD-tuned', 'Tau(AdamW-scope)')
_c_ad_tau = _pc66('Tau(alpha=0)', 'τ(w)')
_c_sc_const = _pc66('WD-tuned', 'Tau(alpha=0)')
_c_sc_adapt = _pc66('Tau(AdamW-scope)', 'τ(w)')

# Optional fifth cell: AdamW's LR-scaled decay restricted to the τ scope. It splits the
# column contrast into its two components — scope at the same schedule (vs WD-tuned) and
# schedule at the same scope (vs Tau(α = 0)). Rendered ONLY if the row is in the 66M CSV.
_HAS_WDW = 'WD-tuned(weights)' in set(_d66['method'].values)
if _HAS_WDW:
    _d_wdw = _dv('WD-tuned(weights)')
    _c_wdw_scope = _pc66('WD-tuned', 'WD-tuned(weights)')
    _c_wdw_sched = _pc66('WD-tuned(weights)', 'Tau(alpha=0)')

_dtab = doc.add_table(rows=1, cols=4)
_dtab.style = 'Table Grid'
_dtab.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['', 'Constant profile', 'Magnitude-adaptive profile',
                         'Row contrast: adaptivity']):
    _dtab.rows[0].cells[_i].text = ""
    _pp = _dtab.rows[0].cells[_i].paragraphs[0]
    _rr = _pp.add_run(_h); _rr.font.size = Pt(9); _rr.font.name = 'Times New Roman'
    _rr.bold = True
    _pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_dtab.rows[0].cells[_i], 'D9D9D9')
add_table_row(_dtab, ['AdamW scope + schedule',
                      f"WD-tuned\n{_d_wd:.2f} ± {_dsd('WD-tuned'):.2f}",
                      f"Tau(AdamW-scope)\n{_d_ad:.2f} ± {_dsd('Tau(AdamW-scope)'):.2f}",
                      f"{_eff_adapt_adamw:+.2f} ({_pfmt(_c_ad_adamw[1])})"])
add_table_row(_dtab, ['τ-decay scope + schedule',
                      f"Tau(α = 0)\n{_d_a0:.2f} ± {_dsd('Tau(alpha=0)'):.2f}",
                      f"τ(w)\n{_d_tau:.2f} ± {_dsd('τ(w)'):.2f}",
                      f"{_eff_adapt_tau:+.2f} ({_pfmt(_c_ad_tau[1])})"])
add_table_row(_dtab, ['Column contrast: scope + schedule',
                      f"{_eff_scope_const:+.2f} ({_pfmt(_c_sc_const[1])})",
                      f"{_eff_scope_adapt:+.2f} ({_pfmt(_c_sc_adapt[1])})",
                      f"interaction {_eff_adapt_tau - _eff_adapt_adamw:+.2f}"],
              bold=True, bg_color='F2F2F2')
if _HAS_WDW:
    add_table_row(_dtab, ['τ-decay scope, AdamW schedule',
                          f"WD-tuned(weights)\n{_d_wdw:.2f} ± {_dsd('WD-tuned(weights)'):.2f}",
                          '—',
                          f"scope: {_d_wdw - _d_wd:+.2f} ({_pfmt(_c_wdw_scope[1])}); "
                          f"schedule: {_d_a0 - _d_wdw:+.2f} ({_pfmt(_c_wdw_sched[1])})"])

_pcap = doc.add_paragraph()
_rcap = _pcap.add_run(
    f'Table 6. Scope × adaptivity factorial at 66M (from-scratch GPT-2, WikiText-2, test '
    f'perplexity, mean ± std, n = {N_SEEDS["gpt2_large_wikitext"]} seeds; unregularized '
    f'Baseline {_d_base:.2f} ± {_dsd("Baseline"):.2f} for reference). Row contrasts isolate '
    f'magnitude adaptivity at a fixed scope and schedule; column contrasts isolate scope and '
    f'schedule at a fixed profile. Negative = improvement; p-values are paired by seed.')
_rcap.font.size = Pt(9); _rcap.font.name = 'Times New Roman'; _rcap.italic = True

_share_adapt = abs(_eff_adapt_tau) / _d_total * 100
_share_scope = abs(_eff_scope_const) / _d_total * 100
_share_wd = abs(_d_base - _d_wd) / _d_total * 100
add_text(
    f"The factorial reads as follows. Moving down the constant column — from AdamW's decay to a "
    f"weight-only, schedule-independent decay at the same constant profile — is worth "
    f"{abs(_eff_scope_const):.2f} PPL. Moving across the AdamW row — switching the profile from "
    f"constant to Fair while keeping AdamW's scope and schedule — is worth "
    f"{abs(_eff_adapt_adamw):.2f} PPL — about "
    f"{abs(_eff_adapt_adamw) / abs(_eff_scope_const) * 100:.0f}% of the scope-and-schedule "
    f"effect — but moving across the τ row is worth only "
    f"{abs(_eff_adapt_tau):.2f} PPL. The two routes end {abs(_d_a0 - _d_ad):.2f} PPL apart "
    f"(Tau(α = 0) at {_d_a0:.2f}, Tau(AdamW-scope) at {_d_ad:.2f}) and the combination is best "
    f"({_d_tau:.2f}); the interaction is therefore large: adaptivity is worth "
    f"{abs(_eff_adapt_tau - _eff_adapt_adamw):.2f} PPL less once the scope is right. Magnitude "
    f"adaptivity is valuable mainly "
    f"when the decay has the wrong scope — sparing the large weights compensates for decaying "
    f"LayerNorm gains and biases at a rate tied to the learning-rate schedule — and once the "
    f"scope is right it is a second-order refinement. This changes how the headline number "
    f"should be read, and we state the reading plainly: the complete τ-decay "
    f"implementation improves perplexity by {abs(_d_total) / _d_base * 100:.1f}% over the "
    f"unregularized Baseline at 66M, of which {_share_wd:.0f}% is having any tuned decay at all, "
    f"{_share_scope:.0f}% is the scope and schedule of the decay, and {_share_adapt:.0f}% is the "
    f"magnitude-adaptive profile measured at the right scope."
)

if _HAS_WDW:
    add_text(
        f"A fifth cell refines the column contrast. WD-tuned(weights) applies AdamW's "
        f"LR-scaled decay restricted to the τ scope, so it isolates the schedule factor: "
        f"against WD-tuned it changes only the scope at the same schedule "
        f"({_d_wdw - _d_wd:+.2f} PPL, {_pfmt(_c_wdw_scope[1])}), and against Tau(α = 0) it "
        f"changes only the schedule at the same scope ({_d_a0 - _d_wdw:+.2f} PPL, "
        f"{_pfmt(_c_wdw_sched[1])})."
    )

add_bold_paragraph("The primary comparison, tested as paired data", size=12, space_before=8)

add_text(
    "Every method on a benchmark is trained on the same seed list in the same order, so runs "
    "are paired by construction: run k of τ(w) and run k of Tau(α = 0) share an "
    "initialisation and a data ordering. Treating them as independent samples, as an unpaired "
    "Welch test does, discards that structure and with it most of the power available at these "
    "sample sizes. We therefore test the primary contrast as paired data, and report a "
    "percentile bootstrap interval on the paired difference alongside the t-test, since at "
    "n = 3–10 the normal-theory interval rests on an assumption the sample cannot verify."
)

_pair_tab = doc.add_table(rows=1, cols=6)
_pair_tab.style = 'Table Grid'
_pair_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['Benchmark', 'n', 'Δ PPL (τ − α=0)', 'Paired t', 'Paired p', 'Bootstrap CI₉₅']):
    _c = _pair_tab.rows[0].cells[_i]; _c.text = ""
    _r = _c.paragraphs[0].add_run(_h); _r.bold = True
    _r.font.size = Pt(9); _r.font.name = 'Times New Roman'
    _c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_c, 'D9D9D9')

_pair_rows = []
for _bn, _label in (('gpt2_wikitext', f'GPT-2 {_MP["small"]:.1f}M (from scratch, 30 epochs)'),
                    ('gpt2_medium_wikitext', f'GPT-2 {_MP["medium"]:.0f}M (from scratch)'),
                    ('gpt2_large_wikitext', f'GPT-2 {_MP["large"]:.0f}M (from scratch)'),
                    ('gpt2_wt103', 'GPT-2 124M / WikiText-103 (transfer-tuned)'),
                    ('smollm2_wikitext', 'SmolLM2-135M (pretrained)')):
    _ta = load_runs(_bn, 'τ(w)', 'test_ppl')
    _aa = load_runs(_bn, 'Tau(alpha=0)', 'test_ppl')
    if not _ta or not _aa:
        continue
    _t, _p, _md, _n = paired_ttest(_ta, _aa)
    _lo, _hi = bootstrap_ci(_ta, _aa, seed=0)
    add_table_row(_pair_tab, [_label, str(_n), f'{_md:+.3f}', f'{_t:+.1f}',
                              '< 0.0001' if _p < 1e-4 else f'{_p:.4f}',
                              f'[{_lo:+.3f}, {_hi:+.3f}]'])
    _pair_rows.append((_bn, _md, _p, _lo, _hi, _t))

_p = doc.add_paragraph()
_r = _p.add_run(
    'Table 7. The pre-registered primary contrast — τ(w) against its implementation-matched '
    'constant-decay ablation — tested as paired data. A negative Δ favours τ(w). The bootstrap '
    'interval is the percentile interval of the paired difference over 10,000 resamples of the '
    'seed pairs. The 2M sweep member is omitted: there the two searches settled at different '
    'rates with no knee inside the weight distribution (Section 5.6).')
_r.font.size = Pt(9); _r.font.name = 'Times New Roman'; _r.italic = True

_pr = {r[0]: r for r in _pair_rows}
_smol = _pr.get('smollm2_wikitext')
_n_neg = sum(1 for r in _pair_rows if r[1] < 0)
add_text(
    f"Two things are worth drawing out. First, the paired analysis is decisive at every scale "
    f"(all p ≤ {max(r[2] for r in _pair_rows):.4f}), which is exactly why the sign pattern "
    f"matters: the increment attributable to magnitude adaptivity is "
    f"{_pr['gpt2_wikitext'][1]:+.2f} PPL at {_MP['small']:.1f}M, "
    f"{_pr['gpt2_medium_wikitext'][1]:+.2f} at {_MP['medium']:.0f}M, "
    f"{_pr['gpt2_large_wikitext'][1]:+.2f} at {_MP['large']:.0f}M and "
    f"{_pr['gpt2_wt103'][1]:+.2f} at 124M on WikiText-103 (transfer-tuned), i.e. in τ(w)'s "
    f"favour on {_n_neg} of the {len(_pair_rows)} rows and against it on the {_MP['small']:.1f}M "
    f"benchmark. The pre-registered contrast is therefore not confirmed as a scale-independent "
    f"effect: it is a small, reliably measurable, scale-dependent refinement whose sign the "
    f"tuning search can flip, and the paper's headline is drawn from the secondary contrast and "
    f"the factorial, as Section 4.3 states."
    + (f" Second, on SmolLM2 the paired difference is {_smol[1]:+.3f} PPL "
       f"(p = {_smol[2]:.4f}), far inside the ±{EQUIV_MARGIN_PPL} PPL equivalence margin, which is "
       f"the cleanest illustration in this study of why we report practical margins alongside "
       f"significance: a detectable difference can be operationally irrelevant." if _smol else "")
)

add_text(
    "Two consequences follow. First, the inferential weight of this paper rests on the τ family "
    "versus tuned AdamW decay and on the factorial that attributes that margin, not on the "
    "adaptivity contrast that gives the method its name; we keep the pre-specification and "
    "report its outcome rather than replacing it after the fact. Second, the scope-and-schedule "
    "finding is itself the result worth reporting: where a decoupled decay is applied, and "
    "whether its rate follows the learning-rate schedule, matter far more under overfitting "
    "pressure than the shape of the shrinkage. That is actionable for practitioners "
    "independently of τ(w), and it is why the factorial of Section 4.2 separates the two "
    "factors rather than reporting only their sum."
)

add_text(
    "The remainder of the discussion develops the underlying mechanisms: direct "
    "training-dynamics evidence of the overfitting delay and of what the tuned profile does "
    "(Section 6.2), why the regime rather than the architecture is the operative variable "
    "(Section 6.3), the hyperparameter sensitivity and the modern-LLM null-effect regime "
    "(Section 6.4), the comparison with non-convex penalties (Section 6.5), the sparsity "
    "patterns produced by τ-decay (Section 6.6), and methodological limitations (Section 6.7)."
)

doc.add_heading("6.2 Mechanism: overfitting delay, and what the tuned profile actually does", level=2)

_hp_m, _hp_l = _best_hp('gpt2_medium_wikitext'), _best_hp('gpt2_large_wikitext')
_hpv = lambda hp, m, k: float((hp.get(m) or {}).get(k, float('nan')))
_rho_m_tau, _del_m_tau = _hpv(_hp_m, 'τ(w)', 'rho'), _hpv(_hp_m, 'τ(w)', 'delta')
_rho_m_a0 = _hpv(_hp_m, 'Tau(alpha=0)', 'rho')
_rho_l_tau, _del_l_tau = _hpv(_hp_l, 'τ(w)', 'rho'), _hpv(_hp_l, 'τ(w)', 'delta')
_rho_l_a0 = _hpv(_hp_l, 'Tau(alpha=0)', 'rho')
_rho_l_ad, _del_l_ad = _hpv(_hp_l, 'Tau(AdamW-scope)', 'rho'), _hpv(_hp_l, 'Tau(AdamW-scope)', 'delta')
_rate = lambda rho, delta, w: rho / (1.0 + w / delta)
_n_ins_l_lo = min((v['n'] for v in _INST['large'].values()), default=0)
_n_ins_l_hi = max((v['n'] for v in _INST['large'].values()), default=0)
_n_ins_l = (f"{_n_ins_l_lo}" if _n_ins_l_lo == _n_ins_l_hi else f"{_n_ins_l_lo}–{_n_ins_l_hi}")
_n_ins_m = max((v['n'] for v in _INST['medium'].values()), default=0)
_ins_l_present = [m for m in ('Baseline', 'WD-tuned', 'Tau(AdamW-scope)', 'Tau(alpha=0)', 'τ(w)')
                  if m in _INST['large']]

add_text(
    f"Section 5.6 attributes the gain to overfitting delay; this section shows the dynamics "
    f"behind it, using instrumented runs of the {_MP['medium']:.0f}M, {_MP['large']:.0f}M and 124M "
    f"WikiText-103 from-scratch benchmarks in which every training run logged, at each epoch, the "
    f"train and validation perplexity together with summary statistics of the weight-magnitude "
    f"distribution (median, mean, maximum, L2 norm and the fraction of near-zero weights) over "
    f"all decayed parameters. Only runs of the final tuned configuration enter the analysis "
    f"(at {_MP['large']:.0f}M all {len(_INST['large'])} methods of Table 2 × {_n_ins_l} "
    f"seeds — Tau(α = 0) and Tau(AdamW-scope) carry 9 of the 10 seeds because one seed's "
    f"instrumentation predates the early-stopping fix; at {_MP['medium']:.0f}M the core six "
    f"× {_n_ins_m})."
)

_ovf_gap_l = min(_iv('large', m, 'best_val') for m in ('Baseline', 'WD-tuned')) - _iv('large', 'τ(w)', 'last_val')
_drift_l = _iv('large', 'Baseline', 'last_val') - _iv('large', 'Baseline', 'best_val')
add_text(
    f"Overfitting delay. At {_MP['large']:.0f}M — the most over-capacitated point — the delay is "
    f"directly visible (Figure 6A): the Baseline's validation perplexity bottoms out at epoch "
    f"{_iv('large', 'Baseline', 'best_ep'):.0f} and then deteriorates, finishing "
    f"{_drift_l:.1f} PPL above its own optimum, and WD-tuned turns at epoch "
    f"{_iv('large', 'WD-tuned', 'best_ep'):.0f}, while Tau(α = 0) and τ(w) are still improving "
    f"when the 12-epoch budget ends (best epoch {_iv('large', 'Tau(alpha=0)', 'best_ep'):.0f} and "
    f"{_iv('large', 'τ(w)', 'best_ep'):.0f} on the seed-mean curves) and Tau(AdamW-scope) turns "
    f"only at epoch {_iv('large', 'Tau(AdamW-scope)', 'best_ep'):.0f}; by then τ(w) is "
    f"{_ovf_gap_l:.1f} PPL below the best point either constant-penalty competitor ever reached. "
    f"The instrumented L2 and ElasticNet runs track the Baseline rather than the decays: they "
    f"turn upward at epochs {_iv('large', 'L2', 'best_ep'):.0f} and "
    f"{_iv('large', 'ElasticNet', 'best_ep'):.0f} with best validation perplexities of "
    f"{_iv('large', 'L2', 'best_val'):.1f} and {_iv('large', 'ElasticNet', 'best_val'):.1f} "
    f"against the Baseline's {_iv('large', 'Baseline', 'best_val'):.1f} — the trajectory-level "
    f"counterpart of their near-null tuned coefficients (Section 5.6). "
    f"At {_MP['medium']:.0f}M the budget ends before any method turns upward (every best epoch is "
    f"12 of 12), which is why the gain over the Baseline there is "
    f"{_gap('gpt2_medium_wikitext', 'Baseline'):.1f} PPL rather than the "
    f"{_gap('gpt2_large_wikitext', 'Baseline'):.1f} PPL of the 66M point: the same ordering "
    f"appears in the final validation perplexities, and τ(w) is the only method that combines a "
    f"small generalization gap ({_iv('medium', 'τ(w)', 'gap'):.1f} vs "
    f"{_iv('medium', 'Baseline', 'gap'):.1f} for the Baseline at the validation optimum) with the "
    f"best validation perplexity — L2 achieves a smaller gap ({_iv('medium', 'L2', 'gap'):.1f}) only "
    f"at the price of the worst validation performance, i.e. by under-fitting. At the 124M "
    f"WikiText-103 endpoint, by contrast, no method — Baseline included — turns upward within the "
    f"8-epoch budget: every validation curve is still falling when training ends, confirming from "
    f"the training dynamics themselves that this operating point lies outside the overfitting "
    f"regime."
)

add_figure('mechanism_overfitting_delay_large.png',
           f'Figure 6A. Train vs validation perplexity, GPT-2 large ({_MP["large"]:.0f}M) from '
           f'scratch, mean over {_n_ins_l} seeds. Left: full training on a log scale (dashed = '
           f'train, solid = validation). Right: zoom on the late-epoch validation curves: the '
           f'Baseline and WD-tuned show a clear overfitting upturn (Baseline at epoch '
           f'{_iv("large", "Baseline", "best_ep"):.0f}), while τ(w) and Tau(α=0) are still '
           f'improving when the budget ends.')

_dm = lambda m: _iv('medium', m, 'd_med')
_dx = lambda m: _iv('medium', m, 'd_max')
_dml = lambda m: _iv('large', m, 'd_med')
_dxl = lambda m: _iv('large', m, 'd_max')
add_text(
    f"The weight-distribution signature — or rather its absence. Between the first and the last "
    f"epoch at {_MP['medium']:.0f}M, every method lets the median weight magnitude grow: the "
    f"Baseline by {_dm('Baseline'):+.0f}%, WD-tuned by {_dm('WD-tuned'):+.0f}%, Tau(α = 0) by "
    f"{_dm('Tau(alpha=0)'):+.0f}%, L2 by {_dm('L2'):+.0f}%, ElasticNet by "
    f"{_dm('ElasticNet'):+.0f}% and τ(w) — the least — by {_dm('τ(w)'):+.0f}%; the maximum grows "
    f"by {_dx('Baseline'):+.0f}% under the Baseline, {_dx('WD-tuned'):+.0f}% under WD-tuned, "
    f"{_dx('Tau(alpha=0)'):+.0f}% under the ablation and {_dx('τ(w)'):+.0f}% under τ(w) "
    f"(Figure 6B). Under the canonical (ρ, δ) search the tuned τ(w) has δ = {_del_m_tau:.2f} at "
    f"{_MP['medium']:.0f}M and δ = {_del_l_tau:.2f} at {_MP['large']:.0f}M, far above the bulk of "
    f"the distribution (final median |w| ≈ {_iv('medium', 'τ(w)', 'med'):.3f} and "
    f"{_iv('large', 'τ(w)', 'med'):.3f}); the tuned rule therefore decays the bulk at essentially "
    f"the constant rate ρ and is graded only in the extreme tail; its weight dynamics differ "
    f"from the ablation's only in the direction the tuned rates predict — a somewhat stronger "
    f"decay of the bulk, a weaker one of the tail. At {_MP['large']:.0f}M the median "
    f"falls by {abs(_dml('τ(w)')):.0f}% under τ(w) and by {abs(_dml('Tau(alpha=0)')):.0f}% under "
    f"Tau(α = 0) while growing by {_dml('Baseline'):+.0f}% under the Baseline and "
    f"{_dml('WD-tuned'):+.0f}% under WD-tuned; the maximum grows by {_dxl('τ(w)'):+.0f}% under "
    f"τ(w) against {_dxl('Tau(alpha=0)'):+.0f}% under the ablation and {_dxl('Baseline'):+.0f}% "
    f"under the Baseline. Tau(AdamW-scope), whose knee (δ = {_del_l_ad:.2f}) sits closer to the "
    f"bulk and whose nominal ρ = {_rho_l_ad:.1e} is what a schedule-tied decay needs in order to "
    f"bite, is the one method that compresses the median sharply ({_dml('Tau(AdamW-scope)'):+.0f}%) "
    f"while letting the maximum grow ({_dxl('Tau(AdamW-scope)'):+.0f}%). A search that settles "
    f"on a knee close to the bulk would give τ(w) a distinctive signature of its own — a falling "
    f"median with a rising maximum; with the redundant dimension removed and the budget "
    f"equalised (Section 3), the search prefers a knee above the bulk and no such signature "
    f"appears; we "
    f"report the picture as measured."
)

add_figure('mechanism_weight_dynamics_medium.png',
           f'Figure 6B. Weight-magnitude dynamics, GPT-2 medium ({_MP["medium"]:.0f}M) from scratch '
           f'(mean over {_n_ins_m} seeds). Top: median |w| (left) and max |w| (right). Bottom: '
           f'fraction of near-zero weights and total L2 norm. With the tuned knee above the bulk, '
           f'τ(w) differs from the constant-rate ablation only as its tuned rates predict: it '
           f'restrains the growth of the median slightly more and that of the maximum slightly less.')

_r_med_m = _rate(_rho_m_tau, _del_m_tau, _iv('medium', 'τ(w)', 'med')) / _rho_m_tau * 100
_r_max_m = _rate(_rho_m_tau, _del_m_tau, _iv('medium', 'τ(w)', 'max')) / _rho_m_tau * 100
_r_med_l = _rate(_rho_l_tau, _del_l_tau, _iv('large', 'τ(w)', 'med')) / _rho_l_tau * 100
_r_max_l = _rate(_rho_l_tau, _del_l_tau, _iv('large', 'τ(w)', 'max')) / _rho_l_tau * 100
_r_max_ad = _rate(_rho_l_ad, _del_l_ad, _iv('large', 'Tau(AdamW-scope)', 'max')) / _rho_l_ad * 100
add_text(
    f"The tuned operating point. Figure 6C plots the per-step rate ρ/(1 + |w|/δ) implied by the "
    f"tuned {_MP['medium']:.0f}M hyperparameters against the magnitude-independent rates of "
    f"WD-tuned and Tau(α = 0). The median weight decays at {_r_med_m:.0f}% of the small-weight rate "
    f"and the largest weight at {_r_max_m:.0f}% ({_r_med_l:.0f}% and {_r_max_l:.0f}% at "
    f"{_MP['large']:.0f}M): the profile is flat over the bulk and graded only where the weights are "
    f"largest. The comparison with the constant-rate ablation makes the point quantitatively. "
    f"τ(w)'s small-weight rate ρ = {_rho_m_tau:.1e} is {_rho_m_tau / _rho_m_a0:.2f}× the ablation's "
    f"ρ = {_rho_m_a0:.1e} at {_MP['medium']:.0f}M ({_rho_l_tau / _rho_l_a0:.2f}× at "
    f"{_MP['large']:.0f}M), while its rate on the largest weights is "
    f"{_rate(_rho_m_tau, _del_m_tau, _iv('medium', 'τ(w)', 'max')) / _rho_m_a0:.2f}× "
    f"({_rate(_rho_l_tau, _del_l_tau, _iv('large', 'τ(w)', 'max')) / _rho_l_a0:.2f}×): the search "
    f"has bought a slightly stronger decay of the bulk at the price of a weaker decay of the tail, "
    f"and the net effect on perplexity is the fraction of a point of Table 7, in either direction. "
    f"Where adaptivity does substantial work is the AdamW-scope cell of Table 6: there the decay "
    f"also touches LayerNorm gains and biases and follows the schedule, and sparing the largest "
    f"weights (δ = {_del_l_ad:.2f}, which cuts their rate to {_r_max_ad:.0f}% of ρ) recovers "
    f"{abs(_eff_adapt_adamw):.2f} of the {abs(_eff_scope_const):.2f} PPL that the wrong scope "
    f"costs. That is the mechanism the factorial identifies: adaptivity as a compensator for "
    f"scope, not as an independent source of generalization."
)

add_figure('mechanism_adaptivity_medium.png',
           f'Figure 6C. The decay profile implied by the tuned hyperparameters (GPT-2 medium, '
           f'{_MP["medium"]:.0f}M). Left: per-step decay rate ρ/(1 + |w|/δ) vs weight magnitude, '
           f'against the magnitude-independent rates of WD-tuned and Tau(α=0); vertical lines mark '
           f'the median and maximum of the trained τ(w) weight distribution, both far below and '
           f'just above the knee δ = {_del_m_tau:.2f} respectively. Right: the implicit Fair '
           f'penalty — quadratic near the origin, linear in the tail — against L2 and L1.')

doc.add_heading("6.3 Why from-scratch autoregressive LMs, and not transformers as such?", level=2)

add_text(
    "The affinity of the τ family for from-scratch autoregressive language modelling warrants "
    "discussion, because the obvious architectural reading is the wrong one. Three factors "
    "account for it:"
)

p = doc.add_paragraph()
run = p.add_run("What is decayed. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"The factorial of Table 6 locates the largest single effect in the scope of the decay. "
    f"AdamW's decoupled decay is applied to every parameter handed to the optimizer — biases "
    f"included — and is multiplied by the current learning rate, so it fades with the schedule; "
    f"the τ family leaves the biases alone and decays the weight tensors at a constant per-step rate. "
    f"At 66M the difference is worth {abs(_eff_scope_const):.2f} PPL at a constant profile. The "
    f"structural fact that matters is therefore not the heterogeneity of weight magnitudes across "
    f"attention and feed-forward matrices, which magnitude adaptivity was meant to exploit, but "
    f"the presence of parameters that should not be decayed at all and of a schedule that "
    f"removes the decay precisely when an over-capacity model starts to overfit."
)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("Interaction with Adam. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    "Transformer training typically uses Adam or AdamW, which maintain per-parameter adaptive "
    "learning rates. The post-update application of τ-decay avoids interfering with Adam's "
    "momentum and variance estimates, allowing both mechanisms to operate in their respective "
    "domains — gradient adaptation (Adam) and decoupled regularization (τ-decay). By contrast, "
    "loss-based penalties (L1, L2, SCAD, MCP, LSP) modify the gradient signal directly and are "
    "rescaled by the adaptive step sizes, which is consistent with their poor showing on the "
    "from-scratch language models: at 7M none of them separates from the unregularized Baseline "
    "and two of them degrade it, and at 66M the tuning itself switches them off — the PSO "
    f"drives L2 and ElasticNet to (near) zero effect while the decoupled decays gain "
    f"{_d_base - _d_wd:.0f}–{_d_base - _d_tau:.0f} PPL (Section 5.6)."
)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("Overfitting pressure. ")
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(
    f"The regime evidence of Sections 5.6 and 6.2 locates the advantage where the unregularized "
    f"model's validation optimum precedes the end of training: {_MP['small']:.1f}M on 30 epochs "
    f"and {_MP['large']:.0f}M on 12 epochs of WikiText-2, but not {_MP['medium']:.0f}M on 12 "
    f"epochs, not {_MP['tiny']:.1f}M, and not 124M on the hundred-times-larger WikiText-103. A "
    f"language model trained from scratch on a small corpus is simply the setting in which that "
    f"pressure is most acute; a vision transformer on CIFAR-10 with augmentation, a CNN, or a "
    f"pretrained LLM fine-tuned for three epochs is not. In those settings the tuning search "
    f"drives the decay towards zero and the τ family degrades to parity with tuned decoupled "
    f"weight decay rather than into harm. The affinity is with over-capacity from-scratch "
    f"autoregressive-LM training, and that is how we state it throughout; establishing an "
    f"architectural effect proper would require the controlled design described in Section 5.8."
)
run.font.name = 'Times New Roman'

doc.add_heading("6.4 Hyperparameter sensitivity and the modern-LLM null effect", level=2)

_sm_fam_spread = max(abs(_sm_tau - _sm_base), abs(_sm_wd - _sm_base), abs(_sm_ab - _sm_base))
_rho_tuned = {s: _hpv(_best_hp(k), 'τ(w)', 'rho') for s, k in _SCALE_KEYS.items()}
_rho_lo = min(v for v in (_rho_tuned['small'], _rho_tuned['medium'], _rho_tuned['large']) if v == v)
_rho_hi = max(v for v in (_rho_tuned['small'], _rho_tuned['medium'], _rho_tuned['large']) if v == v)
_l1_lam = _hpv(_best_hp('gpt2_wikitext'), 'L1', 'lambda_val')
add_text(
    f"A genuine null-effect regime does exist: modern small LLMs already saturated by "
    f"pretraining, represented here by SmolLM2-135M, where the decay-family configurations "
    f"finish within {_sm_fam_spread:.2f} PPL of the unregularized Baseline (Section 5.5). "
    f"When pretraining has already exhausted the redundancy that regularization could exploit, "
    f"the tuning search drives the decay towards zero (ρ = {_rho_tuned['smollm2']:.1e} on "
    f"SmolLM2 and ρ = {_rho_tuned['tiny']:.1e} at 2M, against {_rho_lo:.1e}–{_rho_hi:.1e} on the "
    f"three over-capacity scales) and τ-decay degrades gracefully to the Baseline rather than "
    f"harming performance. Within the regime, the winning settings lie in a narrow band of small "
    f"ρ — a factor of {_rho_hi / _rho_lo:.0f} across a tenfold range of model size — and the knee "
    f"δ settles above the bulk of the weight distribution (Section 6.2), so the method is not "
    f"delicately tuned. The tuning caveat cuts the other way for the loss-side penalties: the "
    f"one-dimensional searches (12 evaluations on a log-scale λ) are saturated for L2 and for "
    f"the decoupled decays, whose optima are flat, but not for L1 on the {_MP['small']:.1f}M "
    f"GPT-2, where the search settled on λ = {_l1_lam:.1e} and the resulting model finishes "
    f"{_gpt_pen_pct('L1'):+.0f}% above the Baseline. That number is an artefact of the search, "
    f"not a property of L1, and we do not count it as evidence for τ-decay."
)

doc.add_heading("6.5 Comparison with non-convex penalties", level=2)

add_text(
    f"The proposed method occupies an intermediate position between convex and non-convex "
    f"penalties. Like SCAD and MCP, it provides magnitude-dependent shrinkage that reduces bias "
    f"on large weights. Unlike them, it achieves this through a smooth, convex implicit penalty "
    f"rather than piecewise definitions. The practical consequence is visible on the "
    f"from-scratch language models: SCAD and MCP "
    f"{'never separate from the unregularized Baseline' if not _ncx_sep else 'separate from the unregularized Baseline only for ' + _and_join(_ncx_sep)}, "
    f"and LSP — the most aggressive non-convex penalty — collapses to chance on BERT-tiny "
    f"({_lsp_bert:.1f}%) and degrades the {_MP['small']:.1f}M GPT-2 to {_lsp_gpt:.1f} PPL against "
    f"{_gpt_base:.1f} for the Baseline"
    + (f" (on ViT-CIFAR it reaches {_lsp_vit:.2f}%)" if _lsp_vit == _lsp_vit else "")
    + ". The τ family avoids these failure modes entirely."
)

add_text(
    "We would caution against reading that contrast as the measure of the contribution. SCAD, MCP "
    "and LSP are the conventional comparison set for a magnitude-dependent shrinkage, but they are "
    "not the closest one: they are non-convex, flatten the penalty rather than cap the shrinkage "
    "force, and are motivated by variable selection rather than by training dynamics. Outperforming "
    "them establishes that τ-decay is a better-behaved regularizer than a family known to be "
    "ill-suited to deep networks; it does not establish anything about the Fair profile "
    "specifically. The comparison that does is Section 5.10, against the other decoupled decays "
    "derived from robust losses, where the shrinkage profile is the only free variable — and "
    "there the profiles are interchangeable."
)

doc.add_heading("6.6 Sparsity patterns", level=2)

_sp_tau = lambda b: get_val(all_data[b], 'τ(w)', 'sparsity', 'mean')
_sp_classic = [_sp_tau(b) for b in ('sin_regression', 'complex_regression', 'mnist', 'cifar')]
_sp_classic = [v for v in _sp_classic if v == v]
_sp_l1 = lambda b: get_val(all_data[b], 'L1', 'sparsity', 'mean')
add_text(
    f"τ-decay produces little sparsity at the |w| < 10⁻³ threshold: "
    f"{min(_sp_classic):.1f}–{max(_sp_classic):.1f}% on the classic benchmarks, "
    f"{_sp_tau('smollm2_wikitext'):.1f}% on the pretrained SmolLM2, and "
    f"{_sp_tau('gpt2_wikitext'):.1f}%, {_sp_tau('gpt2_medium_wikitext'):.1f}% and "
    f"{_sp_tau('gpt2_large_wikitext'):.1f}% on the {_MP['small']:.1f}M, {_MP['medium']:.0f}M and "
    f"{_MP['large']:.0f}M from-scratch GPT-2 models — far less than L1 "
    f"({_sp_l1('gpt2_wikitext'):.0f}% on the {_MP['small']:.1f}M GPT-2). This is what the "
    f"analysis of Section 6.2 predicts: with the tuned knee above the bulk, the decay "
    f"concentrates the distribution near zero only slightly more than a constant decay does, "
    f"and, as noted in Section 2, a multiplicative shrinkage that vanishes as w → 0 does not "
    f"drive weights to exact zeros in finite time. Sparsity is therefore a by-product of the "
    f"method, not one of its aims, and no compression claim is made in this paper."
)

doc.add_heading("6.7 Limitations", level=2)

add_text(
    f"On traditional feedforward and CNN architectures, τ-decay does not outperform standard "
    f"methods: it ranks mid-field on sin(x) ({_sin['rank']}/{_sin['total']}) and CIFAR-10 "
    f"({_cif['rank']}/{_cif['total']}), and on CIFAR-10 it is significantly behind "
    f"{_cif_best_name} after family-wide correction (Section 5.7). Practitioners working "
    f"exclusively with small, traditional architectures should not adopt τ-decay over "
    f"established methods."
)

add_text(
    f"The magnitude-adaptive profile — the feature that gives the method its name — is a "
    f"second-order and scale-dependent refinement. Against the implementation-matched "
    f"constant-decay ablation it is worth {_AG['large'][2]:+.2f} PPL at {_MP['large']:.0f}M, "
    f"{_AG['medium'][2]:+.2f} at {_MP['medium']:.0f}M and {_AG['small'][2]:+.2f} at "
    f"{_MP['small']:.1f}M, and the factorial shows that its larger apparent value at AdamW's "
    f"scope ({_eff_adapt_adamw:+.2f} PPL) is compensation for that scope. The robust-decay "
    f"head-to-head points the same way: Huber, pseudo-Huber, log-cosh and Fair profiles, and the "
    f"constant-rate ablation, all fall within a few tenths of a point of one another at 66M, so "
    f"the exact saturation profile appears to be immaterial. A reader who takes from this paper "
    f"only the recommendation to apply a decoupled, weight-only, schedule-independent decay "
    f"when training an over-capacity language model from scratch has taken the supported part "
    f"of the contribution."
)

if os.path.exists(_ROBUST_CSV):
    add_text(
        f"The robust-decay head-to-head of Section 5.10 is run at a single operating point — the "
        f"66M from-scratch configuration where the τ-family advantage is largest — with "
        f"n = {_ROBUST_N} seeds for the three competitors, and with those competitors tuned on a "
        f"24-evaluation rather than a 40-evaluation budget. It is therefore sufficient to "
        f"establish that the Fair saturation profile is not distinguishable from the Huber, "
        f"pseudo-Huber and log-cosh profiles where the mechanism is most active — a conclusion "
        f"the smaller budget can only have made harder to reach — but not to map how the "
        f"ordering behaves across scales or corpora. We also compare against the Huber shrinkage "
        f"as a decay rule at our own call site, not against the full AdamHD optimizer of Guo & "
        f"Fan (2025) as published; the comparison isolates the penalty profile at the cost of not "
        f"reproducing their complete training configuration."
    )

add_text(
    f"Sample sizes are the sharpest limitation of this study. Most benchmarks use n = 5 and the "
    f"most expensive ones n = 3, which is too few to estimate a variance stably; the primary and "
    f"secondary contrasts and the factorial were run at n = {N_SEEDS['gpt2_large_wikitext']} at "
    f"66M for exactly this reason, but the rest of the suite was not. Several comparisons on "
    f"BERT-tiny and ViT-CIFAR show sizeable point differences yet do not reach α = 0.05. We "
    f"deliberately do not describe these as cases a larger sample would resolve in τ-decay's "
    f"favour: a larger sample could confirm the effect, return a null, or reverse the sign of the "
    f"estimate, and at n = 3–5 the point estimates are not precise enough to prefer one of those "
    f"outcomes. Effect sizes at these sample sizes should also be read with care even after the "
    f"Hedges correction, because a small denominator inflates them; the very large values "
    f"reported on the from-scratch GPT-2 benchmarks reflect unusually tight seed-to-seed variance "
    f"as much as they reflect the size of the gap."
)

add_text(
    "The scale sweep varies capacity at fixed data, so on its own it could not separate model "
    "size from the capacity-to-data ratio; the data-quantity arm of Section 5.6.1 manipulates "
    "the data directly and sharpens the reading rather than completing it. Its outcome — at "
    "fixed hyperparameters the margin grows with the amount of data — shows that overfitting "
    "pressure marks where the gain appears without scaling it, and it leaves one question "
    "open: because the arm transfers hyperparameters from the full-corpus optimum, "
    "under-regularization at higher pressure cannot yet be separated from a genuine data "
    "dependence, and a per-fraction re-tuning remains to be run."
)

add_text(
    "The 124M WikiText-103 confirmation carries protocol caveats of its own. Compute "
    "constraints (~7 GPU-hours per training run) ruled out full per-method PSO; "
    "hyperparameters were instead transferred from the 66M sweep optimum with "
    "step-count rescaling and validated by a ¼×/1×/4× confirmation sweep, so each "
    "method competes at a well-chosen but not exhaustively tuned operating point — a "
    "condition applied symmetrically to all six methods. The 8-epoch budget "
    "(~0.9B tokens seen) also leaves every method's validation perplexity still "
    "improving at the end of training; conclusions at this scale therefore concern "
    "the compute-bound regime typical of academic budgets, and we make no claim about "
    "asymptotic behaviour under substantially longer training."
)

# ============================================================================
# 7. CONCLUSION
# ============================================================================

doc.add_heading("7. Conclusion", level=1)

add_text(
    "τ-decay is a smooth, decoupled weight-decay update: the gradient of the Fair penalty of "
    "robust statistics, applied after the optimizer step to the weight tensors only (biases excluded), at a rate "
    "independent of the learning-rate schedule, with a relative shrinkage that decreases with "
    "parameter magnitude. It is a member of the family of decoupled decays derived from robust "
    "saturating losses rather than a new class of regularizer, and in the two-parameter form "
    "w ← w − ρ·w/(1 + |w|/δ) it costs one line after optimizer.step, adds constant per-step "
    "overhead, and works with SGD and Adam alike; its δ → ∞ limit is a constant decoupled decay "
    "with the same scope and schedule."
)

add_text(
    f"Its gains are confined to one regime and are large there: autoregressive language models "
    f"trained from scratch under overfitting pressure. On WikiText-2 the τ family — the constant "
    f"decay and τ(w) alike — beats PSO-tuned AdamW weight decay by {_h_gap7:.1f} PPL at "
    f"{_MP['small']:.1f}M parameters and {_h_gap66:.1f} PPL at {_MP['large']:.0f}M "
    f"({abs(_pct('gpt2_wikitext')):.0f}% and {abs(_pct('gpt2_large_wikitext')):.0f}% below the "
    f"unregularized Baseline), by {_h_gap18:.1f} PPL at {_MP['medium']:.0f}M and not at all at "
    f"{_MP['tiny']:.1f}M. What orders these cases is not model size but whether the unregularized "
    f"model overfits inside its budget — a condition that marks where the gain appears rather "
    f"than scaling it: at fixed capacity and hyperparameters the margin grows with the amount "
    f"of data (Section 5.6.1). We are explicit about where the margin comes from, because "
    f"it is not what the method's name suggests: the scope × adaptivity factorial at 66M attributes "
    f"{abs(_eff_scope_const):.2f} PPL of it to the decay's scope and schedule — weight tensors only, "
    f"biases excluded, at a rate independent of the learning-rate schedule — and "
    f"{abs(_eff_adapt_tau):.2f} PPL to magnitude adaptivity at that scope; tested paired by seed, "
    f"the adaptivity increment is {_AG['large'][2]:+.2f} PPL at {_MP['large']:.0f}M, "
    f"{_AG['medium'][2]:+.2f} at {_MP['medium']:.0f}M and {_AG['small'][2]:+.2f} at "
    f"{_MP['small']:.1f}M, where the constant decay wins. Adaptivity mainly compensates a wrong "
    f"decay scope, and the Huber, pseudo-Huber and log-cosh profiles land at parity with the Fair "
    f"one. Practitioners can take the scope-and-schedule lesson independently of τ(w) itself."
)

add_text(
    f"The benefit is regime-dependent, and the regime is narrower than the architecture. It is "
    f"absent at {_MP['tiny']:.1f}M and at {_MP['medium']:.0f}M on 12 epochs, where nothing overfits, "
    f"and it compresses to parity at the 124M/WikiText-103 operating point, where the corpus "
    f"rather than capacity binds. On pretrained models and on the non-autoregressive "
    f"architectures we tested, τ-decay is equivalent to tuned decoupled weight decay within "
    f"pre-specified margins wherever the data can show it, and it loses after family-wide "
    f"correction on {_sl} classic benchmark{'s' if _sl != 1 else ''}"
    + (f" ({_ab_loss_txt})" if _ab_losses else "")
    + ". That is a no-harm result with named exceptions, not a win, and we do not present it as "
    "one. A vision transformer is a transformer and τ-decay does not help there; the affinity is "
    "with over-capacity autoregressive language-model training, not with attention as such."
)

add_text(
    "The honest summary is narrower than the one this work began with, and better supported. "
    "A decoupled, weight-only, schedule-independent decay is a cheap and well-behaved regularizer "
    "whose benefit is concentrated in a specific and identifiable training regime; the Fair "
    "profile is one acceptable way to implement it, and its magnitude adaptivity is a "
    "second-order refinement. The open questions we consider most worth pursuing are the "
    "per-fraction re-tuning of the data-quantity arm, which would separate "
    "under-regularization at higher pressure from a genuine data dependence; a "
    "controlled architectural comparison at matched capacity and token budget; and larger-n "
    "replication of the adaptivity contrast at 7M and 18M, where its sign is the question."
)

# ============================================================================
# DECLARATIONS
# ============================================================================

# ============================================================================
# APPENDIX A — reproducibility (REVIEWER: "completezza metodologica")
# Generated by introspecting the code that actually runs, so it cannot drift from
# the implementation the way a hand-maintained table would.
# ============================================================================
doc.add_page_break()
doc.add_heading("Appendix A. Reproducibility", level=1)

add_text(
    "Every value below is read out of the source at build time rather than transcribed, so "
    "this appendix cannot fall out of step with the code. The repository, including the "
    "exact scripts and the raw per-seed results, is linked under Data Availability."
)

add_bold_paragraph("A.1 Search spaces and hyperparameter tuning", size=12, space_before=8)

add_text(
    "Hyperparameters are tuned by Particle Swarm Optimization (Kennedy & Eberhart, 1995) on "
    "the validation split, independently per method and per benchmark. Each PSO evaluation is "
    "a complete training run. The number of evaluations is set by the dimensionality of the "
    "method's search space — the same budget for every one-parameter method and the same, "
    "larger budget for every method with two or more parameters — so that no method, τ(w) "
    "included, receives more tuning effort than a competitor of equal dimensionality. The "
    "budgets below are read from the code that ran (dimension_aware_budget)."
)

from experiment_utils import (SEARCH_SPACES as _SS, dimension_aware_budget as _dab,
                              PSO_BUDGETS as _PB)
_ss_tab = doc.add_table(rows=1, cols=4)
_ss_tab.style = 'Table Grid'
_ss_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['Method', 'Parameter', 'Range', 'PSO evals (auto)']):
    _c = _ss_tab.rows[0].cells[_i]; _c.text = ""
    _r = _c.paragraphs[0].add_run(_h); _r.bold = True
    _r.font.size = Pt(9); _r.font.name = 'Times New Roman'
    _c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_c, 'D9D9D9')
for _m in ['Baseline', 'L1', 'L2', 'ElasticNet', 'SCAD', 'MCP', 'LSP', 'WD-tuned',
           'Tau(alpha=0)', 'Tau(AdamW-scope)', 'τ(w)',
           'Huber-decay', 'PseudoHuber-decay', 'LogCosh-decay']:
    _sp = _SS.get(_m, {})
    _b = _dab(len(_sp))
    _ev = _b['n_particles'] * _b['n_iterations'] if _sp else 0
    if not _sp:
        add_table_row(_ss_tab, [_m, '— (no tuning)', '—', '0'])
        continue
    for _j, (_k, _v) in enumerate(_sp.items()):
        _rng = (f"[{_v['low']:g}, {_v['high']:g}]"
                + (" log" if _v.get('log_scale') else " linear"))
        add_table_row(_ss_tab, [_m if _j == 0 else '', _k, _rng,
                                str(_ev) if _j == 0 else ''])
_p = doc.add_paragraph()
_r = _p.add_run(
    f"Table A1. Tuned hyperparameters, their search ranges and the PSO budget under the "
    f"dimension-aware policy ({_dab(1)['n_particles']}×{_dab(1)['n_iterations']} evaluations "
    f"for 1-D, {_dab(2)['n_particles']}×{_dab(2)['n_iterations']} for 2-D, "
    f"{_dab(3)['n_particles']}×{_dab(3)['n_iterations']} for 3-D). Benchmarks predating this "
    f"policy used the fixed 'standard' budget of "
    f"{_PB['standard']['n_particles']}×{_PB['standard']['n_iterations']} evaluations. The three "
    f"robust decays of Section 5.10 were tuned under the earlier 24-evaluation two-dimensional "
    f"budget (Section 5.10); the table lists the current policy.")
_r.font.size = Pt(9); _r.font.name = 'Times New Roman'; _r.italic = True

_wd_sweep_wd = {s: _hpv(_best_hp(_SCALE_KEYS[s]), 'WD-tuned', 'wd')
                for s in ('small', 'medium', 'large')}
_wd_single_vals = [_hpv(_best_hp(_bn), 'WD-tuned', 'wd')
                   for _bn, _m_, _mo_, _sc_, _u_ in BENCHMARKS if _bn != 'gpt2_wikitext']
_wd_single_vals = [v for v in _wd_single_vals if v == v]
_wd_single_max = max(_wd_single_vals) if _wd_single_vals else float('nan')
add_text(
    f"One range in Table A1 deserves a note. WD-tuned's range [1e-6, 10] is the extended "
    f"range adopted for the GPT-2 scale sweep after the original upper bound of 0.1 proved "
    f"binding there: with AdamW's per-step decay equal to lr × wd under a linear-to-zero "
    f"schedule, wd ≤ 0.1 caps the decay below the per-step rates the τ family selects, and "
    f"the re-tuned winners lie above the old bound (wd = {_wd_sweep_wd['small']:.3f} at "
    f"{_MP['small']:.1f}M, {_wd_sweep_wd['medium']:.3f} at {_MP['medium']:.0f}M and "
    f"{_wd_sweep_wd['large']:.3f} at {_MP['large']:.0f}M). For the single-scale benchmarks "
    f"other than the GPT-2 sweep the search used the 0.1 upper bound, which was never "
    f"binding there (winning values ≤ {_wd_single_max:.3f})."
)

add_bold_paragraph("A.2 What was NOT tuned", size=12, space_before=8)

add_text(
    "The learning rate is held fixed for every method on a given benchmark and is never part "
    "of any search space; only the regularization hyperparameters are tuned. This is applied "
    "symmetrically, but it is a real limitation rather than a neutral choice: regularization "
    "strength and step size interact, the optimal learning rate need not be the same for a "
    "strongly and a weakly regularized model, and a fixed value may therefore favour whichever "
    "method happens to suit it. Tuning the learning rate jointly would add a dimension to "
    "every method's search and multiply the cost of the whole study; we did not do it, and we "
    "flag it as the most likely residual confound in the comparisons. Also fixed across "
    "methods within a benchmark: optimizer (AdamW) and its β and ε, the warmup and decay "
    "schedule, gradient clipping at 1.0, dropout, batch size, sequence length, maximum epochs "
    "and the early-stopping criterion and patience."
)

add_bold_paragraph("A.3 Per-benchmark configuration", size=12, space_before=8)

_cfg_tab = doc.add_table(rows=1, cols=5)
_cfg_tab.style = 'Table Grid'
_cfg_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
for _i, _h in enumerate(['Benchmark', 'Metric', 'Seeds (n)', 'Roster size', 'WD-tuned / Tau(α=0)']):
    _c = _cfg_tab.rows[0].cells[_i]; _c.text = ""
    _r = _c.paragraphs[0].add_run(_h); _r.bold = True
    _r.font.size = Pt(9); _r.font.name = 'Times New Roman'
    _c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(_c, 'D9D9D9')
for _bn, _metric, _mode, _sc, _unit in BENCHMARKS + SCALE_BENCHMARKS:
    _ms = set(all_data[_bn]['method'])
    add_table_row(_cfg_tab, [
        _bn, f"{_metric} ({'lower' if _mode == 'min' else 'higher'} is better)",
        str(N_SEEDS.get(_bn, 5)), str(len(_ms)),
        'yes' if 'WD-tuned' in _ms else 'no'])
_p = doc.add_paragraph()
_r = _p.add_run(
    "Table A2. Per-benchmark evaluation configuration as recorded in the result files. The "
    "last column makes explicit which benchmarks carry the decoupled-weight-decay competitor "
    "and the α = 0 ablation, which belong to the protocol; Table 1 shows "
    "a dash wherever a method is absent from a benchmark's roster.")
_r.font.size = Pt(9); _r.font.name = 'Times New Roman'; _r.italic = True

add_bold_paragraph("A.4 Software, hardware and determinism", size=12, space_before=8)

add_text(
    "Training runs on a single NVIDIA GPU per benchmark (RTX 3090 24 GB for the from-scratch "
    "scale sweep and the runs reported here; earlier benchmarks used A4500/A5000-class cards), "
    "with PyTorch and transformers pinned in the repository's requirements file. Every run "
    "calls a single seeding routine that sets the Python, NumPy and CUDA generators and "
    "enables deterministic cuDNN kernels, so a given (method, hyperparameters, seed) triple "
    "reproduces exactly; this was verified by re-running the 18M configuration and obtaining "
    "identical numbers. Long runs are journaled per completed training run, so an interrupted "
    "benchmark resumes without repeating finished work and without changing its results."
)

doc.add_heading("Author Contributions", level=2)
add_text("G. Maulucci: conceptualization, methodology, software, investigation, formal analysis, "
         "writing \u2013 original draft, corresponding author. T. Marchetti and M. De Spirito: "
         "writing \u2013 review and editing.")

doc.add_heading("Declaration of Competing Interest", level=2)
add_text("The authors declare no competing financial or personal interests.")

doc.add_heading("Data Availability", level=2)
add_text("Code and data are available at https://github.com/doctormaulux/Regularization_Decay.")

doc.add_heading("Acknowledgments", level=2)
add_text("The authors acknowledge the computational resources provided by RunPod and the "
         "support of Università Cattolica del Sacro Cuore.")

# ============================================================================
# REFERENCES
# ============================================================================

doc.add_heading("References", level=1)

# Elsevier APA-style references, alphabetical by first author, with DOIs / URLs.
refs = [
    "Allal, L. B., Lozhkov, A., Bakouch, E., Bl\u00e1zquez, G. M., Penedo, G., Tunstall, L., Marafioti, A., Kydl\u00ed\u010dek, H., et al. (2025). SmolLM2: When smol goes big \u2014 data-centric training of a small language model. arXiv preprint arXiv:2502.02737. https://doi.org/10.48550/arXiv.2502.02737",
    "Cand\u00e8s, E. J., Wakin, M. B., & Boyd, S. P. (2008). Enhancing sparsity by reweighted l1 minimization. Journal of Fourier Analysis and Applications, 14(5), 877\u2013905. https://doi.org/10.1007/s00041-008-9045-x",
    "Charbonnier, P., Blanc-Féraud, L., Aubert, G., & Barlaud, M. (1994). Two deterministic half-quadratic regularization algorithms for computed imaging. In Proceedings of the 1st IEEE International Conference on Image Processing (ICIP) (Vol. 2, pp. 168–172). IEEE. https://doi.org/10.1109/ICIP.1994.413553",
    "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. In Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics: Human Language Technologies (NAACL-HLT), Volume 1 (pp. 4171\u20134186). Association for Computational Linguistics. https://doi.org/10.18653/v1/N19-1423",
    "Dosovitskiy, A., Beyer, L., Kolesnikov, A., Weissenborn, D., Zhai, X., Unterthiner, T., et al. (2021). An image is worth 16x16 words: Transformers for image recognition at scale. In International Conference on Learning Representations (ICLR). arXiv:2010.11929. https://arxiv.org/abs/2010.11929",
    "Fair, R. C. (1974). On the robust estimation of econometric models. Annals of Economic and Social Measurement, 3(4), 667\u2013677.",
    "Fan, J., & Li, R. (2001). Variable selection via nonconcave penalized likelihood and its oracle properties. Journal of the American Statistical Association, 96(456), 1348\u20131360. https://doi.org/10.1198/016214501753382273",
    "Friedman, J. H. (1991). Multivariate adaptive regression splines. The Annals of Statistics, 19(1), 1\u201367. https://doi.org/10.1214/aos/1176347963",
    "Ghiasi, M. A., Shafahi, A., & Ardekani, R. (2023). Improving robustness with adaptive weight decay. In Advances in Neural Information Processing Systems 36 (NeurIPS). https://proceedings.neurips.cc/paper_files/paper/2023/hash/f9d7d6c695bc983fcfb5b70a5fbdfd2f-Abstract-Conference.html",
    "Guo, F.-M., & Fan, Y. (2025). AdamHD: Decoupled Huber decay regularization for language model pre-training. arXiv preprint arXiv:2511.14721. https://doi.org/10.48550/arXiv.2511.14721",
    "Han, S., Pool, J., Tran, J., & Dally, W. J. (2015). Learning both weights and connections for efficient neural networks. In Advances in Neural Information Processing Systems 28 (NeurIPS) (pp. 1135\u20131143). https://papers.nips.cc/paper/2015/hash/ae0eb3eed39d2bcef4622b2499a05fe6-Abstract.html",
    "Hinton, G., Vinyals, O., & Dean, J. (2015). Distilling the knowledge in a neural network. arXiv preprint arXiv:1503.02531. (Presented at NIPS 2014 Deep Learning Workshop.) https://arxiv.org/abs/1503.02531",
    "Holland, P. W., & Welsch, R. E. (1977). Robust regression using iteratively reweighted least-squares. Communications in Statistics \u2014 Theory and Methods, 6(9), 813\u2013827. https://doi.org/10.1080/03610927708827533",
    "Huber, P. J. (1964). Robust estimation of a location parameter. The Annals of Mathematical Statistics, 35(1), 73–101. https://doi.org/10.1214/aoms/1177703732",
    "Kennedy, J., & Eberhart, R. (1995). Particle swarm optimization. In Proceedings of ICNN'95 \u2014 International Conference on Neural Networks (Vol. 4, pp. 1942\u20131948). IEEE. https://doi.org/10.1109/ICNN.1995.488968",
    "Kingma, D. P., & Ba, J. (2015). Adam: A method for stochastic optimization. In 3rd International Conference on Learning Representations (ICLR). arXiv:1412.6980. https://arxiv.org/abs/1412.6980",
    "Krizhevsky, A. (2009). Learning multiple layers of features from tiny images (Technical Report TR-2009). Department of Computer Science, University of Toronto. https://www.cs.toronto.edu/~kriz/learning-features-2009-TR.pdf",
    "LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-based learning applied to document recognition. Proceedings of the IEEE, 86(11), 2278\u20132324. https://doi.org/10.1109/5.726791",
    "Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization. In 7th International Conference on Learning Representations (ICLR). arXiv:1711.05101. https://arxiv.org/abs/1711.05101",
    "Merity, S., Xiong, C., Bradbury, J., & Socher, R. (2017). Pointer sentinel mixture models. In International Conference on Learning Representations (ICLR). arXiv:1609.07843. https://arxiv.org/abs/1609.07843",
    "Michel, P., Levy, O., & Neubig, G. (2019). Are sixteen heads really better than one? In Advances in Neural Information Processing Systems 32 (NeurIPS) (pp. 14014\u201314024). https://proceedings.neurips.cc/paper/2019/hash/2c601ad9d2ff9bc8b282670cdd54f69f-Abstract.html",
    "Nakamura, K., & Hong, B.-W. (2019). Adaptive weight decay for deep neural networks. arXiv preprint arXiv:1907.08931. https://doi.org/10.48550/arXiv.1907.08931",
    "Radford, A., Wu, J., Child, R., Luan, D., Amodei, D., & Sutskever, I. (2019). Language models are unsupervised multitask learners. OpenAI Technical Report. https://cdn.openai.com/better-language-models/language_models_are_unsupervised_multitask_learners.pdf",
    "Socher, R., Perelygin, A., Wu, J., Chuang, J., Manning, C. D., Ng, A. Y., & Potts, C. (2013). Recursive deep models for semantic compositionality over a sentiment treebank. In Proceedings of the 2013 Conference on Empirical Methods in Natural Language Processing (EMNLP) (pp. 1631\u20131642). Association for Computational Linguistics. https://aclanthology.org/D13-1170/",
    "Srivastava, N., Hinton, G., Krizhevsky, A., Sutskever, I., & Salakhutdinov, R. (2014). Dropout: A simple way to prevent neural networks from overfitting. Journal of Machine Learning Research, 15(1), 1929\u20131958. https://jmlr.org/papers/v15/srivastava14a.html",
    "Tibshirani, R. (1996). Regression shrinkage and selection via the lasso. Journal of the Royal Statistical Society: Series B (Methodological), 58(1), 267\u2013288. https://doi.org/10.1111/j.2517-6161.1996.tb02080.x",
    "Tikhonov, A. N. (1963). Solution of incorrectly formulated problems and the regularization method. Soviet Mathematics Doklady, 4, 1035\u20131038.",
    "Wang, A., Singh, A., Michael, J., Hill, F., Levy, O., & Bowman, S. R. (2018). GLUE: A multi-task benchmark and analysis platform for natural language understanding. In Proceedings of the 2018 EMNLP Workshop BlackboxNLP: Analyzing and Interpreting Neural Networks for NLP (pp. 353\u2013355). Association for Computational Linguistics. https://doi.org/10.18653/v1/W18-5446",
    "You, Y., Gitman, I., & Ginsburg, B. (2017). Large batch training of convolutional networks. arXiv preprint arXiv:1708.03888. https://arxiv.org/abs/1708.03888",
    "You, Y., Li, J., Reddi, S., Hseu, J., Kumar, S., Bhojanapalli, S., Song, X., Demmel, J., Keutzer, K., & Hsieh, C.-J. (2020). Large batch optimization for deep learning: Training BERT in 76 minutes. In International Conference on Learning Representations (ICLR). arXiv:1904.00962. https://arxiv.org/abs/1904.00962",
    "Zhang, C.-H. (2010). Nearly unbiased variable selection under minimax concave penalty. The Annals of Statistics, 38(2), 894\u2013942. https://doi.org/10.1214/09-AOS729",
    "Zou, H., & Hastie, T. (2005). Regularization and variable selection via the elastic net. Journal of the Royal Statistical Society: Series B (Statistical Methodology), 67(2), 301\u2013320. https://doi.org/10.1111/j.1467-9868.2005.00503.x",
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.left_indent = Pt(18)

# ── Typographic minus ───────────────────────────────────────────────────────
# Python's format signs are ASCII hyphens; the manuscript uses U+2212 for negative numbers.
# Applied to every run (body and tables) at the end, so no f-string has to remember it.
import re as _re
_NEG = _re.compile(r'(?<![^\s(=\[,])-(?=\d)')
def _fix_minus(container):
    for _par in container.paragraphs:
        for _run in _par.runs:
            if '-' in _run.text:
                _run.text = _NEG.sub('\u2212', _run.text)
_fix_minus(doc)
for _tbl in doc.tables:
    for _row in _tbl.rows:
        for _cell in _row.cells:
            _fix_minus(_cell)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "articolo.docx")
doc.save(output_path)
print(f"[OK] Saved: {output_path}")
