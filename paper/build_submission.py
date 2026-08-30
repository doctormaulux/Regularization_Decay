"""Assemble the Neural Networks (Elsevier) submission folder.

Copies the built manuscript (articolo.docx) and cover letter, and generates the
separate files Editorial Manager asks for: title page, highlights (3-5 bullets,
<= 85 characters each, enforced), CRediT author statement and declaration of
interests. Numbers in the highlights are read from the result CSVs, never
hard-coded. Output: <project root>/submission/.

Run AFTER build_paper.py (articolo.docx must be current):
    python paper/build_paper.py && python paper/build_submission.py
"""
import os
import shutil
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(SCRIPT_DIR)
FIGURES_DIR = os.path.join(SCRIPT_DIR, 'figures')
OUT_DIR = os.path.join(ROOT_DIR, 'submission')

sys.path.insert(0, ROOT_DIR)
from analysis.stats_utils import load_csv, get_val

# Shortened 2026-08-31 for the journal's ~80-character title limit (must match
# build_paper.py's title block).
TITLE = ("Decoupled Fair Weight Decay for Over-Capacity Autoregressive Language Models")
assert len(TITLE) <= 80, f"Title over 80 characters ({len(TITLE)})"
RUNNING_TITLE = "Decoupled Fair weight decay"
AUTHORS_LINE = ("Giuseppe Maulucci¹* (ORCID: 0000-0002-2154-319X), "
                "Tommaso Marchetti¹, Marco De Spirito¹")
AFFILIATION = "¹ Università Cattolica del Sacro Cuore, Rome, Italy"
CORRESPONDING = ("* Giuseppe Maulucci — giuseppe.maulucci@unicatt.it — "
                 "Largo Francesco Vito 1, 00168 Rome, Italy — "
                 "ORCID https://orcid.org/0000-0002-2154-319X")
KEYWORDS = ("regularization; weight decay; decoupled weight decay; language models; "
            "transformers; overfitting; robust regularization; Fair penalty; "
            "Huber decay; adaptive shrinkage")

# Figures embedded in the manuscript, in order of appearance (build_paper.py).
MANUSCRIPT_FIGURES = [
    'fig1_all_benchmarks.png',
    'fig2_gpt2_detail.png',
    'fig4_transformer_ci.png',
    'fig2b_scale_sweep.png',
    'mechanism_overfitting_delay_wt103.png',
    'fig5_wins_losses.png',
    'fig3_architectural_affinity.png',
    'mechanism_overfitting_delay_large.png',
    'mechanism_weight_dynamics_medium.png',
    'mechanism_adaptivity_medium.png',
]


# ── document helpers (same conventions as build_paper.py) ──────────────────
def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    return doc


def para(doc, text, bold=False, size=11, space_before=0, space_after=6,
         align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'


# ── highlights numbers, read from the regenerated CSVs ─────────────────────
def _tau_gap(bname, method):
    df = load_csv(bname)
    return get_val(df, method, 'test_ppl', 'mean') - get_val(df, 'τ(w)', 'test_ppl', 'mean')


def build_highlights():
    g7 = _tau_gap('gpt2_wikitext', 'WD-tuned')
    g66 = _tau_gap('gpt2_large_wikitext', 'WD-tuned')
    highlights = [
        "τ-decay: a one-line decoupled Fair weight decay applied after the optimizer step",
        f"Beats PSO-tuned AdamW weight decay by {g7:.1f} PPL at 7M and {g66:.1f} PPL at 66M",
        "The gain tracks overfitting pressure, not model size, across 2M–124M GPT-2",
        "A scope × adaptivity factorial ties the gain to decay scope and schedule",
        "Huber, pseudo-Huber and log-cosh decays match Fair: the profile is immaterial",
    ]
    for h in highlights:
        assert len(h) <= 85, f"Highlight over 85 characters ({len(h)}): {h!r}"
    assert 3 <= len(highlights) <= 5

    doc = new_doc()
    para(doc, "Highlights", bold=True, size=13, space_after=10)
    for h in highlights:
        bullet(doc, h)
    doc.save(os.path.join(OUT_DIR, '4_highlights.docx'))
    return highlights


def build_title_page():
    doc = new_doc()
    para(doc, "Title page", bold=True, size=13, space_after=12)
    para(doc, "Article type: Full-length article (original research)")
    para(doc, "Journal section: Learning Systems", space_after=12)
    para(doc, "Title:", bold=True)
    para(doc, TITLE, size=13, space_after=10)
    para(doc, f"Running title: {RUNNING_TITLE}", space_after=12)
    para(doc, "Authors:", bold=True)
    para(doc, AUTHORS_LINE)
    para(doc, "Affiliation:", bold=True, space_before=6)
    para(doc, AFFILIATION)
    para(doc, "Corresponding author:", bold=True, space_before=6)
    para(doc, CORRESPONDING, space_after=12)
    para(doc, "Keywords:", bold=True)
    para(doc, KEYWORDS, italic=True, space_after=12)
    para(doc, "Funding:", bold=True)
    para(doc, "This research did not receive any specific grant from funding agencies "
              "in the public, commercial, or not-for-profit sectors. The authors "
              "acknowledge the computational resources provided by RunPod and the "
              "support of Università Cattolica del Sacro Cuore.")
    para(doc, "Declarations of interest:", bold=True, space_before=6)
    para(doc, "None.")
    para(doc, "Data availability:", bold=True, space_before=6)
    para(doc, "Code and data are available at "
              "https://github.com/doctormaulux/Regularization_Decay.")
    doc.save(os.path.join(OUT_DIR, '2_title_page.docx'))


def build_credit():
    doc = new_doc()
    para(doc, "CRediT author statement", bold=True, size=13, space_after=12)
    para(doc, "Giuseppe Maulucci: Conceptualization, Methodology, Software, "
              "Investigation, Formal analysis, Data curation, Visualization, "
              "Writing – original draft, Writing – review & editing.")
    para(doc, "Tommaso Marchetti: Writing – review & editing.")
    para(doc, "Marco De Spirito: Writing – review & editing.")
    doc.save(os.path.join(OUT_DIR, '5_credit_author_statement.docx'))


def build_declaration():
    doc = new_doc()
    para(doc, "Declaration of interests", bold=True, size=13, space_after=12)
    para(doc, "☒ The authors declare that they have no known competing financial "
              "interests or personal relationships that could have appeared to "
              "influence the work reported in this paper.")
    para(doc, "☐ The authors declare the following financial interests/personal "
              "relationships which may be considered as potential competing "
              "interests: none.", space_before=6)
    doc.save(os.path.join(OUT_DIR, '6_declaration_of_interests.docx'))


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    fig_out = os.path.join(OUT_DIR, 'figures')
    os.makedirs(fig_out, exist_ok=True)

    manuscript_src = os.path.join(SCRIPT_DIR, 'articolo.docx')
    cover_src = os.path.join(SCRIPT_DIR, 'cover_letter.docx')
    shutil.copy2(cover_src, os.path.join(OUT_DIR, '1_cover_letter.docx'))
    shutil.copy2(manuscript_src, os.path.join(OUT_DIR, '3_manuscript.docx'))

    build_title_page()
    highlights = build_highlights()
    build_credit()
    build_declaration()

    for i, fname in enumerate(MANUSCRIPT_FIGURES, start=1):
        src = os.path.join(FIGURES_DIR, fname)
        shutil.copy2(src, os.path.join(fig_out, f'Figure_{i}_{fname}'))

    print(f"Submission folder assembled at: {OUT_DIR}")
    print("Highlights (length check passed):")
    for h in highlights:
        print(f"  [{len(h):2d}] {h}")
    for f in sorted(os.listdir(OUT_DIR)):
        print(" ", f)


if __name__ == '__main__':
    main()
