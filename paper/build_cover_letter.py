"""Generate the cover letter for the Neural Networks (Elsevier) submission."""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)


def add_para(text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    return p


# ── Sender block ──
add_para("Giuseppe Maulucci", bold=True)
add_para("Università Cattolica del Sacro Cuore")
add_para("Largo Francesco Vito 1, 00168 Rome, Italy")
add_para("Email: giuseppe.maulucci@unicatt.it")
add_para("ORCID: 0000-0002-2154-319X", space_after=12)

# ── Date ──
add_para(date.today().strftime("%d %B %Y"), space_after=12)

# ── Recipient ──
add_para("Prof. DeLiang Wang and Prof. Kenji Doya", bold=True)
add_para("Co-Editors-in-Chief")
add_para("Neural Networks")
add_para("Elsevier", space_after=18)

# ── Subject ──
add_para(
    "Re: Submission of original research article — "
    "“Decoupled Fair Weight Decay for Over-Capacity Autoregressive Language Models”",
    bold=True, space_after=18
)

# ── Body ──
add_para("Dear Editors,", space_after=12)

add_para(
    "We are pleased to submit the original research article cited above for consideration "
    "in Neural Networks. We would like the manuscript to be assigned to the Learning Systems "
    "section, which is the most appropriate fit for its scope."
)

add_para(
    "The paper studies τ-decay, a decoupled weight-decay update w ← w − ρ·w/(1 + |w|/δ) "
    "applied after the optimizer step to weight matrices only, at a rate independent of the "
    "learning-rate schedule; its implicit penalty is the Fair function of robust statistics, "
    "and its δ → ∞ limit is a constant decoupled decay with the same scope and schedule, which "
    "serves as the implementation-matched ablation. We compare it with seven regularizers "
    "(L1, L2, ElasticNet, SCAD, MCP, the log-sum penalty and PSO-tuned AdamW weight decay) on "
    "twelve benchmarks — regression, CNNs, a vision transformer, an encoder on SST-2, a "
    "pretrained 135M language model, a from-scratch GPT-2 scale sweep from 2M to 66M "
    "parameters and a 124M GPT-2 trained from scratch on WikiText-103 — under one protocol: "
    "per-method particle-swarm tuning with a dimension-matched budget, verified best-epoch "
    "early stopping, equivalence tests against pre-specified margins and, at the headline "
    "scale, ten seeds with paired tests."
)

add_para(
    "The finding is regime-specific and, we believe, of practical value. In from-scratch "
    "autoregressive language modelling under overfitting pressure, a weight-only, "
    "schedule-independent decoupled decay beats tuned AdamW weight decay by 2.4 perplexity "
    "points at 7M and 2.4 at 66M (9% and 10% below the unregularized baseline), by 1.4 "
    "points at 18M and not at all at 2M; what separates these cases is overfitting pressure "
    "rather than model size. A scope × adaptivity factorial at 66M attributes the margin to "
    "the decay's scope and schedule, and shows that magnitude adaptivity — the element that "
    "gives τ-decay its name — is a second-order, scale-dependent refinement that mainly "
    "compensates a wrong scope. Everywhere else the method degrades gracefully to parity "
    "with tuned weight decay, and the whole robust-decay family (Huber, pseudo-Huber, "
    "log-cosh) performs alike."
)

add_para(
    "We believe the work fits Neural Networks for three reasons. First, it isolates, with a "
    "factorial design and matched ablations, which properties of a decoupled decay matter for "
    "over-parameterized sequence models. Second, the empirical protocol is unusually strict for "
    "the field: every method is tuned with the same effort, every reported number is the "
    "verified best-epoch model, and claims of parity are tested as equivalence rather than "
    "inferred from non-significance. Third, the conclusion is stated with its limits — where "
    "the method helps, by how much, and where it does not — which we hope makes it useful to "
    "practitioners as well as to readers interested in the mechanics of regularization."
)

add_para(
    "We confirm that the manuscript reports original work; that it has not been published "
    "elsewhere; that it is not currently under consideration by another journal; and that all "
    "authors have approved the submitted version. Code and data are publicly available at "
    "https://github.com/doctormaulux/Regularization_Decay to support reproducibility."
)

add_para(
    "We declare no competing financial or personal interests."
)

add_para(
    "Thank you for considering this submission. We would be happy to address any "
    "questions or to provide supplementary material as needed.",
    space_after=24
)

add_para("Sincerely,", space_after=24)
add_para("Giuseppe Maulucci (corresponding author, on behalf of the authors: G. Maulucci, T. Marchetti, M. De Spirito)", bold=True)
add_para("Università Cattolica del Sacro Cuore")
add_para("giuseppe.maulucci@unicatt.it · ORCID 0000-0002-2154-319X")

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "cover_letter.docx")
doc.save(out)
print(f"[OK] Saved: {out}")
